"""Audit Word and HTML report consistency after generation."""
import re
import sys
import os

try:
    from docx import Document
except ImportError:
    print("python-docx not installed, skipping Word audit")
    Document = None


def audit_html(path):
    """Extract section headings and key content from HTML report."""
    with open(path, "r", encoding="utf-8") as f:
        html = f.read()
    headings = re.findall(r"<h2>(.*?)</h2>", html)
    return headings, html


def audit_word(path):
    """Extract section headings and key content from Word report."""
    if Document is None:
        return [], ""
    doc = Document(path)
    headings = []
    full_text = []
    for para in doc.paragraphs:
        if para.style and "Heading" in (para.style.name or ""):
            headings.append(para.text)
        full_text.append(para.text)
    # Also get table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return headings, "\n".join(full_text)


def main():
    tmpdir = sys.argv[1] if len(sys.argv) > 1 else None
    if tmpdir is None:
        # Try to find the latest report_test_* directory in temp
        import glob
        import tempfile
        pattern = os.path.join(tempfile.gettempdir(), "report_test_*")
        candidates = sorted(glob.glob(pattern), key=os.path.getmtime, reverse=True)
        if candidates:
            tmpdir = candidates[0]
        else:
            print("No report_test_* directory found in temp. Run test_report_gen.py first.")
            return 1
    html_path = os.path.join(tmpdir, "step3_report", "Report.html")
    docx_path = os.path.join(tmpdir, "step3_report", "Report.docx")

    html_headings, html_text = audit_html(html_path)
    word_headings, word_text = audit_word(docx_path)

    print("=== HTML Section Headings ===")
    for h in html_headings:
        print("  " + h)

    print("\n=== Word Section Headings ===")
    for h in word_headings:
        print("  " + h)

    # Check content presence
    content_checks = [
        ("Results table", "Outlet T"),
        ("Approach text", "SRK EOS"),
        ("Conclusions text", "safe operation"),
        ("Discussion: observation", "25.3"),
        ("Discussion: mechanism", "JT cooling"),
        ("Discussion: recommendation", "Proceed with current design"),
        ("Benchmark: source", "NIST"),
        ("Benchmark: PASS", "PASS"),
        ("Benchmark: FAIL", "FAIL"),
        ("Uncertainty: method", "Monte Carlo"),
        ("Uncertainty: P10", "-10"),
        ("Uncertainty: P50", "150"),
        ("Uncertainty: P90", "320"),
        ("Uncertainty: tornado", "Gas Price"),
        ("Risk: overall level", "Medium"),
        ("Risk: high risk item", "Gas price drop"),
        ("Risk: low risk item", "Corrosion"),
        ("Risk: mitigation", "Long-term contracts"),
        ("References", "Smith"),
    ]

    print("\n=== Content Consistency Check ===")
    print(f"{'Check':<35} {'HTML':>6} {'Word':>6}")
    print("-" * 50)
    mismatches = 0
    for label, needle in content_checks:
        in_html = needle in html_text
        in_word = needle in word_text
        match = "OK" if in_html == in_word else "DIFF"
        if match == "DIFF":
            mismatches += 1
        html_mark = "YES" if in_html else "NO"
        word_mark = "YES" if in_word else "NO"
        print(f"  {label:<33} {html_mark:>6} {word_mark:>6}  {match}")

    print(f"\n{mismatches} mismatches found")
    if mismatches > 0:
        return 1
    print("Word and HTML reports are consistent!")
    return 0


if __name__ == "__main__":
    sys.exit(main())
