"""
new_task.py - Create a new task-solving workspace and task folders.

This script lives in devtools/ (tracked in git) and is always available
after cloning. It auto-creates the task_solve/ folder structure on first run.

Usage:
    python devtools/new_task.py "JT cooling for rich gas"
    python devtools/new_task.py "TEG dehydration sizing" --type B
    python devtools/new_task.py "hydrate formation temperature" --type A --author "Your Name"
    python devtools/new_task.py --setup              # just create task_solve/ without a task
    python devtools/new_task.py --list               # list existing tasks
"""
import os
import shutil
import sys
from datetime import date


# ── Paths ────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
TASK_SOLVE_DIR = os.path.join(PROJECT_ROOT, "task_solve")
TEMPLATE_DIR = os.path.join(TASK_SOLVE_DIR, "TASK_TEMPLATE")

TASK_TYPES = {
    "A": "Property",
    "B": "Process",
    "C": "PVT",
    "D": "Standards",
    "E": "Feature",
    "F": "Design",
    "G": "Workflow",
}


# ══════════════════════════════════════════════════════════
# Embedded templates — these are the "source of truth" so
# new users get them even though task_solve/ is gitignored.
# ══════════════════════════════════════════════════════════

WORKSPACE_README = r"""# AI-Supported Task Solving While Developing

This folder is a **local working area** for solving engineering tasks using the
3-step AI-assisted workflow. It is in `.gitignore` — nothing here is committed.
Each task gets its own subfolder with scope & research notes, simulation results,
figures, and reports.

## Quick Start

Open VS Code Copilot Chat and type:

```
@solve.task JT cooling for rich gas at 100 bara
```

That's it. The agent creates the folder, researches the topic, builds and runs
a simulation, validates the results, and generates reports (Word + HTML) — all
in one session. You solve advanced engineering tasks while simultaneously
improving the NeqSim toolbox.

> **Alternative:** If you prefer a manual step-by-step approach, run
> `python devtools/new_task.py "your task"` and follow the prompts in the
> generated README.

---

## Prerequisites

| Requirement | Install / Setup | Used In |
|-------------|----------------|---------|
| Python 3.8+ | [python.org](https://python.org) | All steps |
| Java JDK 8+ | Bundled via `devtools/` setup | Step 2 simulations |
| NeqSim dev tools | `pip install -e devtools/` (from repo root) | Step 2 — boots the JVM and gives you `neqsim_dev_setup` |
| VS Code + GitHub Copilot Chat | VS Code marketplace | All steps |
| python-docx | `pip install python-docx matplotlib` | Step 3 reports |
| Google NotebookLM (optional) | [notebooklm.google.com](https://notebooklm.google.com) | Step 1 research (or use Copilot — see below) |

> **Important:** For the task-solving workflow use `pip install -e devtools/` — this
> installs the `neqsim_dev_setup` helper that boots the JVM from your local build.
> Do **not** use `pip install neqsim` here (that installs the released package, not
> your working copy).

**Quick setup (one-time):**

```powershell
cd path/to/neqsim
pip install -e devtools/          # installs neqsim_dev_setup for Jupyter
pip install python-docx matplotlib # for Word reports and plots
```

---

## Who Is This For?

### Path A: Process Engineer (I just want answers)

You have an engineering question — "What's the hydrate temperature for this
gas?" or "Size a 3-stage compressor train." You don't want to learn Java or git.

1. Open VS Code Copilot Chat
2. Type: `@solve.task hydrate temperature for wet gas at 100 bara`
3. The agent creates a folder, runs the simulation, and gives you results
4. Find reports in `task_solve/.../step3_report/`

### Path B: Developer (I want to extend NeqSim)

You're solving a task AND improving the NeqSim codebase. When the API is
missing something, you add it mid-task — new methods, equipment, or models.

1. Type: `@solve.task add JT coefficient method` (or run `python devtools/new_task.py` for manual control)
2. Work through all 3 steps — the agent flags API gaps as it goes
3. Add the missing Java code, rebuild, and the notebook picks it up immediately
4. Promote reusable code back into `src/main/`, `src/test/`, or `examples/`

### Path C: Researcher (I need a technical assessment)

You're producing a deliverable — a report, technology screening, or design
study. The 3 steps map directly to a professional workflow.

1. Type: `@solve.task field development concept selection for deepwater gas`
2. Review and refine the scope and research notes the agent produces
3. Iterate on analysis — the agent refines until results validate
4. Run report generation to produce Word + HTML deliverables

### Path D: Other AI Tools (OpenAI Codex, Claude Code, Cursor, etc.)

The workflow is **not tied to VS Code Copilot Chat**. The script, folder
structure, templates, and report generator work from any terminal. Any AI
coding agent that can read files and run commands can drive the workflow.

**How to start a task from OpenAI Codex (or any AI agent):**

1. Run the setup: `python devtools/new_task.py "your task" --type B`
2. Point the agent to the guide:
   ```
   Read docs/development/TASK_SOLVING_GUIDE.md for the full workflow.
   Read the task README at task_solve/YYYY-MM-DD_your_task/README.md.
   Follow the 3-step workflow: fill task_spec.md, create a notebook
   in step2_analysis/, then run step3_report/generate_report.py.
   ```
3. The AI agent reads the templates, fills in the task spec, creates notebooks,
   and runs the report generator — same output, different tool.

**What works everywhere** (no VS Code required):
- `python devtools/new_task.py` — creates task folders
- `task_spec.md` — scope document (plain markdown)
- Jupyter notebooks — work in any Python environment
- `python step3_report/generate_report.py` — generates Word + HTML
- `git` + `gh pr create` — contributing back via PR

**What's VS Code Copilot-specific** (optional convenience):
- `@solve.task` agent — automates the full workflow (the agent reads
  `.github/agents/solve.task.agent.md` for its instructions)
- Specialist agents (`@thermo.fluid`, `@solve.process`, etc.)

---

## Common Task Examples

| Type | Example Tasks |
|------|--------------|
| **A - Property** | Density of CO2 at 200 bar; viscosity of MEG-water; JT coefficient for rich gas |
| **B - Process** | TEG dehydration unit; 3-stage compression; HP/LP separation train |
| **C - PVT** | CME test for reservoir oil; CCE at 100C; swelling test with CO2 injection |
| **D - Standards** | Wobbe index per ISO 6976; hydrocarbon dew point; AGA flow measurement |
| **E - Feature** | Add anti-surge to compressor; fix CPA solver for CO2-water; new property method |
| **F - Design** | Pipeline wall thickness per DNV; separator mechanical design; PSV sizing |
| **G - Workflow** | Field development concept selection; technology screening; design basis |

---

## Adaptive Complexity — One Workflow, Any Scale

The framework adapts automatically. You don't need to configure anything — the
agent scales its depth based on what you ask for:

| Scale | Example | What Happens |
|-------|---------|-------------|
| **Quick** | "density of CO2 at 200 bar" | Minimal task_spec, one notebook cell, brief summary — done in minutes |
| **Standard** | "TEG dehydration for 50 MMSCFD" | Full task_spec, complete notebook, Word + HTML reports |
| **Comprehensive** | "field development concept selection per NORSOK" | Detailed task_spec with all standards, multiple notebooks per discipline, full HTML report with navigation |

**The same `@solve.task` command handles all of these.** The agent reads your
request and decides how deep to go. Specify standards ("per DNV-OS-F101") and
deliverables ("with sensitivity analysis and cost estimate") to guide depth.

### Guiding the Analysis

You control the scope through your request — the more you specify, the deeper
the analysis. Compare:

- **Simple:** `@solve.task hydrate temperature for wet gas at 100 bara`
  → Quick calculation, one-page result

- **Medium:** `@solve.task hydrate temperature for wet gas at 100 bara, per NORSOK P-001, with inhibitor dosing curve`
  → Standard analysis with standards compliance and sensitivity plot

- **Full study:** `@solve.task field development flow assurance assessment per NORSOK P-001 and DNV-RP-F109, covering hydrate, wax, corrosion, and slugging for 50 km subsea tieback, deliver phase envelopes, inhibitor curves, pipeline profiles, and design basis document`
  → Multi-notebook comprehensive study with full deliverable set

---

## The 3-Step Workflow

```
 STEP 1                    STEP 2                    STEP 3
 Scope & Research          Analysis & Evaluation     Report

 Define standards,         Build simulation,         Word + HTML
 methods, deliverables     run, validate, iterate    deliverables

 Google NotebookLM         NeqSim API +              python-docx
  or Copilot               GitHub Copilot            + HTML template
 + open sources
                           Iteration is implicit:
 Build knowledge base      refine until validated
```

### Step 1: Scope & Research

This step has two parts: **define the scope** (what to do) and **research** (gather background).

#### Part A: Task Specification (scope)

Before any analysis, define what governs the work:

- **Applicable standards**: Which codes and standards apply? (e.g., NORSOK P-001,
  ISO 6976, DNV-OS-F101, API 520, ASME B31.3, company TR documents)
- **Calculation methods/models**: Which EOS, correlations, or pipe flow models
  to use? (e.g., SRK-CPA for polar systems, Beggs & Brill for multiphase flow,
  OLGA-style thermal-hydraulic)
- **Required deliverables**: What must the final output include? (e.g., phase
  envelopes, pressure profiles, sizing calculations, sensitivity plots, VFP tables)
- **Acceptance criteria**: Tolerances, design factors, safety margins, convergence
  targets (e.g., mass balance < 0.1%, design factor per DNV = 0.72)
- **Operating envelope**: Range of conditions to cover (pressures, temperatures,
  flow rates, compositions)

Fill in `step1_scope_and_research/task_spec.md` — this is the "brief" that
guides everything in Step 2.

#### Part B: Research (background knowledge)

**Providing literature papers and reference documents:**

If you have PDF papers, standards documents, lab reports, or other background
material, place them in the `step1_scope_and_research/references/` folder.
Then summarise their key contributions in `notes.md` under "Literature &
Reference Documents". See `references/README.md` for naming conventions and
tips on how the AI can use these files.

Use **either** Google NotebookLM or Copilot — or both:

**Option A: Google NotebookLM** (best for deep literature review)
- Upload PDFs from `step1_scope_and_research/references/` to NotebookLM
- Ask questions across all your sources at once
- Get cited answers with references back to source pages
- Paste findings into `notes.md`

**Option B: GitHub Copilot in VS Code** (best for code-adjacent research)
- Open Copilot Chat and ask research questions directly
- Copilot can search the web, read repo docs, and summarise findings
- For PDFs, extract key sections as text or summarise them in `notes.md`

**Copilot research workflow:**

1. Open VS Code Copilot Chat (Ctrl+Shift+I)
2. Paste this prompt:
   ```
   I'm researching [TOPIC] for a NeqSim task.
   Search the web and this repository for:
   1. Key physical principles and governing equations
   2. Typical operating ranges and design rules of thumb
   3. Relevant industry standards (API, ASME, ISO, NORSOK, DNV)
   4. What NeqSim classes/methods already exist for this
   Write the findings to step1_scope_and_research/notes.md in my task folder.
   ```
3. Review and refine — ask follow-up questions
4. Save the final notes in `step1_scope_and_research/`

### Step 2: Analysis & Evaluation

This step combines building, running, and validating the simulation in one
iterative flow. You don't need to separate "analysis" from "evaluation" — it's
a natural loop:

1. Build the simulation (notebook or Java test)
2. Run it and inspect results
3. Check physics (mass/energy balance, reasonable ranges)
4. Compare against reference data from Step 1
5. If results are off → adjust and rerun (iteration is implicit)
6. When satisfied → save final results and figures

All simulation code, results, and validation notes go to `step2_analysis/`.

### Step 3: Report (Word + HTML)

The deliverables are a **Word report** (`.docx`) and optionally an **HTML report**.

- Word report: formal document for sharing/review, generated via `generate_report.py`
- HTML report: interactive, navigable document — ideal for large workflows with
  many sections, embedded plots, and linked references
- Run: `python step3_report/generate_report.py`
- Both formats embed all figures from `figures/` directory

---

## VS Code Agent Quick Reference

| Agent | Best For | Example |
|-------|----------|---------|
| `@solve.task` | **Full 3-step workflow** (does everything) | "JT cooling for rich gas at 100 bara" |
| `@thermo.fluid` | Fluid setup, EOS, flash, properties | "Density of CO2-methane mix at 200 bar" |
| `@solve.process` | Complete simulation -> notebook | "TEG dehydration for 50 MMSCFD" |
| `@pvt.simulation` | PVT lab experiments | "CME test at 100C for this oil" |
| `@gas.quality` | Gas standards (ISO, GPA) | "Wobbe index per ISO 6976" |
| `@mechanical.design` | Wall thickness, sizing | "20-inch pipe per DNV-OS-F101" |
| `@flow.assurance` | Hydrates, wax, corrosion | "Hydrate curve for wet gas at 100 bara" |
| `@safety.depressuring` | Blowdown, PSV, fire | "Fire-case blowdown for HP separator" |

---

## Contributing Back

### Minimum (everyone should do this)
Add a task log entry to `docs/development/TASK_LOG.md`.

### Medium (if you wrote a useful notebook)
Copy your notebook to `examples/notebooks/`.

### Full (if you extended the API)
Write a test in `src/test/java/neqsim/` and run `mvnw.cmd test`.

### Create a Pull Request (if you want to contribute code or examples)

When your task produces reusable outputs — new methods, tests, notebooks, or
documentation — create a PR directly from the task:

```powershell
# 1. Create a branch from your current branch
git checkout -b task/your-task-name

# 2. Stage the files you want to contribute
git add src/test/java/neqsim/...             # tests
git add examples/notebooks/your_notebook.ipynb # notebook
git add docs/development/TASK_LOG.md          # task log entry

# 3. Commit and push
git commit -m "Add [description] from task solving workflow"
git push -u origin task/your-task-name

# 4. Create the PR (requires GitHub CLI: https://cli.github.com/)
gh pr create --title "Add [description]" --body "From task: [task title]"
```

> **Tip:** The `@solve.task` agent can do this for you — just ask
> "create a PR with the reusable outputs from this task".

---

## Troubleshooting

| Problem | Solution |
|---------|----------|
| `python` not found | Use `python3` or install Python from python.org |
| Copilot Chat doesn't know NeqSim | Start prompt with "Read CONTEXT.md for orientation" |
| Simulation gives wrong numbers | Check EOS choice, mixing rule, units (Kelvin vs Celsius!) |
| Want to share task with colleague | Zip the task folder and send it - it's self-contained |

---

## Large Workflows (e.g., Field Development, Class A Studies)

For complex tasks that span multiple engineering disciplines (field development
concept selection, design basis studies, technology screening, Class A/B
estimates), the framework scales naturally:

1. **Step 1 (Scope)** becomes critical — define ALL standards, methods, and
   deliverables upfront in `task_spec.md`. For Class A studies, this may
   reference 10+ standards and produce a detailed work breakdown.
2. **Step 2 (Analysis)** can contain multiple notebooks, each covering a
   sub-analysis (e.g., `01_reservoir_fluid.ipynb`, `02_pipeline_sizing.ipynb`,
   `03_process_train.ipynb`, `04_flow_assurance.ipynb`, `05_cost_estimation.ipynb`)
3. **Step 3 (Report)** produces a comprehensive HTML document with navigation
   sidebar, linking all sub-analyses — plus a Word summary for formal distribution

The task type **G (Workflow)** is intended for these multi-discipline studies.

### Example: Field Development Concept Selection

```
@solve.task field development concept selection for 200 MMSCFD deepwater gas,
  per NORSOK P-001, Z-013, L-001, and DNV-OS-F101.
  Evaluate subsea tieback vs. FPSO vs. fixed platform.
  Deliver: reservoir fluid characterization, process train sizing,
  pipeline hydraulics, flow assurance assessment (hydrate + wax + corrosion),
  mechanical design summary, CAPEX/OPEX ranking, and recommendation report.
```

This generates 5-8 notebooks, a full task_spec.md referencing all standards,
and a navigable HTML report — a complete engineering study.
"""

TASK_README = r"""# Task: [Title]

**Date:** YYYY-MM-DD
**Type:** A (Property) | B (Process) | C (PVT) | D (Standards) | E (Feature) | F (Design) | G (Workflow)
**Status:** In Progress | Complete

## Problem Statement

[Describe the engineering question or task]

---

## Step 1: Scope & Research

### Part A: Task Specification

Fill in `step1_scope_and_research/task_spec.md` before starting analysis.

- [ ] Applicable standards defined
- [ ] Calculation methods/models specified
- [ ] Required deliverables listed
- [ ] Acceptance criteria set
- [ ] Operating envelope defined

### Part B: Research

- [ ] Literature search completed
- [ ] Reference documents placed in `step1_scope_and_research/references/`
- [ ] Reference data collected
- [ ] Key sources documented in `step1_scope_and_research/notes.md`

**Providing literature papers and reference documents:**

If you have PDF papers, standards documents, lab reports, or other background
material, place them in the `step1_scope_and_research/references/` folder.
Then summarise their key contributions in `notes.md` under "Literature &
Reference Documents". See `references/README.md` for naming conventions and
tips on how the AI can use these files.

**Option A — Google NotebookLM** (upload PDFs, get cited answers):

```
I need to understand [TOPIC] for oil & gas process engineering.
Give me:
1. Key physical principles and governing equations
2. Typical operating ranges and design rules of thumb
3. Relevant industry standards (API, ASME, ISO, NORSOK, DNV)
4. Common correlations used in practice
5. Known limitations or edge cases
```

**Option B — VS Code Copilot Chat** (Ctrl+Shift+I, web search + repo context):

```
I'm researching [TOPIC] for a NeqSim task.
Read the task specification in step1_scope_and_research/task_spec.md first.
Search the web and this repository for:
1. Key physical principles and governing equations
2. Requirements from the specified standards
3. What NeqSim classes/methods already exist for this
Write the findings to step1_scope_and_research/notes.md in my task folder.
```

---

## Step 2: Analysis & Evaluation

- [ ] NeqSim simulation written (notebook or test)
- [ ] Results extracted and saved
- [ ] Results validated against references and acceptance criteria
- [ ] Physics checks passed (mass/energy balance, ranges)
- [ ] API gaps identified (if any)

The analysis and evaluation happen in one iterative loop — build, run, validate,
refine until results meet the acceptance criteria from Step 1.

**AI prompt - paste into VS Code Copilot Chat:**

```
I'm working on a task in task_solve/[THIS_FOLDER]/.
Read the task specification in step1_scope_and_research/task_spec.md.
Read the research notes in step1_scope_and_research/notes.md.

Task: [DESCRIBE YOUR TASK]

Create a Jupyter notebook in step2_analysis/ that:
1. Sets up the fluid system with appropriate EOS (per task spec)
2. Builds the process flowsheet
3. Runs the simulation
4. Validates results against acceptance criteria
5. Extracts key results and saves figures to figures/
```

**Which VS Code agent to use:**

| Task Type | Agent | Example prompt |
|-----------|-------|----------------|
| Fluid properties | `@thermo.fluid` | "Create a CPA fluid for gas with 5% MEG" |
| Process simulation | `@solve.process` | "3-stage compression from 5 to 150 bara" |
| PVT study | `@pvt.simulation` | "CME test for reservoir fluid at 100C" |
| Gas quality | `@gas.quality` | "Wobbe index per ISO 6976 for this gas" |
| Mechanical design | `@mechanical.design` | "Wall thickness for 20-inch pipe per DNV" |
| Flow assurance | `@flow.assurance` | "Hydrate formation curve for wet gas" |
| Safety | `@safety.depressuring` | "Fire-case blowdown for HP separator" |

---

## Step 3: Report (Word + HTML)

The deliverables are a **Word report** (`.docx`) and optionally an **HTML report**.

- [ ] Figures saved to `figures/`
- [ ] `generate_report.py` customized and runs end-to-end
- [ ] Report(s) generated in `step3_report/`
- [ ] All required deliverables from task spec produced
- [ ] Task logged in `docs/development/TASK_LOG.md`

**To generate reports:**

```powershell
pip install python-docx matplotlib    # one-time setup
python step3_report/generate_report.py
```

**AI prompt - paste into VS Code Copilot Chat:**

```
Customize step3_report/generate_report.py for this task.
Read the task spec in step1_scope_and_research/task_spec.md for required deliverables.
Fill in the report sections with actual results from step2_analysis/.
Embed all figures from figures/.
Generate both Word (.docx) and HTML output.
```

---

## Key Results

[Summary of findings - fill in when complete]

---

## Saving Results for the Report (results.json)

The report generator (`step3_report/generate_report.py`) auto-reads a `results.json`
file from the task root. Add a cell at the end of your notebook to save results:

```python
import json, os, pathlib

# Resolve task directory from the notebook's own location
# (os.getcwd() is unreliable in VS Code notebooks — it returns workspace root)
NOTEBOOK_DIR = pathlib.Path(globals().get(
    "__vsc_ipynb_file__", os.path.abspath("step2_analysis/notebook.ipynb")
)).resolve().parent
TASK_DIR = NOTEBOOK_DIR.parent
FIGURES_DIR = TASK_DIR / "figures"
FIGURES_DIR.mkdir(exist_ok=True)

results = {
    "key_results": {
        # Add your key numerical results here (units in the key name)
        "outlet_temperature_C": 25.3,
        "pressure_drop_bar": 5.2,
        "power_kW": 1250.0,
    },
    "validation": {
        # Each check: True = pass, False = fail, or a numeric value
        "mass_balance_error_pct": 0.01,
        "energy_balance_error_pct": 0.5,
        "temperature_in_range": True,
        "pressure_positive": True,
        "acceptance_criteria_met": True,
    },
    "approach": "Used SRK EOS with classic mixing rule. Process: ...",
    "conclusions": "The analysis shows that ...",
    # Optional: custom figure captions (map filename -> caption text)
    "figure_captions": {
        # "my_plot.png": "Temperature and pressure profiles during simulation",
    },
    # Optional: key equations used in the analysis (rendered in reports)
    "equations": [
        # {"label": "Energy Balance", "latex": "Q = m C_p \\\\Delta T"},
    ],
    # Optional: custom tables (rendered in both Word and HTML)
    "tables": [
        # {"title": "Sensitivity Analysis",
        #  "headers": ["Parameter", "Base Case", "Low", "High"],
        #  "rows": [["Pressure (bar)", 60.0, 40.0, 80.0],
        #           ["Temperature (C)", 25.0, 15.0, 35.0]]}
    ],
    # Optional: references (rendered as numbered list in References section)
    "references": [
        # {"id": "Smith2019", "text": "Smith, J. (2019). CNG Tank Thermal Analysis. J. Energy Storage, 25, 100-115."},
        # {"id": "API521", "text": "API 521, 7th Edition (2020). Pressure-Relieving and Depressuring Systems."},
    ],
}

results_path = str(TASK_DIR / "results.json")
with open(results_path, "w") as f:
    json.dump(results, f, indent=2)
print(f"Results saved to {results_path}")
```

**Saving figures** — use `FIGURES_DIR` (set above) so plots end up in the right place:

```python
import matplotlib.pyplot as plt

fig, ax = plt.subplots()
ax.plot(x, y)
ax.set_xlabel("X Label")
ax.set_ylabel("Y Label")
fig.savefig(str(FIGURES_DIR / "my_plot.png"), dpi=150, bbox_inches="tight")
plt.show()
```

When you run `python step3_report/generate_report.py`, the Results and Validation
sections are auto-populated from this file. Figures in `figures/` are also embedded.
Equations from results.json are rendered with KaTeX in HTML and as images in Word.

---

## Contribute Back - Reusable Outputs

When your task is done, promote valuable work back into the repo so others
benefit. Check what applies:

- [ ] **Test** - Copy simulation to `src/test/java/neqsim/` (proves it keeps working)
- [ ] **Notebook** - Copy notebook to `examples/notebooks/` (others can rerun it)
- [ ] **API extension** - New methods added to `src/main/java/neqsim/`
- [ ] **Documentation** - Guide or recipe added to `docs/`
- [ ] **Task log** - Entry added to `docs/development/TASK_LOG.md`

Don't worry if you can't do all of these. Even just the task log entry helps
the next person (or AI session) find your solution.

### Create a Pull Request

If this task produced reusable code, tests, or notebooks, create a PR to
contribute them back:

```powershell
# Create a feature branch
git checkout -b task/[SHORT_NAME]

# Stage reusable outputs (pick what applies)
git add src/test/java/neqsim/...              # new tests
git add examples/notebooks/...                 # example notebooks
git add docs/...                               # documentation
git add docs/development/TASK_LOG.md           # task log entry

# Commit and push
git commit -m "Add [description] from task: [TITLE]"
git push -u origin task/[SHORT_NAME]

# Create PR (requires GitHub CLI)
gh pr create --title "Add [description]" \\
  --body "From task-solving workflow: [TITLE]"
```

> **Tip:** Ask the `@solve.task` agent to do this:
> "create a PR with the test and notebook from this task"
"""

STEP1_NOTES = """# Step 1: Research Notes

## Sources

| # | Source | Type | Key Finding |
|---|--------|------|-------------|
| 1 | | | |

## Literature & Reference Documents

Place PDF papers, standards, and technical reports in the `references/` folder
next to this file. Then summarise each document's key contributions here.

**How to add a reference:**
1. Copy the PDF/document to `step1_scope_and_research/references/`
2. Add a row to the Sources table above
3. Write a brief summary below noting the relevant equations, data, or design rules

### Paper/Document Summaries

<!--
For each reference document, add a subsection like this:

### Smith (2019) — CNG Tank Thermal Analysis
- **File:** `references/Smith_2019_CNG_Tank_Thermal_Analysis.pdf`
- **Relevance:** Provides heat transfer correlations for compressed gas filling
- **Key equations:** Eq. 12 — convective HTC = 15-25 W/m2K for turbulent fill
- **Key data:** Table 3 — experimental fill temperatures vs. time
- **Limitations:** Only covers Type III tanks, ambient temperature 20C
-->

## Background

[Summary of the engineering context]

## Key Data / Correlations

[Reference values, correlations, experimental data]

## Open Questions

- [ ]
"""

REFERENCES_README = """# References Folder

Place literature papers, standards documents, and other reference material here.

## What to put in this folder

- **PDF papers** -- journal articles, conference papers, technical reports
- **Standards excerpts** -- relevant sections from ASME, API, DNV, ISO, NORSOK, etc.
- **Company documents** -- TR documents, design basis, operating philosophy
- **Data sheets** -- equipment data sheets, material certificates
- **Lab reports** -- PVT reports, fluid analysis, corrosion test results

## How the AI uses these files

1. **Google NotebookLM (recommended for PDFs):** Upload the PDFs from this
   folder to NotebookLM. It can read, cross-reference, and cite multiple
   documents at once. Ask it targeted questions and paste the answers into
   `notes.md`.

2. **VS Code Copilot Chat:** Copilot can read text-based files (`.txt`, `.md`,
   `.csv`) placed here. For PDFs, extract key sections as text or summarise
   them in `notes.md` so the AI can reference the content.

3. **Manual notes:** Read the papers yourself and capture key equations,
   data points, and design rules in `notes.md` under the "Literature &
   Reference Documents" section.

## Naming convention

Use descriptive filenames that include author/org and year:

```
Smith_2019_CNG_Tank_Thermal_Analysis.pdf
API_521_6th_Ed_Relief_Systems.pdf
DNV-ST-F101_2021_Submarine_Pipelines.pdf
Equinor_TR2000_Pressure_Vessel_Design.pdf
Lab_Report_Fluid_Analysis_2024.pdf
```

## Documenting references

After placing files here, add each to the Sources table in `notes.md`:

| # | Source | Type | Key Finding |
|---|--------|------|-------------|
| 1 | Smith (2019) -- see `references/Smith_2019_CNG_Tank_...pdf` | Paper | Heat transfer coefficient 15-25 W/m2K |
| 2 | API 521 6th Ed | Standard | Relief sizing per Section 5.4 |

And add structured entries to the `references` list in `results.json` so they
appear in the final report (see `results.json` schema in AGENTS.md).
"""

TASK_SPEC = """# Task Specification

This file defines the scope and requirements that guide the analysis in Step 2.
Fill this in before starting any simulation work.

## Applicable Standards

List the codes, standards, and company requirements that govern this task.

| Standard | Scope | Key Requirements |
|----------|-------|-----------------|
| | | |

Examples: NORSOK P-001, ISO 6976, DNV-OS-F101, API 520, ASME B31.3, Equinor TR1414

## Calculation Methods & Models

Specify which methods, equations of state, and correlations to use.

- **Equation of State:** [e.g., SRK, PR, SRK-CPA for polar systems]
- **Pipe flow model:** [e.g., Beggs & Brill, OLGA-style, single-phase]
- **Heat transfer:** [e.g., adiabatic, U-value based, ambient loss]
- **Other correlations:** [e.g., API 520 for PSV sizing, NORSOK M-001 for materials]

## Required Deliverables

What must the final output include? Check all that apply and add specifics.

- [ ] Phase envelope / phase diagram
- [ ] Pressure-temperature profiles
- [ ] Equipment sizing calculations
- [ ] Sensitivity analysis (specify parameters)
- [ ] Comparison with reference/experimental data
- [ ] VFP tables
- [ ] Material selection
- [ ] Cost estimation
- [ ] Other: [specify]

## Acceptance Criteria

Define what \"good enough\" means for this task.

- **Mass balance tolerance:** [e.g., < 0.1%]
- **Energy balance tolerance:** [e.g., < 1%]
- **Design factor:** [e.g., 0.72 per DNV]
- **Safety margin:** [e.g., 10% on design pressure]
- **Convergence:** [e.g., solver residual < 1e-6]
- **Other:** [specify]

## Operating Envelope

Define the range of conditions to be covered.

| Parameter | Min | Design | Max | Unit |
|-----------|-----|--------|-----|------|
| Pressure | | | | bara |
| Temperature | | | | C |
| Flow rate | | | | kg/hr |
| Composition | | | | mol% |

## Input Data

Reference any input data files, lab reports, or composition tables.

- [ ] Fluid composition: [source]
- [ ] Operating conditions: [source]
- [ ] Equipment data: [source]
- [ ] Literature papers: [place PDFs in `references/` folder, summarise in `notes.md`]
- [ ] Other: [specify]

## Reference Fluid Compositions

Pick a starting composition or define your own. All values in mol%.

### Typical Lean Pipeline Gas

| Component | mol% |
|-----------|------|
| methane | 85.0 |
| ethane | 7.0 |
| propane | 3.0 |
| i-butane | 0.5 |
| n-butane | 0.5 |
| i-pentane | 0.1 |
| n-pentane | 0.1 |
| nitrogen | 1.5 |
| CO2 | 2.3 |

### Typical Rich Gas (with condensate potential)

| Component | mol% |
|-----------|------|
| methane | 75.0 |
| ethane | 10.0 |
| propane | 5.0 |
| i-butane | 1.5 |
| n-butane | 2.0 |
| i-pentane | 0.5 |
| n-pentane | 0.5 |
| n-hexane | 0.3 |
| nitrogen | 1.0 |
| CO2 | 4.2 |

### Typical Wet Gas (with water for hydrate/dehydration studies)

| Component | mol% |
|-----------|------|
| methane | 80.0 |
| ethane | 6.0 |
| propane | 3.0 |
| n-butane | 1.0 |
| CO2 | 2.0 |
| nitrogen | 1.0 |
| water | 7.0 |

### Typical CO2-Rich Stream (CCS)

| Component | mol% |
|-----------|------|
| CO2 | 95.0 |
| nitrogen | 2.0 |
| methane | 1.5 |
| water | 1.0 |
| H2S | 0.5 |

> **Note:** Adapt these to your specific project data. For oil systems with
> C7+ fractions, use the `@thermo.fluid` agent or define TBP/plus fractions.
"""

STEP2_NOTES = """# Step 2: Analysis & Validation Notes

## Analysis Log

### Run 1 - YYYY-MM-DD

**Setup:**

**Results:**

**Validation against acceptance criteria:**

**Status:** Pass / Needs refinement

---

## Validation Summary

Fill this table as you validate each check. This maps directly to the
`validation` section in `results.json` and auto-populates the report.

| Check | Status | Value / Note |
|-------|--------|--------------|
| Mass balance (in = out +/- tolerance) | | |
| Energy balance | | |
| Temperatures in reasonable range | | |
| Pressures positive | | |
| Densities in expected range | | |
| Results consistent with literature | | |
| Acceptance criteria from task_spec met | | |
| Sensitivity to key parameters checked | | |
| All deliverables from task_spec produced | | |

## Sensitivity Analysis (if applicable)

Document any parameter sweeps performed:

| Parameter Varied | Range | Effect on Output | Conclusion |
|------------------|-------|------------------|------------|
| | | | |

## Comparison with Reference Data (if applicable)

| Source | Parameter | Reference Value | NeqSim Value | Deviation |
|--------|-----------|----------------|--------------|----------|
| | | | | |

## results.json Status

- [ ] results.json saved from notebook (see task README for pattern)
- [ ] key_results section populated with all key outputs
- [ ] validation section populated with all checks above
- [ ] approach and conclusions fields filled in
- [ ] figure_captions populated with custom captions for key plots
- [ ] equations populated with key equations used in the analysis
- [ ] Figures saved to figures/ using absolute TASK_DIR path (not os.getcwd!)
"""

GENERATE_REPORT = '''"""
generate_report.py - Generate Word and HTML reports for this task.

Usage:
    pip install python-docx matplotlib   (one-time setup)
    python step3_report/generate_report.py

This script AUTO-READS data from the task folder:
  - step1_scope_and_research/task_spec.md  -> populates Scope & Standards
  - results.json (task root)               -> populates Results + Validation
  - figures/*.png                          -> embeds all plots
  - results.json "equations"               -> renders equations (KaTeX/images)
  - results.json "figure_captions"         -> custom captions for figures

It produces:
  - step3_report/Report.docx  (Word document for formal distribution)
  - step3_report/Report.html  (navigable HTML with sidebar, KaTeX equations)

If results.json or task_spec.md are missing, the report uses placeholder text.
Customize MANUAL_SECTIONS below for content that can't be auto-generated.
"""
import os
import sys
import glob
import json
import base64
import io
from datetime import date

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# Optional: matplotlib for rendering equations to images (Word report)
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

# ── Paths ────────────────────────────────────────────────
TASK_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(TASK_DIR, "figures")
REPORT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_FILE = os.path.join(REPORT_DIR, "Report.docx")
HTML_FILE = os.path.join(REPORT_DIR, "Report.html")
RESULTS_FILE = os.path.join(TASK_DIR, "results.json")
TASK_SPEC_FILE = os.path.join(TASK_DIR, "step1_scope_and_research", "task_spec.md")

# ── Configuration (edit these) ───────────────────────────
TITLE = "Task Report"           # <-- Change to your task title
AUTHOR = ""                     # <-- Your name
TASK_DATE = date.today().isoformat()

# ── Manual sections (edit content for your specific task) ─
# These are used when auto-read data is not available.
# If results.json exists, sections 5-6 are auto-populated.
MANUAL_SECTIONS = {
    "executive_summary": (
        "[Replace with a 3-5 sentence summary of the task, approach, "
        "and key findings.]"
    ),
    "problem_description": (
        "[Describe the engineering question or task that was solved.]"
    ),
    "approach": (
        "[Describe the methodology: EOS used, process configuration, "
        "simulation setup, key assumptions.]"
    ),
    "conclusions": (
        "[Summarize key findings and provide recommendations.]"
    ),
    "references": (
        "[List references from step1_scope_and_research/notes.md.]"
    ),
}


# ══════════════════════════════════════════════════════════
# Auto-read functions
# ══════════════════════════════════════════════════════════

def load_results():
    """Load results.json if it exists. Returns dict or None."""
    if os.path.exists(RESULTS_FILE):
        with open(RESULTS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        print("  Loaded results.json ({} keys)".format(len(data)))
        return data
    print("  No results.json found (using manual sections)")
    return None


def load_task_spec():
    """Load task_spec.md and extract standards/methods/criteria sections."""
    if not os.path.exists(TASK_SPEC_FILE):
        print("  No task_spec.md found (using placeholder for scope)")
        return None
    with open(TASK_SPEC_FILE, "r", encoding="utf-8") as f:
        content = f.read()
    print("  Loaded task_spec.md ({} chars)".format(len(content)))
    return content


def extract_spec_section(spec_text, heading):
    """Extract a section from task_spec.md by heading."""
    if not spec_text:
        return ""
    lines = spec_text.split("\\n")
    capturing = False
    result = []
    for line in lines:
        if line.startswith("## ") and heading.lower() in line.lower():
            capturing = True
            continue
        elif line.startswith("## ") and capturing:
            break
        elif capturing:
            result.append(line)
    text = "\\n".join(result).strip()
    # Skip if still placeholder
    if text and "| | | |" not in text and "[e.g.," not in text:
        return text
    return ""


def _md_table_to_html(lines):
    """Convert markdown table lines to an HTML table string."""
    if len(lines) < 2:
        return ""
    # Parse header
    header_cells = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    # Skip separator line (line 1)
    html = '<table class="scope-table"><thead><tr>'
    for cell in header_cells:
        html += "<th>{}</th>".format(_md_inline(cell))
    html += "</tr></thead><tbody>\\n"
    for row_line in lines[2:]:
        cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
        html += "<tr>"
        for cell in cells:
            html += "<td>{}</td>".format(_md_inline(cell))
        html += "</tr>\\n"
    html += "</tbody></table>"
    return html


def _md_inline(text):
    """Convert inline markdown (bold) to HTML."""
    import re as _re
    # **bold**
    text = _re.sub(r"\\*\\*(.+?)\\*\\*", r"<strong>\\1</strong>", text)
    return text


def _md_list_to_html(lines):
    """Convert markdown bullet list lines to an HTML list."""
    html = "<ul>\\n"
    for line in lines:
        item = line.lstrip("- ").strip()
        html += "  <li>{}</li>\\n".format(_md_inline(item))
    html += "</ul>"
    return html


def scope_content_to_html(content):
    """Convert scope section content (from task_spec.md) to styled HTML.

    Handles markdown tables, bold text, bullet lists, and sub-headings.
    """
    lines = content.split("\\n")
    html_parts = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Sub-heading (e.g., "Applicable Standards:")
        if (line.strip().endswith(":") and not line.strip().startswith("-")
                and not line.strip().startswith("|") and not line.strip().startswith("*")):
            html_parts.append("<h3>{}</h3>".format(_md_inline(line.strip())))
            i += 1
            continue

        # Markdown table (starts with |)
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # Check if second line is separator (|---|)
            if len(table_lines) >= 2 and set(table_lines[1].replace("|", "").replace("-", "").replace(":", "").strip()) <= set(""):
                html_parts.append(_md_table_to_html(table_lines))
            else:
                # Not a real table, just text with pipes
                for tl in table_lines:
                    html_parts.append("<p>{}</p>".format(_md_inline(tl.strip())))
            continue

        # Bullet list (starts with -)
        if line.strip().startswith("- "):
            list_lines = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                list_lines.append(lines[i])
                i += 1
            html_parts.append(_md_list_to_html(list_lines))
            continue

        # Regular paragraph
        html_parts.append("<p>{}</p>".format(_md_inline(line.strip())))
        i += 1

    return "\\n".join(html_parts)


def render_scope_to_word(doc, content):
    """Render scope section content (from task_spec.md) into a Word document.

    Parses markdown tables into Word tables, bold text into runs, and
    bullet lists into formatted paragraphs.
    """
    lines = content.split("\\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Sub-heading (e.g., "Applicable Standards:")
        if (line.strip().endswith(":") and not line.strip().startswith("-")
                and not line.strip().startswith("|") and not line.strip().startswith("*")):
            doc.add_heading(line.strip(), level=2)
            i += 1
            continue

        # Markdown table
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                _md_table_to_word(doc, table_lines)
            else:
                for tl in table_lines:
                    doc.add_paragraph(tl.strip())
            continue

        # Bullet list
        if line.strip().startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                item_text = lines[i].strip()[2:]  # Remove "- "
                p = doc.add_paragraph(style="List Bullet")
                _add_bold_runs(p, item_text)
                i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_bold_runs(p, line.strip())
        i += 1


def _md_table_to_word(doc, table_lines):
    """Convert markdown table lines to a styled Word table."""
    header_cells = [c.strip() for c in table_lines[0].strip().strip("|").split("|")]
    data_rows = []
    for row_line in table_lines[2:]:  # skip header and separator
        cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
        data_rows.append(cells)
    add_word_table(doc, header_cells, data_rows)


def _add_bold_runs(paragraph, text):
    """Add text with **bold** sections as separate runs."""
    import re as _re
    parts = _re.split(r"(\\*\\*.+?\\*\\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def get_figures():
    """Collect all PNG/SVG figures from the figures/ directory."""
    pngs = sorted(glob.glob(os.path.join(FIG_DIR, "*.png")))
    svgs = sorted(glob.glob(os.path.join(FIG_DIR, "*.svg")))
    return pngs + svgs


def get_figure_caption(fig_path, results, fig_index):
    """Get a caption for a figure: custom from results.json or auto-generated."""
    fig_name = os.path.basename(fig_path)
    captions = {}
    if results:
        captions = results.get("figure_captions", {})
    if fig_name in captions:
        return "Figure {}: {}".format(fig_index, captions[fig_name])
    # Auto-generate from filename
    auto = fig_name.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").title()
    return "Figure {}: {}".format(fig_index, auto)


def get_equations(results):
    """Get equations from results.json. Returns list of {label, latex}."""
    if not results:
        return []
    return results.get("equations", [])


def render_equation_to_image(latex_str, output_path):
    """Render a LaTeX equation to a high-quality PNG image using matplotlib.

    Uses display-style math, large font, and 300 DPI for crisp rendering
    in Word documents. Returns True if the image was created, False otherwise.
    """
    if not HAS_MATPLOTLIB:
        return False
    try:
        fig = plt.figure(figsize=(8, 1.2))
        fig.text(
            0.5, 0.5,
            "${}$".format(latex_str),
            fontsize=24, ha="center", va="center",
            math_fontfamily="cm",
        )
        fig.savefig(output_path, dpi=300, bbox_inches="tight",
                    pad_inches=0.15, facecolor="white", edgecolor="none")
        plt.close(fig)
        return True
    except Exception as e:
        print("  Warning: could not render equation: {}".format(e))
        return False


def _parse_key_name(key):
    """Parse a key_results key into (label, unit). Splits on last known unit suffix."""
    unit_suffixes = [
        ("_pct", "%"), ("_percent", "%"),
        ("_bar", "bar"), ("_bara", "bara"), ("_barg", "barg"),
        ("_psi", "psi"), ("_psia", "psia"),
        ("_C", "\u00b0C"), ("_K", "K"), ("_F", "\u00b0F"),
        ("_kg", "kg"), ("_g", "g"), ("_lb", "lb"),
        ("_m3", "m\u00b3"), ("_m2", "m\u00b2"), ("_m", "m"),
        ("_mm", "mm"), ("_cm", "cm"), ("_km", "km"),
        ("_ft", "ft"), ("_in", "in"),
        ("_kW", "kW"), ("_MW", "MW"), ("_W", "W"),
        ("_kJ", "kJ"), ("_MJ", "MJ"), ("_J", "J"),
        ("_kg_hr", "kg/hr"), ("_kg_s", "kg/s"),
        ("_m3_hr", "m\u00b3/hr"), ("_m3_s", "m\u00b3/s"),
        ("_Sm3_day", "Sm\u00b3/day"), ("_Sm3_hr", "Sm\u00b3/hr"),
        ("_hours", "hours"), ("_hr", "hr"), ("_min", "min"), ("_s", "s"),
        ("_rpm", "rpm"), ("_Hz", "Hz"),
    ]
    for suffix, unit in unit_suffixes:
        if key.endswith(suffix):
            name_part = key[:len(key) - len(suffix)]
            label = name_part.replace("_", " ").title()
            return label, unit
    return key.replace("_", " ").title(), ""


def format_results_table(results):
    """Format key_results dict as a text table."""
    key_results = results.get("key_results", {})
    if not key_results:
        return "[No key_results in results.json]"
    lines = []
    for key, value in key_results.items():
        label, unit = _parse_key_name(key)
        if isinstance(value, float):
            val_str = "{:.4g}".format(value)
        else:
            val_str = str(value)
        if unit:
            lines.append("{}: {} {}".format(label, val_str, unit))
        else:
            lines.append("{}: {}".format(label, val_str))
    return "\\n".join(lines)


def format_validation_table(results):
    """Format validation checks as a text table."""
    validation = results.get("validation", {})
    if not validation:
        return "[No validation data in results.json]"
    lines = ["Validation Summary:", ""]
    for check, outcome in validation.items():
        label = check.replace("_", " ").title()
        if isinstance(outcome, bool):
            status = "PASS" if outcome else "FAIL"
        elif isinstance(outcome, (int, float)):
            status = "{:.4g}".format(outcome)
        else:
            status = str(outcome)
        lines.append("  {}: {}".format(label, status))
    return "\\n".join(lines)


def format_validation_html(results):
    """Format validation checks as an HTML table."""
    validation = results.get("validation", {})
    if not validation:
        return "<p><em>No validation data in results.json</em></p>"
    rows = ""
    for check, outcome in validation.items():
        label = check.replace("_", " ").title()
        if isinstance(outcome, bool):
            status = "PASS" if outcome else "FAIL"
            css_class = ' class="pass"' if outcome else ' class="fail"'
        elif isinstance(outcome, (int, float)):
            status = "{:.4g}".format(outcome)
            css_class = ""
        else:
            status = str(outcome)
            css_class = ""
        rows += \'<tr><td>{}</td><td{}>{}</td></tr>\\n\'.format(
            label, css_class, status)
    return \'<table class="validation-table"><thead><tr><th>Check</th><th>Result</th></tr></thead><tbody>\\n{}</tbody></table>\'.format(rows)


def format_results_html(results):
    """Format key_results dict as a styled HTML table with units column."""
    key_results = results.get("key_results", {})
    if not key_results:
        return ""
    rows = ""
    for key, value in key_results.items():
        label, unit = _parse_key_name(key)
        if isinstance(value, float):
            val_str = "{:.4g}".format(value)
        else:
            val_str = str(value)
        rows += \'<tr><td>{}</td><td class="num">{}</td><td>{}</td></tr>\\n\'.format(
            label, val_str, unit)
    return (
        \'<table class="results-table"><thead>\'
        \'<tr><th>Parameter</th><th>Value</th><th>Unit</th></tr>\'
        \'</thead><tbody>\\n{}</tbody></table>\'.format(rows)
    )


def format_custom_tables_html(results):
    """Format custom tables from results.json 'tables' key as HTML."""
    tables = results.get("tables", [])
    if not tables:
        return ""
    html_parts = []
    for tbl in tables:
        title = tbl.get("title", "")
        headers = tbl.get("headers", [])
        data_rows = tbl.get("rows", [])
        if not headers or not data_rows:
            continue
        h = ""
        if title:
            h += \'<h3>{}</h3>\\n\'.format(title)
        h += \'<table class="custom-table"><thead><tr>\'
        for col in headers:
            h += \'<th>{}</th>\'.format(col)
        h += \'</tr></thead><tbody>\\n\'
        for row in data_rows:
            h += "<tr>"
            for i, cell in enumerate(row):
                css = \' class="num"\' if i > 0 and isinstance(cell, (int, float)) else ""
                if isinstance(cell, float):
                    h += \'<td{}>{:.4g}</td>\'.format(css, cell)
                else:
                    h += \'<td{}>{}</td>\'.format(css, cell)
            h += "</tr>\\n"
        h += "</tbody></table>"
        html_parts.append(h)
    return "\\n".join(html_parts)


def format_references_html(results):
    """Format the references list from results.json as a styled HTML ordered list."""
    refs = results.get("references", [])
    if not refs:
        return ""
    h = \'<ol class="reference-list">\\n\'
    for ref in refs:
        ref_id = ref.get("id", "")
        ref_text = ref.get("text", "")
        if ref_id:
            h += \'  <li id="ref-{}"><strong>[{}]</strong> {}</li>\\n\'.format(
                ref_id, ref_id, ref_text)
        else:
            h += \'  <li>{}</li>\\n\'.format(ref_text)
    h += \'</ol>\'
    return h


def add_word_table(doc, headers, data_rows, col_widths=None):
    """Add a professionally styled table to a Word document.

    Args:
        doc: Document object.
        headers: list of column header strings.
        data_rows: list of lists (each inner list = one row of cell values).
        col_widths: optional list of Inches widths per column.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(text)
        # Style header: bold, white text on dark blue background
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(
            \'<w:shd {} w:fill="2F5496"/>\'.format(nsdecls(\'w\'))
        )
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for row_data in data_rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Apply column widths if specified
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    doc.add_paragraph("")  # spacing after table
    return table


def add_results_word_table(doc, results):
    """Add key_results as a styled Word table with units column."""
    key_results = results.get("key_results", {})
    if not key_results:
        return
    headers = ["Parameter", "Value", "Unit"]
    data_rows = []
    for key, value in key_results.items():
        label, unit = _parse_key_name(key)
        if isinstance(value, float):
            val_str = "{:.4g}".format(value)
        else:
            val_str = str(value)
        data_rows.append([label, val_str, unit])
    add_word_table(doc, headers, data_rows,
                   col_widths=[Inches(3.0), Inches(1.5), Inches(1.5)])


def add_validation_word_table(doc, results):
    """Add validation checks as a styled Word table."""
    validation = results.get("validation", {})
    if not validation:
        return
    headers = ["Check", "Result"]
    data_rows = []
    for check, outcome in validation.items():
        label = check.replace("_", " ").title()
        if isinstance(outcome, bool):
            status = "PASS" if outcome else "FAIL"
        elif isinstance(outcome, (int, float)):
            status = "{:.4g}".format(outcome)
        else:
            status = str(outcome)
        data_rows.append([label, status])
    table = add_word_table(doc, headers, data_rows,
                           col_widths=[Inches(4.0), Inches(2.0)])
    # Color-code PASS/FAIL cells
    for row in table.rows[1:]:
        cell = row.cells[1]
        text = cell.text.strip()
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                if text == "PASS":
                    run.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)
                elif text == "FAIL":
                    run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)


def add_custom_word_tables(doc, results):
    """Add custom tables from results.json 'tables' key."""
    tables = results.get("tables", [])
    if not tables:
        return
    for tbl in tables:
        title = tbl.get("title", "")
        headers = tbl.get("headers", [])
        data_rows = tbl.get("rows", [])
        if not headers or not data_rows:
            continue
        if title:
            doc.add_heading(title, level=2)
        # Format numeric values
        formatted_rows = []
        for row in data_rows:
            formatted = []
            for cell in row:
                if isinstance(cell, float):
                    formatted.append("{:.4g}".format(cell))
                else:
                    formatted.append(str(cell))
            formatted_rows.append(formatted)
        add_word_table(doc, headers, formatted_rows)


# ══════════════════════════════════════════════════════════
# Build sections (auto-populated where possible)
# ══════════════════════════════════════════════════════════

def build_sections(results, task_spec):
    """Build report sections, auto-populating from results.json and task_spec.md."""
    sections = []

    # 1. Executive Summary
    sections.append({
        "heading": "1. Executive Summary",
        "content": MANUAL_SECTIONS["executive_summary"],
    })

    # 2. Problem Description
    sections.append({
        "heading": "2. Problem Description",
        "content": MANUAL_SECTIONS["problem_description"],
    })

    # 3. Scope and Standards (auto-populated from task_spec.md)
    scope_parts = []
    standards = extract_spec_section(task_spec, "Applicable Standards")
    if standards:
        scope_parts.append("Applicable Standards:\\n" + standards)
    methods = extract_spec_section(task_spec, "Calculation Methods")
    if methods:
        scope_parts.append("Calculation Methods:\\n" + methods)
    criteria = extract_spec_section(task_spec, "Acceptance Criteria")
    if criteria:
        scope_parts.append("Acceptance Criteria:\\n" + criteria)
    envelope = extract_spec_section(task_spec, "Operating Envelope")
    if envelope:
        scope_parts.append("Operating Envelope:\\n" + envelope)

    scope_content = "\\n\\n".join(scope_parts) if scope_parts else (
        "[Auto-populated from task_spec.md when filled in. "
        "Edit step1_scope_and_research/task_spec.md and re-run.]"
    )
    sections.append({
        "heading": "3. Scope and Standards",
        "content": scope_content,
        "has_scope": True,
    })

    # 4. Approach
    approach = MANUAL_SECTIONS["approach"]
    if results and results.get("approach"):
        approach = results["approach"]
    sections.append({
        "heading": "4. Approach",
        "content": approach,
        "has_equations": True,
    })

    # 5. Results (auto-populated from results.json)
    if results and results.get("key_results"):
        results_text = format_results_table(results)
    else:
        results_text = (
            "[Auto-populated from results.json when created by notebook. "
            "Save results with the pattern shown in the task README.]"
        )
    sections.append({
        "heading": "5. Results",
        "content": results_text,
        "has_figures": True,
    })

    # 6. Validation Summary (auto-populated from results.json)
    if results and results.get("validation"):
        validation_text = format_validation_table(results)
    else:
        validation_text = (
            "[Auto-populated from results.json validation section. "
            "Add validation checks to your notebook results output.]"
        )
    sections.append({
        "heading": "6. Validation Summary",
        "content": validation_text,
    })

    # 7. Conclusions and Recommendations
    conclusions = MANUAL_SECTIONS["conclusions"]
    if results and results.get("conclusions"):
        conclusions = results["conclusions"]
    sections.append({
        "heading": "7. Conclusions and Recommendations",
        "content": conclusions,
    })

    # 8. References (auto-populated from results.json if available)
    refs_content = MANUAL_SECTIONS["references"]
    if results and results.get("references"):
        ref_lines = []
        for i, ref in enumerate(results["references"], 1):
            ref_id = ref.get("id", "")
            ref_text = ref.get("text", "")
            if ref_id:
                ref_lines.append("[{}] {}".format(i, ref_text))
            else:
                ref_lines.append("[{}] {}".format(i, ref_text))
        refs_content = "\\n".join(ref_lines)
    sections.append({
        "heading": "8. References",
        "content": refs_content,
        "has_references": True,
    })

    return sections


# ══════════════════════════════════════════════════════════
# Word report
# ══════════════════════════════════════════════════════════

def build_word_report(sections, results=None):
    """Build the Word document with numbered figures, captions, and equations."""
    doc = Document()

    # Title page
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph("")
    doc.add_paragraph("Author: {}".format(AUTHOR or "(not specified)"))
    doc.add_paragraph("Date: {}".format(TASK_DATE))
    doc.add_page_break()

    # Add all sections
    for section in sections:
        doc.add_heading(section["heading"], level=1)

        # Results section: use Word table instead of plain text
        if section.get("has_figures") and results and results.get("key_results"):
            add_results_word_table(doc, results)
            # Custom tables
            if results.get("tables"):
                add_custom_word_tables(doc, results)
        elif section.get("has_figures"):
            # No results data — show placeholder text
            for para_text in section["content"].split("\\n\\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        elif section.get("has_scope"):
            # Scope section: parse markdown tables, bold, and lists
            render_scope_to_word(doc, section["content"])
        elif "Validation" in section["heading"] and results and results.get("validation"):
            # Validation section: use Word table
            add_validation_word_table(doc, results)
        else:
            # Regular text content
            for para_text in section["content"].split("\\n\\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

        # Embed figures after Results section
        if section.get("has_figures"):
            figures = get_figures()
            if figures:
                for fig_idx, fig_path in enumerate(figures, 1):
                    caption_text = get_figure_caption(fig_path, results, fig_idx)
                    doc.add_picture(fig_path, width=Inches(6.0))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = doc.add_paragraph(caption_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].font.size = Pt(9)
                    caption.runs[0].font.italic = True
                    doc.add_paragraph("")
            else:
                doc.add_paragraph(
                    "[No figures found in figures/ directory. "
                    "Save plots as PNG files there and re-run this script.]"
                )

        # Embed equations after Approach section
        if section.get("has_equations"):
            equations = get_equations(results)
            if equations:
                doc.add_heading("Key Equations", level=2)
                eq_img_dir = os.path.join(REPORT_DIR, "_eq_images")
                if not os.path.exists(eq_img_dir):
                    os.makedirs(eq_img_dir)
                for eq_idx, eq in enumerate(equations, 1):
                    label = eq.get("label", "Equation {}".format(eq_idx))
                    latex = eq.get("latex", "")
                    if not latex:
                        continue
                    # Try to render equation as image
                    eq_img_path = os.path.join(eq_img_dir, "eq_{}.png".format(eq_idx))
                    if render_equation_to_image(latex, eq_img_path):
                        doc.add_paragraph("")
                        doc.add_picture(eq_img_path, width=Inches(5.5))
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption = doc.add_paragraph(
                            "Equation {}: {}".format(eq_idx, label)
                        )
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.runs[0].font.size = Pt(9)
                        caption.runs[0].font.italic = True
                    else:
                        # Fallback: text representation
                        doc.add_paragraph("{}: {}".format(label, latex))
                    doc.add_paragraph("")

    # Save
    doc.save(DOCX_FILE)
    print("Word report saved: {}".format(DOCX_FILE))


# ══════════════════════════════════════════════════════════
# HTML report
# ══════════════════════════════════════════════════════════

def build_html_report(sections, results=None):
    """Build an HTML report with embedded figures, KaTeX equations, and navigation."""
    figures = get_figures()

    # Build figure HTML with base64-embedded images and numbered captions
    figure_html = ""
    if figures:
        for fig_idx, fig_path in enumerate(figures, 1):
            fig_name = os.path.basename(fig_path)
            caption_text = get_figure_caption(fig_path, results, fig_idx)
            # Determine MIME type
            if fig_path.endswith(".svg"):
                mime = "image/svg+xml"
            else:
                mime = "image/png"
            with open(fig_path, "rb") as f:
                img_data = base64.b64encode(f.read()).decode("utf-8")
            figure_html += """
            <div class="figure">
                <img src="data:{};base64,{}" alt="{}">
                <p class="caption">{}</p>
            </div>
            """.format(mime, img_data, caption_text, caption_text)
    else:
        figure_html = "<p><em>No figures found in figures/ directory.</em></p>"

    # Build equation HTML (KaTeX rendering with embedded image fallbacks)
    equation_html = ""
    equations = get_equations(results)
    if equations:
        equation_html += '<h3>Key Equations</h3>\\n'
        # Pre-render equation images for offline fallback
        eq_img_dir = os.path.join(REPORT_DIR, "_eq_images")
        if not os.path.exists(eq_img_dir):
            os.makedirs(eq_img_dir)
        for eq_idx, eq in enumerate(equations, 1):
            label = eq.get("label", "Equation {}".format(eq_idx))
            latex = eq.get("latex", "")
            if not latex:
                continue
            # Render fallback image
            fallback_img = ""
            eq_img_path = os.path.join(eq_img_dir, "eq_{}.png".format(eq_idx))
            if render_equation_to_image(latex, eq_img_path):
                with open(eq_img_path, "rb") as imgf:
                    img_b64 = base64.b64encode(imgf.read()).decode("utf-8")
                fallback_img = (
                    '<img class="eq-fallback" '
                    'src="data:image/png;base64,{}" '
                    'alt="{}" style="display:none; max-width:90%;">'.format(
                        img_b64, label)
                )
            equation_html += """
            <div class="equation-block">
                <div class="equation katex-eq">$${}$$</div>
                {}
                <p class="equation-label">Equation {}: {}</p>
            </div>
            """.format(latex, fallback_img, eq_idx, label)

    # Build validation HTML
    validation_html = ""
    if results and results.get("validation"):
        validation_html = format_validation_html(results)

    # Build key results HTML table
    results_table_html = ""
    if results and results.get("key_results"):
        results_table_html = format_results_html(results)

    # Build custom tables HTML
    custom_tables_html = ""
    if results and results.get("tables"):
        custom_tables_html = format_custom_tables_html(results)

    # Build section HTML and navigation
    nav_items = ""
    section_html = ""
    for section in sections:
        section_id = section["heading"].lower().replace(" ", "-").replace(".", "")
        nav_items += \'    <li><a href="#{}">{}</a></li>\\n\'.format(
            section_id, section["heading"]
        )
        # Convert scope section markdown to HTML
        if section.get("has_scope"):
            content = scope_content_to_html(section["content"])
        else:
            content = section["content"].replace("\\n", "<br>")

        # Insert auto-generated HTML for special sections
        if section.get("has_figures"):
            if results_table_html:
                content = results_table_html + custom_tables_html + figure_html
            else:
                content += figure_html

        if section.get("has_equations") and equation_html:
            content += equation_html

        if "Validation" in section["heading"] and validation_html:
            content = validation_html

        if section.get("has_references") and results and results.get("references"):
            content = format_references_html(results)

        section_html += """
        <section id="{}">
            <h2>{}</h2>
            <div>{}</div>
        </section>
        """.format(section_id, section["heading"], content)

    # KaTeX CDN for equation rendering (only if equations exist)
    katex_head = ""
    katex_body_script = ""
    if equations:
        katex_head = """
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.css">
    <script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.js"></script>
    <script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/contrib/auto-render.min.js"></script>"""
        katex_body_script = """
    <script>
        document.addEventListener("DOMContentLoaded", function() {
            if (typeof renderMathInElement === "function") {
                renderMathInElement(document.body, {
                    delimiters: [
                        {left: "$$", right: "$$", display: true},
                        {left: "$", right: "$", display: false}
                    ],
                    throwOnError: false
                });
            } else {
                // KaTeX not available (offline) — show fallback images
                var eqs = document.querySelectorAll(".katex-eq");
                for (var i = 0; i < eqs.length; i++) {
                    eqs[i].style.display = "none";
                }
                var imgs = document.querySelectorAll(".eq-fallback");
                for (var j = 0; j < imgs.length; j++) {
                    imgs[j].style.display = "inline";
                }
            }
        });
    </script>"""

    html = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>{katex_head}
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
               display: flex; line-height: 1.6; color: #333; }}
        nav {{ width: 260px; min-height: 100vh; background: #f5f5f5; padding: 1.5rem;
              position: fixed; overflow-y: auto; border-right: 1px solid #ddd; }}
        nav h3 {{ margin-bottom: 1rem; color: #555; font-size: 0.9rem;
                  text-transform: uppercase; letter-spacing: 0.05em; }}
        nav ul {{ list-style: none; }}
        nav li {{ margin-bottom: 0.5rem; }}
        nav a {{ color: #0366d6; text-decoration: none; font-size: 0.9rem; }}
        nav a:hover {{ text-decoration: underline; }}
        main {{ margin-left: 260px; max-width: 900px; padding: 2rem 3rem; }}
        h1 {{ margin-bottom: 0.5rem; color: #1a1a1a; }}
        h2 {{ margin-top: 2rem; margin-bottom: 1rem; color: #1a1a1a;
             border-bottom: 1px solid #eee; padding-bottom: 0.3rem; }}
        h3 {{ margin-top: 1.5rem; margin-bottom: 0.5rem; color: #333; }}
        .meta {{ color: #666; margin-bottom: 2rem; }}
        section {{ margin-bottom: 2rem; }}
        .figure {{ text-align: center; margin: 1.5rem 0; }}
        .figure img {{ max-width: 100%; border: 1px solid #ddd; border-radius: 4px; }}
        .caption {{ font-size: 0.85rem; color: #666; font-style: italic;
                    margin-top: 0.3rem; }}
        .equation-block {{ margin: 1.5rem 0; text-align: center; }}
        .equation {{ font-size: 1.2rem; padding: 0.5rem 0; }}
        .equation-label {{ font-size: 0.85rem; color: #666; font-style: italic;
                           margin-top: 0.2rem; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1.5rem 0;
                font-size: 0.92rem; box-shadow: 0 1px 3px rgba(0,0,0,0.08); }}
        thead th {{ background: #2F5496; color: #fff; font-weight: 600;
                    padding: 0.6rem 0.75rem; text-align: left;
                    border: 1px solid #2a4a85; }}
        tbody td {{ border: 1px solid #e0e0e0; padding: 0.5rem 0.75rem;
                    text-align: left; }}
        tbody tr:nth-child(even) {{ background: #f8f9fa; }}
        tbody tr:hover {{ background: #e9ecef; }}
        td.num {{ text-align: right; font-variant-numeric: tabular-nums; }}
        .pass {{ color: #28a745; font-weight: bold; }}
        .fail {{ color: #dc3545; font-weight: bold; }}
        .results-table {{ max-width: 600px; }}
        .validation-table {{ max-width: 500px; }}
        .custom-table {{ margin-top: 0.5rem; }}
        .scope-table {{ margin: 0.5rem 0 1rem 0; }}
        section h3 {{ color: #2F5496; margin-top: 1.2rem; margin-bottom: 0.4rem;
            font-size: 1.1rem; border-bottom: 1px solid #ddd; padding-bottom: 0.2rem; }}
        section ul {{ margin: 0.3rem 0 0.8rem 1.5rem; }}
        section ul li {{ margin-bottom: 0.3rem; }}
        .reference-list {{ list-style: none; padding-left: 0; }}
        .reference-list li {{ margin-bottom: 0.6rem; padding: 0.4rem 0.6rem;
            border-left: 3px solid #2F5496; background: #f8f9fa; }}
        .reference-list li strong {{ color: #2F5496; }}
        @media (max-width: 768px) {{
            nav {{ position: static; width: 100%; min-height: auto; }}
            main {{ margin-left: 0; padding: 1rem; }}
        }}
    </style>
</head>
<body>
    <nav>
        <h3>Contents</h3>
        <ul>
{nav}
        </ul>
        <hr style="margin: 1rem 0;">
        <p style="font-size: 0.8rem; color: #999;">Generated {date}</p>
    </nav>
    <main>
        <h1>{title}</h1>
        <p class="meta">Author: {author} | Date: {date}</p>
{sections}
    </main>{katex_body_script}
</body>
</html>""".format(
        title=TITLE,
        author=AUTHOR or "(not specified)",
        date=TASK_DATE,
        nav=nav_items,
        sections=section_html,
        katex_head=katex_head,
        katex_body_script=katex_body_script,
    )

    with open(HTML_FILE, "w", encoding="utf-8") as f:
        f.write(html)
    print("HTML report saved: {}".format(HTML_FILE))


# ══════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Generating reports for: {}".format(TITLE))
    print("")

    # Auto-read task data
    results = load_results()
    task_spec = load_task_spec()

    # Build sections (auto-populated where data exists)
    sections = build_sections(results, task_spec)

    # Generate both formats
    print("")
    build_word_report(sections, results)
    build_html_report(sections, results)
    print("")
    print("Both reports generated.")
    print("  Open Report.html in a browser for navigable view.")
    print("  Open Report.docx for formal distribution.")
    if not results:
        print("")
        print("TIP: Create results.json in the task root to auto-populate")
        print("     the Results and Validation sections. See the task README")
        print("     for the results.json pattern.")
'''


# ══════════════════════════════════════════════════════════
# Functions
# ══════════════════════════════════════════════════════════

def slugify(title):
    """Convert a title to a folder-safe slug."""
    slug = title.lower().strip()
    for ch in ",:;!?()[]{}'\"/\\.":
        slug = slug.replace(ch, "")
    slug = slug.replace(" ", "_").replace("-", "_")
    while "__" in slug:
        slug = slug.replace("__", "_")
    return slug.strip("_")


def _write_file(path, content):
    """Write content to a file, creating parent dirs as needed."""
    parent = os.path.dirname(path)
    if parent and not os.path.exists(parent):
        os.makedirs(parent)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


def setup_workspace():
    """
    Create the task_solve/ folder with README and TASK_TEMPLATE.

    Safe to call multiple times — skips files that already exist.
    Returns True if anything was created.
    """
    created = False

    # Main README
    readme = os.path.join(TASK_SOLVE_DIR, "README.md")
    if not os.path.exists(readme):
        _write_file(readme, WORKSPACE_README)
        created = True

    # Template structure
    template_files = {
        os.path.join(TEMPLATE_DIR, "README.md"): TASK_README,
        os.path.join(TEMPLATE_DIR, "step1_scope_and_research", "task_spec.md"): TASK_SPEC,
        os.path.join(TEMPLATE_DIR, "step1_scope_and_research", "notes.md"): STEP1_NOTES,
        os.path.join(TEMPLATE_DIR, "step1_scope_and_research", "references", "README.md"): REFERENCES_README,
        os.path.join(TEMPLATE_DIR, "step2_analysis", "notes.md"): STEP2_NOTES,
        os.path.join(TEMPLATE_DIR, "step2_analysis", ".gitkeep"): "",
        os.path.join(TEMPLATE_DIR, "step3_report", "generate_report.py"): GENERATE_REPORT,
        os.path.join(TEMPLATE_DIR, "figures", ".gitkeep"): "",
    }

    for path, content in template_files.items():
        if not os.path.exists(path):
            _write_file(path, content)
            created = True

    return created


def create_task(title, task_type="B", author=""):
    """Create a new task folder from the template."""
    # Ensure workspace exists
    if not os.path.exists(TEMPLATE_DIR):
        print("Setting up task_solve/ workspace for the first time...")
        setup_workspace()
        print("")

    today = date.today().isoformat()
    folder_name = "{}_{}".format(today, slugify(title))
    task_dir = os.path.join(TASK_SOLVE_DIR, folder_name)

    if os.path.exists(task_dir):
        print("ERROR: Folder already exists: {}".format(task_dir))
        sys.exit(1)

    # Copy template
    shutil.copytree(TEMPLATE_DIR, task_dir)

    # Fill in the README
    readme_path = os.path.join(task_dir, "README.md")
    with open(readme_path, "r", encoding="utf-8") as f:
        content = f.read()

    type_label = "{} ({})".format(task_type, TASK_TYPES.get(task_type, ""))
    content = content.replace("[Title]", title)
    content = content.replace("YYYY-MM-DD", today)
    content = content.replace(
        "A (Property) | B (Process) | C (PVT) | D (Standards) | E (Feature) | F (Design) | G (Workflow)",
        type_label,
    )
    content = content.replace("[THIS_FOLDER]", folder_name)
    if author:
        content = content.replace(
            "**Status:**",
            "**Author:** {}\n**Status:**".format(author),
        )

    with open(readme_path, "w", encoding="utf-8") as f:
        f.write(content)

    # Fill in the step1 notes
    notes_path = os.path.join(task_dir, "step1_scope_and_research", "notes.md")
    with open(notes_path, "r", encoding="utf-8") as f:
        notes = f.read()
    notes = notes.replace(
        "[Summary of the engineering context]",
        "Task: {}".format(title),
    )
    with open(notes_path, "w", encoding="utf-8") as f:
        f.write(notes)

    # Fill in generate_report.py
    report_path = os.path.join(task_dir, "step3_report", "generate_report.py")
    with open(report_path, "r", encoding="utf-8") as f:
        report = f.read()
    report = report.replace('TITLE = "Task Report"', 'TITLE = "{}"'.format(title))
    if author:
        report = report.replace('AUTHOR = ""', 'AUTHOR = "{}"'.format(author))
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report)

    print("Created: task_solve/{}".format(folder_name))
    print("")
    print("Next steps (pick one):")
    print("")
    print("  Recommended - Let Copilot do everything:")
    print("    Open VS Code Copilot Chat and type:")
    print("")
    print("    @solve.task {}".format(title))
    print("")
    print("  Alternative - Follow prompts manually:")
    print("    Open task_solve/{}/README.md".format(folder_name))
    print("")
    return task_dir


def list_tasks():
    """List existing task folders."""
    if not os.path.exists(TASK_SOLVE_DIR):
        print("No task_solve/ folder yet. Run: python devtools/new_task.py --setup")
        return

    entries = sorted(os.listdir(TASK_SOLVE_DIR))
    tasks = [
        e for e in entries
        if os.path.isdir(os.path.join(TASK_SOLVE_DIR, e))
        and e != "TASK_TEMPLATE"
    ]

    if not tasks:
        print("No tasks yet. Create one with: python devtools/new_task.py \"your task\"")
    else:
        print("Tasks in task_solve/:")
        for t in tasks:
            readme = os.path.join(TASK_SOLVE_DIR, t, "README.md")
            status = ""
            if os.path.exists(readme):
                with open(readme, "r", encoding="utf-8") as f:
                    for line in f:
                        if "**Status:**" in line:
                            status = line.split("**Status:**")[-1].strip()
                            break
            print("  {} {}".format(t, "[{}]".format(status) if status else ""))


def main():
    if len(sys.argv) < 2 or sys.argv[1] in ("-h", "--help"):
        print(__doc__)
        sys.exit(0)

    if sys.argv[1] == "--setup":
        if setup_workspace():
            print("Created task_solve/ workspace with README and template.")
        else:
            print("task_solve/ workspace already exists.")
        print("\nCreate a task: python devtools/new_task.py \"your task title\"")
        return

    if sys.argv[1] == "--list":
        list_tasks()
        return

    title = sys.argv[1]
    task_type = "B"
    author = ""

    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == "--type" and i + 1 < len(sys.argv):
            task_type = sys.argv[i + 1].upper()
            i += 2
        elif sys.argv[i] == "--author" and i + 1 < len(sys.argv):
            author = sys.argv[i + 1]
            i += 2
        else:
            i += 1

    if task_type not in TASK_TYPES:
        print("WARNING: Unknown task type '{}'. Valid: {}".format(
            task_type, ", ".join(sorted(TASK_TYPES.keys()))))
        print("Using type B (Process) as default.")
        task_type = "B"

    create_task(title, task_type, author)


if __name__ == "__main__":
    main()
