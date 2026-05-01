"""
generate_report.py - Generate Word/HTML reports and scientific papers for this task.

Usage:
    pip install python-docx matplotlib   (one-time setup)
    python step3_report/generate_report.py            # Technical report only
    python step3_report/generate_report.py --paper     # Also generate scientific paper
    python step3_report/generate_report.py --paper-only  # Scientific paper only

This script AUTO-READS data from the task folder:
    - study_config.yaml                    -> defines depth, notebook plan, quality gates
  - step1_scope_and_research/task_spec.md  -> populates Scope & Standards
  - results.json (task root)               -> populates Results + Validation
  - figures/*.png                          -> embeds all plots
  - results.json "equations"               -> renders equations (KaTeX/images)
  - results.json "figure_captions"         -> custom captions for figures

It produces:
  - step3_report/Report.docx  (Word document for formal distribution)
  - step3_report/Report.html  (navigable HTML with sidebar, KaTeX equations)
  - step3_report/Paper.docx   (scientific paper in Word format, with --paper)
  - step3_report/Paper.html   (scientific paper in HTML format, with --paper)

If results.json or task_spec.md are missing, the report uses placeholder text.
Customize MANUAL_SECTIONS below for content that can't be auto-generated.
"""
import os
import sys
import glob
import json
import base64
import io
import sqlite3
from datetime import date

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# Optional: matplotlib for rendering equations to images (Word report)
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

# ── Paths ────────────────────────────────────────────────
TASK_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(TASK_DIR, "figures")
REPORT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_FILE = os.path.join(REPORT_DIR, "Report.docx")
HTML_FILE = os.path.join(REPORT_DIR, "Report.html")
PAPER_DOCX_FILE = os.path.join(REPORT_DIR, "Paper.docx")
PAPER_HTML_FILE = os.path.join(REPORT_DIR, "Paper.html")
RESULTS_FILE = os.path.join(TASK_DIR, "results.json")
TASK_SPEC_FILE = os.path.join(TASK_DIR, "step1_scope_and_research", "task_spec.md")
STUDY_CONFIG_FILE = os.path.join(TASK_DIR, "study_config.yaml")

# ── Configuration (edit these) ───────────────────────────
TITLE = "Task Report"           # <-- Change to your task title
AUTHOR = ""                     # <-- Your name
TASK_DATE = date.today().isoformat()

# ── Paper-specific configuration (edit for scientific paper output) ──
PAPER_TITLE = ""                # <-- Leave empty to use TITLE
PAPER_AUTHORS = []              # <-- e.g. [{"name": "J. Doe", "affiliation": "NTNU"}]
PAPER_KEYWORDS = []             # <-- e.g. ["thermodynamics", "process simulation"]
PAPER_JOURNAL = ""              # <-- e.g. "Journal of Natural Gas Science and Engineering"
PAPER_ACKNOWLEDGMENTS = ""      # <-- e.g. "Funded by Research Council of Norway"

# ── Report metadata (cover page and revision history) ────
DOC_NUMBER = ""                 # <-- e.g. "REP-2026-001" (auto-generated from task folder if blank)
REVISION = "0"                  # <-- Current revision number
REVISION_HISTORY = [
    # {"rev": "0", "date": "2026-07-04", "description": "Initial issue", "author": ""},
]
CLASSIFICATION = "Open"         # <-- "Open", "Internal", "Confidential"

# ── Manual sections (edit content for your specific task) ─
# These are used when auto-read data is not available.
# If results.json exists, sections 5-6 are auto-populated.
MANUAL_SECTIONS = {
    "executive_summary": (
        "[Replace with a 3-5 sentence summary of the task, approach, "
        "and key findings.]"
    ),
    "problem_description": (
        "[Describe the engineering question or task that was solved.]"
    ),
    "approach": (
        "[Describe the methodology: EOS used, process configuration, "
        "simulation setup, key assumptions.]"
    ),
    "conclusions": (
        "[Summarize key findings and provide recommendations.]"
    ),
    "references": (
        "[List references from step1_scope_and_research/notes.md.]"
    ),
}

# ── Scientific paper sections (edit for paper output) ────
# These map to standard engineering paper sections.
# Auto-populated fields from results.json override placeholders.
PAPER_SECTIONS = {
    "abstract": (
        "[Replace with a 150-300 word abstract summarizing the problem, "
        "methodology, key results, and conclusions.]"
    ),
    "introduction": (
        "[Replace with 2-4 paragraphs covering:\n"
        "- Background and motivation\n"
        "- Brief literature review\n"
        "- Problem statement and objectives\n"
        "- Paper organization (optional)]"
    ),
    "methodology": (
        "[Replace with a description of the methodology:\n"
        "- Thermodynamic model and equation of state\n"
        "- Process simulation setup and assumptions\n"
        "- Numerical methods and convergence criteria\n"
        "- Key equations and correlations used]"
    ),
    "results_discussion": (
        "[Replace with results discussion if auto-populated "
        "data from results.json is insufficient.]"
    ),
    "conclusions": (
        "[Replace with concise, numbered conclusions and "
        "recommendations for future work.]"
    ),
    "acknowledgments": (
        ""
    ),
}


# ══════════════════════════════════════════════════════════
# Auto-read functions
# ══════════════════════════════════════════════════════════

def load_results():
    """Load results.json if it exists. Returns dict or None."""
    if os.path.exists(RESULTS_FILE):
        with open(RESULTS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        print("  Loaded results.json ({} keys)".format(len(data)))
        return data
    print("  No results.json found (using manual sections)")
    return None


def load_task_spec():
    """Load task_spec.md and extract standards/methods/criteria sections."""
    if not os.path.exists(TASK_SPEC_FILE):
        print("  No task_spec.md found (using placeholder for scope)")
        return None
    with open(TASK_SPEC_FILE, "r", encoding="utf-8") as f:
        content = f.read()
    print("  Loaded task_spec.md ({} chars)".format(len(content)))
    return content


def _strip_yaml_comment(line):
    """Strip YAML comments while preserving hashes inside quoted strings."""
    result = []
    quote = None
    for character in line:
        if character in ('"', "'"):
            if quote == character:
                quote = None
            elif quote is None:
                quote = character
        if character == "#" and quote is None:
            break
        result.append(character)
    return "".join(result).rstrip()


def _parse_yaml_value(value):
    """Parse the simple scalar values used by study_config.yaml."""
    cleaned = _strip_yaml_comment(value).strip()
    if not cleaned:
        return ""
    if ((cleaned.startswith('"') and cleaned.endswith('"'))
            or (cleaned.startswith("'") and cleaned.endswith("'"))):
        return cleaned[1:-1]
    lowered = cleaned.lower()
    if lowered == "true":
        return True
    if lowered == "false":
        return False
    try:
        return int(cleaned)
    except ValueError:
        return cleaned


def _section_lines(config_text, section):
    """Return lines belonging to a top-level YAML section."""
    lines = config_text.splitlines()
    section_marker = "{}:".format(section)
    capturing = False
    result = []
    for line in lines:
        stripped = line.strip()
        if stripped == section_marker and not line.startswith(" "):
            capturing = True
            continue
        if capturing and stripped and not line.startswith(" "):
            break
        if capturing:
            result.append(line)
    return result


def _parse_section_scalars(lines):
    """Parse scalar keys directly under a YAML section."""
    parsed = {}
    child_indents = [len(line) - len(line.lstrip()) for line in lines
                     if line.strip() and line.startswith(" ")]
    if not child_indents:
        return parsed
    child_indent = min(child_indents)
    for line in lines:
        cleaned = _strip_yaml_comment(line)
        if not cleaned.strip() or not cleaned.startswith(" "):
            continue
        indent = len(cleaned) - len(cleaned.lstrip())
        if indent != child_indent:
            continue
        if ":" not in cleaned:
            continue
        key, value = cleaned.strip().split(":", 1)
        value = value.strip()
        if value:
            parsed[key] = _parse_yaml_value(value)
    return parsed


def _parse_scalar_list(lines, key):
    """Parse a scalar list under an indented YAML key."""
    values = []
    capturing = False
    key_indent = 0
    for line in lines:
        cleaned = _strip_yaml_comment(line)
        stripped = cleaned.strip()
        if stripped == "{}:".format(key):
            capturing = True
            key_indent = len(cleaned) - len(cleaned.lstrip())
            continue
        if capturing:
            indent = len(cleaned) - len(cleaned.lstrip())
            if stripped and indent <= key_indent:
                break
            if stripped.startswith("- "):
                values.append(_parse_yaml_value(stripped[2:]))
    return values


def _parse_notebook_plan(lines):
    """Parse notebooks.plan entries from study_config.yaml."""
    plan = []
    current = None
    capturing = False
    plan_indent = 0
    for line in lines:
        cleaned = _strip_yaml_comment(line)
        stripped = cleaned.strip()
        if stripped == "plan:":
            capturing = True
            plan_indent = len(cleaned) - len(cleaned.lstrip())
            continue
        if capturing:
            indent = len(cleaned) - len(cleaned.lstrip())
            if stripped and indent <= plan_indent:
                break
            if stripped.startswith("- file:"):
                if current:
                    plan.append(current)
                current = {"file": _parse_yaml_value(stripped.split(":", 1)[1])}
            elif current and stripped.startswith("purpose:"):
                current["purpose"] = _parse_yaml_value(stripped.split(":", 1)[1])
    if current:
        plan.append(current)
    return plan


def _parse_mapping_list(lines, key):
    """Parse a list of simple mappings under an indented YAML key."""
    values = []
    current = None
    capturing = False
    key_indent = 0
    for line in lines:
        cleaned = _strip_yaml_comment(line)
        stripped = cleaned.strip()
        if stripped == "{}:".format(key):
            capturing = True
            key_indent = len(cleaned) - len(cleaned.lstrip())
            continue
        if capturing:
            indent = len(cleaned) - len(cleaned.lstrip())
            if stripped and indent <= key_indent:
                break
            if not stripped:
                continue
            if stripped.startswith("- "):
                if current:
                    values.append(current)
                current = {}
                item = stripped[2:].strip()
                if ":" in item:
                    item_key, item_value = item.split(":", 1)
                    current[item_key.strip()] = _parse_yaml_value(item_value)
            elif current is not None and ":" in stripped:
                item_key, item_value = stripped.split(":", 1)
                current[item_key.strip()] = _parse_yaml_value(item_value)
    if current:
        values.append(current)
    return values


def load_study_config():
    """Load study_config.yaml and return the subset used by this generator."""
    if not os.path.exists(STUDY_CONFIG_FILE):
        print("  No study_config.yaml found (using inferred task depth)")
        return {}
    with open(STUDY_CONFIG_FILE, "r", encoding="utf-8") as config_file:
        text = config_file.read()
    config = {}
    for section in ["study", "inputs", "notebooks", "report", "quality_gates"]:
        lines = _section_lines(text, section)
        config[section] = _parse_section_scalars(lines)
    config["report"]["formats"] = _parse_scalar_list(
        _section_lines(text, "report"), "formats")
    config["report"]["required_sections"] = _parse_scalar_list(
        _section_lines(text, "report"), "required_sections")
    config["notebooks"]["plan"] = _parse_notebook_plan(
        _section_lines(text, "notebooks"))
    config["inputs"]["documents"] = _parse_mapping_list(
        _section_lines(text, "inputs"), "documents")
    print("  Loaded study_config.yaml")
    return config


def extract_spec_section(spec_text, heading):
    """Extract a section from task_spec.md by heading."""
    if not spec_text:
        return ""
    lines = spec_text.split("\n")
    capturing = False
    result = []
    for line in lines:
        if line.startswith("## ") and heading.lower() in line.lower():
            capturing = True
            continue
        elif line.startswith("## ") and capturing:
            break
        elif capturing:
            result.append(line)
    text = "\n".join(result).strip()
    # Skip if still placeholder
    if text and "| | | |" not in text and "[e.g.," not in text:
        return text
    return ""


def _as_bool(value):
    """Interpret YAML-like values as booleans."""
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in ("true", "yes", "required")


def _as_int(value, default=0):
    """Interpret YAML-like values as integers."""
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _is_required(value):
    """Return true when a config value means required."""
    return str(value).strip().lower() == "required" or value is True


def _manual_section_filled(name):
    """Return true when a manual section has been filled in."""
    content = MANUAL_SECTIONS.get(name, "")
    return bool(content and not content.lstrip().startswith("["))


def _required_section_available(section, results, task_spec):
    """Check whether a configured report section has enough input data."""
    normalized = str(section).strip().lower().replace("-", "_").replace(" ", "_")
    if normalized == "executive_summary":
        return True
    if normalized in ("scope", "scope_and_standards"):
        if not task_spec:
            return False
        return bool(extract_spec_section(task_spec, "Applicable Standards")
                    or extract_spec_section(task_spec, "Calculation Methods")
                    or extract_spec_section(task_spec, "Acceptance Criteria")
                    or extract_spec_section(task_spec, "Operating Envelope"))
    if normalized in ("methodology", "approach"):
        return bool((results and results.get("approach"))
                    or _manual_section_filled("approach"))
    if normalized == "results":
        return bool(results and results.get("key_results"))
    if normalized == "discussion":
        return bool(results and results.get("figure_discussion"))
    if normalized == "validation":
        return bool(results and results.get("validation"))
    if normalized == "benchmark_validation":
        return bool(results and results.get("benchmark_validation"))
    if normalized in ("uncertainty", "uncertainty_analysis"):
        return bool(results and results.get("uncertainty"))
    if normalized in ("risk", "risk_assessment", "risk_evaluation"):
        return bool(results and (results.get("risk_evaluation") or results.get("risks")))
    if normalized in ("conclusion", "conclusions", "conclusions_and_recommendations"):
        return bool((results and results.get("conclusions"))
                    or _manual_section_filled("conclusions"))
    if normalized == "references":
        return bool((results and results.get("references"))
                    or _manual_section_filled("references"))
    return True


def _load_runner_jobs(runner_db_path):
    """Return runner job rows from runner.db or an error message."""
    try:
        connection = sqlite3.connect(runner_db_path, timeout=5.0)
        connection.row_factory = sqlite3.Row
        try:
            rows = connection.execute(
                "SELECT job_id, script, job_type, status, error_message "
                "FROM jobs ORDER BY created_at"
            ).fetchall()
        finally:
            connection.close()
    except sqlite3.Error as error:
        return [], str(error)
    return [dict(row) for row in rows], None


def _runner_script_name(job_row):
    """Return the basename of a runner job script path."""
    return os.path.basename(str(job_row.get("script") or ""))


def _runner_status_summary(job_rows):
    """Return a compact status summary for report warnings."""
    counts = {}
    for job_row in job_rows:
        status = str(job_row.get("status") or "unknown")
        counts[status] = counts.get(status, 0) + 1
    return ", ".join(["{}={}".format(status, counts[status])
                      for status in sorted(counts)])


def _validate_runner_execution(notebooks, planned_notebooks, existing_notebooks):
    """Return warnings for incomplete NeqSim Runner notebook execution."""
    warnings = []
    execution_engine = str(notebooks.get("execution_engine", "")).strip().lower()
    notebook_execution_required = _as_bool(notebooks.get("execution_required", False))
    require_successful_jobs = _as_bool(notebooks.get("require_successful_jobs", True))
    if not (notebook_execution_required and execution_engine == "neqsim_runner"):
        return warnings
    if not existing_notebooks:
        return warnings

    runner_output_path = os.path.join(TASK_DIR, "runner_output")
    runner_db_path = os.path.join(TASK_DIR, "runner.db")
    if not os.path.exists(runner_output_path):
        warnings.append("Notebook execution engine is neqsim_runner, but runner_output/ is missing.")
    if not os.path.exists(runner_db_path):
        warnings.append("Notebook execution engine is neqsim_runner, but runner.db is missing.")
        return warnings

    job_rows, load_error = _load_runner_jobs(runner_db_path)
    if load_error:
        warnings.append("Could not inspect runner.db for notebook status: {}".format(
            load_error))
        return warnings
    if not job_rows:
        warnings.append("runner.db exists, but it contains no recorded runner jobs.")
        return warnings

    expected_notebooks = planned_notebooks or [os.path.basename(path)
                                               for path in existing_notebooks]
    expected_script_names = set()
    for notebook_file in expected_notebooks:
        notebook_name = os.path.basename(str(notebook_file))
        expected_script_names.add(notebook_name)
        expected_script_names.add(os.path.splitext(notebook_name)[0] + ".py")

    notebook_rows = []
    for row in job_rows:
        job_type = str(row.get("job_type") or "script")
        script_name = _runner_script_name(row)
        if job_type == "notebook" or script_name in expected_script_names:
            notebook_rows.append(row)
    if not notebook_rows:
        warnings.append(
            "runner.db contains jobs ({}), but none match the planned notebooks.".format(
                _runner_status_summary(job_rows)))
        return warnings

    if require_successful_jobs:
        unsuccessful = [row for row in notebook_rows
                        if str(row.get("status") or "") != "success"]
        if unsuccessful:
            details = []
            for row in unsuccessful[:5]:
                details.append("{}={} ({})".format(
                    row.get("job_id"), row.get("status"), _runner_script_name(row)))
            if len(unsuccessful) > 5:
                details.append("... {} more".format(len(unsuccessful) - 5))
            warnings.append("NeqSim Runner notebook jobs are not all successful: {}.".format(
                "; ".join(details)))

        successful_notebooks = set([_runner_script_name(row) for row in notebook_rows
                                    if str(row.get("status") or "") == "success"])
        for notebook_file in expected_notebooks:
            notebook_name = os.path.basename(str(notebook_file))
            script_name = os.path.splitext(notebook_name)[0] + ".py"
            notebook_path = os.path.join(TASK_DIR, "step2_analysis", str(notebook_file))
            if (os.path.exists(notebook_path)
                    and notebook_name not in successful_notebooks
                    and script_name not in successful_notebooks):
                warnings.append(
                    "Planned notebook has no successful runner job: step2_analysis/{}".format(
                        notebook_file))
    return warnings


def validate_study_config(config, results, task_spec):
    """Return warnings for missing deliverables required by study_config.yaml."""
    if not config:
        return []

    warnings = []
    inputs = config.get("inputs", {})
    notebooks = config.get("notebooks", {})
    report = config.get("report", {})
    quality_gates = config.get("quality_gates", {})

    documents = inputs.get("documents", [])
    documents_required = _as_bool(inputs.get("documents_required", False))
    if documents_required and not documents:
        warnings.append("inputs.documents_required is true, but no input documents are listed.")
    for document in documents:
        document_path = document.get("path")
        document_required = documents_required or _as_bool(document.get("required", False))
        if not document_path or not document_required:
            continue
        if os.path.isabs(str(document_path)):
            resolved_path = str(document_path)
        else:
            resolved_path = os.path.join(TASK_DIR, str(document_path))
        if not os.path.exists(resolved_path):
            warnings.append("Required input document path is missing: {}".format(document_path))
        elif os.path.isdir(resolved_path):
            files = [name for name in os.listdir(resolved_path)
                     if os.path.isfile(os.path.join(resolved_path, name))]
            if not files:
                warnings.append("Required input document directory is empty: {}".format(
                    document_path))

    planned_notebooks = [entry.get("file") for entry in notebooks.get("plan", [])
                         if entry.get("file")]
    existing_notebooks = glob.glob(os.path.join(TASK_DIR, "step2_analysis", "*.ipynb"))
    minimum_count = _as_int(notebooks.get("minimum_count"), 0)
    if minimum_count and len(existing_notebooks) < minimum_count:
        warnings.append(
            "Configured notebook minimum is {}, but {} notebook(s) exist.".format(
                minimum_count, len(existing_notebooks)))
    for notebook_file in planned_notebooks:
        notebook_path = os.path.join(TASK_DIR, "step2_analysis", notebook_file)
        if not os.path.exists(notebook_path):
            warnings.append("Planned notebook is missing: step2_analysis/{}".format(
                notebook_file))

    warnings.extend(_validate_runner_execution(notebooks, planned_notebooks,
                                               existing_notebooks))

    if _as_bool(quality_gates.get("require_results_json", False)) and not results:
        warnings.append("quality_gates.require_results_json is true, but results.json is missing.")
    if _is_required(quality_gates.get("benchmark_validation")) and not (
            results and results.get("benchmark_validation")):
        warnings.append("Benchmark validation is required, but results.json has no benchmark_validation section.")
    if _is_required(quality_gates.get("uncertainty_analysis")) and not (
            results and results.get("uncertainty")):
        warnings.append("Uncertainty analysis is required, but results.json has no uncertainty section.")
    if _is_required(quality_gates.get("risk_register")) and not (
            results and (results.get("risk_evaluation") or results.get("risks"))):
        warnings.append("Risk register is required, but results.json has no risk_evaluation or risks section.")
    if _is_required(quality_gates.get("figure_discussion")) and not (
            results and results.get("figure_discussion")):
        warnings.append("Figure discussion is required, but results.json has no figure_discussion section.")
    if _is_required(quality_gates.get("consistency_checker")):
        consistency_path = os.path.join(TASK_DIR, "consistency_report.json")
        if not os.path.exists(consistency_path):
            warnings.append("Consistency checker is required, but consistency_report.json is missing.")

    minimum_figures = _as_int(quality_gates.get("minimum_figures"), 0)
    figure_count = len(glob.glob(os.path.join(FIG_DIR, "*.png")))
    if minimum_figures and figure_count < minimum_figures:
        warnings.append("Configured figure minimum is {}, but {} PNG figure(s) exist.".format(
            minimum_figures, figure_count))

    for section in report.get("required_sections", []):
        if not _required_section_available(section, results, task_spec):
            warnings.append("Required report section lacks source data: {}".format(section))

    return warnings


def _md_table_to_html(lines):
    """Convert markdown table lines to an HTML table string."""
    if len(lines) < 2:
        return ""
    # Parse header
    header_cells = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    # Skip separator line (line 1)
    html = '<table class="scope-table"><thead><tr>'
    for cell in header_cells:
        html += "<th>{}</th>".format(_md_inline(cell))
    html += "</tr></thead><tbody>\n"
    for row_line in lines[2:]:
        cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
        html += "<tr>"
        for cell in cells:
            html += "<td>{}</td>".format(_md_inline(cell))
        html += "</tr>\n"
    html += "</tbody></table>"
    return html


def _md_inline(text):
    """Convert inline markdown (bold) to HTML."""
    import re as _re
    # **bold**
    text = _re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    return text


def _md_list_to_html(lines):
    """Convert markdown bullet list lines to an HTML list."""
    html = "<ul>\n"
    for line in lines:
        item = line.lstrip("- ").strip()
        html += "  <li>{}</li>\n".format(_md_inline(item))
    html += "</ul>"
    return html


def scope_content_to_html(content):
    """Convert scope section content (from task_spec.md) to styled HTML.

    Handles markdown tables, bold text, bullet lists, and sub-headings.
    """
    lines = content.split("\n")
    html_parts = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Sub-heading (e.g., "Applicable Standards:")
        if (line.strip().endswith(":") and not line.strip().startswith("-")
                and not line.strip().startswith("|") and not line.strip().startswith("*")):
            html_parts.append("<h3>{}</h3>".format(_md_inline(line.strip())))
            i += 1
            continue

        # Markdown table (starts with |)
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # Check if second line is separator (|---|)
            if len(table_lines) >= 2 and set(table_lines[1].replace("|", "").replace("-", "").replace(":", "").strip()) <= set(""):
                html_parts.append(_md_table_to_html(table_lines))
            else:
                # Not a real table, just text with pipes
                for tl in table_lines:
                    html_parts.append("<p>{}</p>".format(_md_inline(tl.strip())))
            continue

        # Bullet list (starts with -)
        if line.strip().startswith("- "):
            list_lines = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                list_lines.append(lines[i])
                i += 1
            html_parts.append(_md_list_to_html(list_lines))
            continue

        # Regular paragraph
        html_parts.append("<p>{}</p>".format(_md_inline(line.strip())))
        i += 1

    return "\n".join(html_parts)


def render_scope_to_word(doc, content):
    """Render scope section content (from task_spec.md) into a Word document.

    Parses markdown tables into Word tables, bold text into runs, and
    bullet lists into formatted paragraphs.
    """
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Sub-heading (e.g., "Applicable Standards:")
        if (line.strip().endswith(":") and not line.strip().startswith("-")
                and not line.strip().startswith("|") and not line.strip().startswith("*")):
            doc.add_heading(line.strip(), level=2)
            i += 1
            continue

        # Markdown table
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                _md_table_to_word(doc, table_lines)
            else:
                for tl in table_lines:
                    doc.add_paragraph(tl.strip())
            continue

        # Bullet list
        if line.strip().startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                item_text = lines[i].strip()[2:]  # Remove "- "
                p = doc.add_paragraph(style="List Bullet")
                _add_bold_runs(p, item_text)
                i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_bold_runs(p, line.strip())
        i += 1


def _md_table_to_word(doc, table_lines):
    """Convert markdown table lines to a styled Word table."""
    header_cells = [c.strip() for c in table_lines[0].strip().strip("|").split("|")]
    data_rows = []
    for row_line in table_lines[2:]:  # skip header and separator
        cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
        data_rows.append(cells)
    add_word_table(doc, header_cells, data_rows)


def _add_bold_runs(paragraph, text):
    """Add text with **bold** sections as separate runs."""
    import re as _re
    parts = _re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def get_figures():
    """Collect all PNG/SVG figures from the figures/ directory."""
    pngs = sorted(glob.glob(os.path.join(FIG_DIR, "*.png")))
    svgs = sorted(glob.glob(os.path.join(FIG_DIR, "*.svg")))
    return pngs + svgs


def get_figure_caption(fig_path, results, fig_index):
    """Get a caption for a figure: custom from results.json or auto-generated."""
    fig_name = os.path.basename(fig_path)
    captions = {}
    if results:
        captions = results.get("figure_captions", {})
    if fig_name in captions:
        return "Figure {}: {}".format(fig_index, captions[fig_name])
    # Auto-generate from filename
    auto = fig_name.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").title()
    return "Figure {}: {}".format(fig_index, auto)


def get_equations(results):
    """Get equations from results.json. Returns list of {label, latex}."""
    if not results:
        return []
    return results.get("equations", [])


def render_equation_to_image(latex_str, output_path):
    """Render a LaTeX equation to a high-quality PNG image using matplotlib.

    Uses display-style math, large font, and 300 DPI for crisp rendering
    in Word documents. Returns True if the image was created, False otherwise.
    """
    if not HAS_MATPLOTLIB:
        return False
    try:
        fig = plt.figure(figsize=(8, 1.2))
        fig.text(
            0.5, 0.5,
            "${}$".format(latex_str),
            fontsize=24, ha="center", va="center",
            math_fontfamily="cm",
        )
        fig.savefig(output_path, dpi=300, bbox_inches="tight",
                    pad_inches=0.15, facecolor="white", edgecolor="none")
        plt.close(fig)
        return True
    except Exception as e:
        print("  Warning: could not render equation: {}".format(e))
        return False


def _parse_key_name(key):
    """Parse a key_results key into (label, unit). Splits on last known unit suffix."""
    unit_suffixes = [
        ("_pct", "%"), ("_percent", "%"),
        ("_bar", "bar"), ("_bara", "bara"), ("_barg", "barg"),
        ("_psi", "psi"), ("_psia", "psia"),
        ("_C", "°C"), ("_K", "K"), ("_F", "°F"),
        ("_kg", "kg"), ("_g", "g"), ("_lb", "lb"),
        ("_m3", "m³"), ("_m2", "m²"), ("_m", "m"),
        ("_mm", "mm"), ("_cm", "cm"), ("_km", "km"),
        ("_ft", "ft"), ("_in", "in"),
        ("_kW", "kW"), ("_MW", "MW"), ("_W", "W"),
        ("_kJ", "kJ"), ("_MJ", "MJ"), ("_J", "J"),
        ("_kg_hr", "kg/hr"), ("_kg_s", "kg/s"),
        ("_m3_hr", "m³/hr"), ("_m3_s", "m³/s"),
        ("_Sm3_day", "Sm³/day"), ("_Sm3_hr", "Sm³/hr"),
        ("_hours", "hours"), ("_hr", "hr"), ("_min", "min"), ("_s", "s"),
        ("_rpm", "rpm"), ("_Hz", "Hz"),
    ]
    for suffix, unit in unit_suffixes:
        if key.endswith(suffix):
            name_part = key[:len(key) - len(suffix)]
            label = name_part.replace("_", " ").title()
            return label, unit
    return key.replace("_", " ").title(), ""


def format_results_table(results):
    """Format key_results dict as a text table."""
    key_results = results.get("key_results", {})
    if not key_results:
        return "[No key_results in results.json]"
    lines = []
    for key, value in key_results.items():
        label, unit = _parse_key_name(key)
        if isinstance(value, float):
            val_str = "{:.4g}".format(value)
        else:
            val_str = str(value)
        if unit:
            lines.append("{}: {} {}".format(label, val_str, unit))
        else:
            lines.append("{}: {}".format(label, val_str))
    return "\n".join(lines)


def format_validation_table(results):
    """Format validation checks as a text table."""
    validation = results.get("validation", {})
    if not validation:
        return "[No validation data in results.json]"
    lines = ["Validation Summary:", ""]
    for check, outcome in validation.items():
        label = check.replace("_", " ").title()
        if isinstance(outcome, bool):
            status = "PASS" if outcome else "FAIL"
        elif isinstance(outcome, (int, float)):
            status = "{:.4g}".format(outcome)
        else:
            status = str(outcome)
        lines.append("  {}: {}".format(label, status))
    return "\n".join(lines)


def format_validation_html(results):
    """Format validation checks as an HTML table."""
    validation = results.get("validation", {})
    if not validation:
        return "<p><em>No validation data in results.json</em></p>"
    rows = ""
    for check, outcome in validation.items():
        label = check.replace("_", " ").title()
        if isinstance(outcome, bool):
            status = "PASS" if outcome else "FAIL"
            css_class = ' class="pass"' if outcome else ' class="fail"'
        elif isinstance(outcome, (int, float)):
            status = "{:.4g}".format(outcome)
            css_class = ""
        else:
            status = str(outcome)
            css_class = ""
        rows += '<tr><td>{}</td><td{}>{}</td></tr>\n'.format(
            label, css_class, status)
    return '<table class="validation-table"><thead><tr><th>Check</th><th>Result</th></tr></thead><tbody>\n{}</tbody></table>'.format(rows)


def format_results_html(results):
    """Format key_results dict as a styled HTML table with units column."""
    key_results = results.get("key_results", {})
    if not key_results:
        return ""
    rows = ""
    for key, value in key_results.items():
        label, unit = _parse_key_name(key)
        if isinstance(value, float):
            val_str = "{:.4g}".format(value)
        else:
            val_str = str(value)
        rows += '<tr><td>{}</td><td class="num">{}</td><td>{}</td></tr>\n'.format(
            label, val_str, unit)
    return (
        '<table class="results-table"><thead>'
        '<tr><th>Parameter</th><th>Value</th><th>Unit</th></tr>'
        '</thead><tbody>\n{}</tbody></table>'.format(rows)
    )


def format_custom_tables_html(results):
    """Format custom tables from results.json 'tables' key as HTML."""
    tables = results.get("tables", [])
    if not tables:
        return ""
    html_parts = []
    for tbl in tables:
        title = tbl.get("title", "")
        headers = tbl.get("headers", [])
        data_rows = tbl.get("rows", [])
        if not headers or not data_rows:
            continue
        h = ""
        if title:
            h += '<h3>{}</h3>\n'.format(title)
        h += '<table class="custom-table"><thead><tr>'
        for col in headers:
            h += '<th>{}</th>'.format(col)
        h += '</tr></thead><tbody>\n'
        for row in data_rows:
            h += "<tr>"
            for i, cell in enumerate(row):
                css = ' class="num"' if i > 0 and isinstance(cell, (int, float)) else ""
                if isinstance(cell, float):
                    h += '<td{}>{:.4g}</td>'.format(css, cell)
                else:
                    h += '<td{}>{}</td>'.format(css, cell)
            h += "</tr>\n"
        h += "</tbody></table>"
        html_parts.append(h)
    return "\n".join(html_parts)


def format_references_html(results):
    """Format the references list from results.json as a styled HTML ordered list."""
    refs = results.get("references", [])
    if not refs:
        return ""
    h = '<ol class="reference-list">\n'
    for ref in refs:
        ref_id = ref.get("id", "")
        ref_text = ref.get("text", "")
        if ref_id:
            h += '  <li id="ref-{}"><strong>[{}]</strong> {}</li>\n'.format(
                ref_id, ref_id, ref_text)
        else:
            h += '  <li>{}</li>\n'.format(ref_text)
    h += '</ol>'
    return h


def format_risk_html(results):
    """Format risk_evaluation from results.json as a styled HTML risk table.

    Renders a summary card with overall risk level, then a professional
    table with color-coded risk levels and mitigation measures.
    """
    re_data = results.get("risk_evaluation", {})
    if not re_data:
        return ""
    risks = re_data.get("risks", [])
    matrix = re_data.get("risk_matrix_used", "5x5 (ISO 31000)")
    overall = re_data.get("overall_risk_level", "N/A")
    high_count = re_data.get("high_risk_count",
                             sum(1 for r in risks if r.get("risk_level") == "High"))
    medium_count = re_data.get("medium_risk_count",
                               sum(1 for r in risks if r.get("risk_level") == "Medium"))
    low_count = sum(1 for r in risks if r.get("risk_level") == "Low")

    h = '<div class="risk-summary-card">\n'
    h += '<p>Risk assessment performed using <strong>{}</strong> framework.</p>\n'.format(
        matrix)
    h += '<p>Overall risk level: <span class="risk-badge risk-{}">{}</span>'.format(
        overall.lower(), overall)
    h += '&nbsp;&mdash;&nbsp;'
    parts = []
    if high_count:
        parts.append('<span class="risk-badge risk-high">{} High</span>'.format(high_count))
    if medium_count:
        parts.append('<span class="risk-badge risk-medium">{} Medium</span>'.format(medium_count))
    if low_count:
        parts.append('<span class="risk-badge risk-low">{} Low</span>'.format(low_count))
    h += ", ".join(parts) if parts else ""
    h += '</p>\n</div>\n'

    if risks:
        h += '<table class="risk-table"><thead><tr>'
        h += '<th>ID</th><th>Category</th><th>Risk Description</th>'
        h += '<th>Likelihood</th><th>Consequence</th><th>Risk Level</th>'
        h += '<th>Mitigation</th>'
        h += '</tr></thead><tbody>\n'
        for risk in risks:
            level = risk.get("risk_level", "").lower()
            h += '<tr>'
            h += '<td><strong>{}</strong></td>'.format(risk.get("id", ""))
            h += '<td>{}</td>'.format(risk.get("category", ""))
            h += '<td>{}</td>'.format(risk.get("description", ""))
            h += '<td>{}</td>'.format(risk.get("likelihood", ""))
            h += '<td>{}</td>'.format(risk.get("consequence", ""))
            h += '<td class="risk-level risk-{}">{}</td>'.format(level, risk.get("risk_level", ""))
            h += '<td class="mitigation-cell">{}</td>'.format(risk.get("mitigation", ""))
            h += '</tr>\n'
        h += '</tbody></table>\n'
    return h


def format_uncertainty_html(results):
    """Format uncertainty analysis from results.json as styled HTML tables.

    Renders: input parameters table, P10/P50/P90 output table, tornado table.
    """
    unc = results.get("uncertainty", {})
    if not unc:
        return ""
    h = ""

    # Summary paragraph
    h += '<div class="uncertainty-summary">\n'
    h += '<p><strong>{}</strong> with <strong>{}</strong> simulations'.format(
        unc.get("method", "Monte Carlo analysis"),
        unc.get("n_simulations", "N/A"))
    engine = unc.get("simulation_engine", "")
    if engine:
        h += ' using {}'.format(engine)
    h += '.</p>\n</div>\n'

    # Input parameters table
    params = unc.get("input_parameters", [])
    if params:
        h += '<h3>Input Parameter Ranges</h3>\n'
        h += '<table class="uncertainty-table"><thead><tr>'
        h += '<th>Parameter</th><th>Unit</th><th>Low</th><th>Base</th>'
        h += '<th>High</th><th>Distribution</th>'
        h += '</tr></thead><tbody>\n'
        for p in params:
            h += '<tr>'
            h += '<td>{}</td>'.format(p.get("name", ""))
            h += '<td>{}</td>'.format(p.get("unit", ""))
            h += '<td class="num">{}</td>'.format(_fmt_num(p.get("low", "")))
            h += '<td class="num">{}</td>'.format(_fmt_num(p.get("base", "")))
            h += '<td class="num">{}</td>'.format(_fmt_num(p.get("high", "")))
            h += '<td>{}</td>'.format(p.get("distribution", ""))
            h += '</tr>\n'
        h += '</tbody></table>\n'

    # Output P10/P50/P90 table — handle both single-output and multi-output formats
    out_param = unc.get("output_parameter", "")
    out_params = unc.get("output_parameters", {})
    if out_param or out_params:
        h += '<h3>Output Distribution (P10 / P50 / P90)</h3>\n'
        h += '<table class="uncertainty-table"><thead><tr>'
        h += '<th>Output Parameter</th><th>P10</th><th>P50</th><th>P90</th>'
        h += '</tr></thead><tbody>\n'
        if out_param:
            h += '<tr><td>{}</td>'.format(out_param)
            h += '<td class="num">{}</td>'.format(_fmt_num(unc.get("p10", "")))
            h += '<td class="num"><strong>{}</strong></td>'.format(_fmt_num(unc.get("p50", "")))
            h += '<td class="num">{}</td></tr>\n'.format(_fmt_num(unc.get("p90", "")))
        for key, val in out_params.items():
            label = key.replace("_", " ").title()
            h += '<tr><td>{}</td>'.format(label)
            h += '<td class="num">{}</td>'.format(_fmt_num(val.get("p10", "")))
            h += '<td class="num"><strong>{}</strong></td>'.format(_fmt_num(val.get("p50", "")))
            h += '<td class="num">{}</td></tr>\n'.format(_fmt_num(val.get("p90", "")))
        h += '</tbody></table>\n'

    # Tornado sensitivity table
    tornado = unc.get("tornado", [])
    if tornado:
        h += '<h3>Sensitivity Ranking (Tornado)</h3>\n'
        h += '<table class="tornado-table"><thead><tr>'
        # Detect column names from first tornado entry
        first = tornado[0]
        cols = [k for k in first.keys() if k != "parameter"]
        h += '<th>Parameter</th>'
        for col in cols:
            h += '<th>{}</th>'.format(col.replace("_", " ").title())
        h += '</tr></thead><tbody>\n'
        for item in tornado:
            h += '<tr><td>{}</td>'.format(item.get("parameter", ""))
            for col in cols:
                h += '<td class="num">{}</td>'.format(_fmt_num(item.get(col, "")))
            h += '</tr>\n'
        h += '</tbody></table>\n'

    return h


def format_benchmark_html(results):
    """Format benchmark_validation from results.json as a styled HTML table."""
    bv = results.get("benchmark_validation", {})
    if not bv:
        return ""
    h = '<table class="benchmark-table"><thead><tr>'
    h += '<th>Test</th><th>Description</th><th>Status</th><th>Details</th>'
    h += '</tr></thead><tbody>\n'
    for key, val in bv.items():
        label = key.replace("_", " ").title()
        desc = val.get("description", key)
        status = val.get("status", "N/A")
        css = ' class="pass"' if status == "PASS" else (' class="fail"' if status == "FAIL" else "")
        # Gather numeric details
        details = []
        for dk, dv in val.items():
            if dk in ("description", "status"):
                continue
            dl = dk.replace("_", " ").title()
            if isinstance(dv, float):
                details.append("{}: {:.4g}".format(dl, dv))
            else:
                details.append("{}: {}".format(dl, dv))
        h += '<tr>'
        h += '<td><strong>{}</strong></td>'.format(label)
        h += '<td>{}</td>'.format(desc)
        h += '<td{}><strong>{}</strong></td>'.format(css, status)
        h += '<td>{}</td>'.format("; ".join(details) if details else "")
        h += '</tr>\n'
    h += '</tbody></table>\n'
    return h


def _fmt_num(value):
    """Format a numeric value for display in tables."""
    if isinstance(value, float):
        if abs(value) >= 1000 or (abs(value) < 0.01 and value != 0):
            return "{:.4g}".format(value)
        return "{:.4g}".format(value)
    return str(value)


def add_risk_word_table(doc, results):
    """Add risk evaluation as a styled Word table with color-coded risk levels."""
    re_data = results.get("risk_evaluation", {})
    if not re_data:
        return
    risks = re_data.get("risks", [])
    matrix = re_data.get("risk_matrix_used", "5x5 (ISO 31000)")
    overall = re_data.get("overall_risk_level", "N/A")
    high_count = re_data.get("high_risk_count",
                             sum(1 for r in risks if r.get("risk_level") == "High"))
    medium_count = re_data.get("medium_risk_count",
                               sum(1 for r in risks if r.get("risk_level") == "Medium"))

    # Summary paragraph
    p = doc.add_paragraph()
    p.add_run("Risk assessment using ").font.size = Pt(10)
    r = p.add_run("{} framework".format(matrix))
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(". Overall risk level: ").font.size = Pt(10)
    r2 = p.add_run(overall)
    r2.bold = True
    r2.font.size = Pt(10)
    if overall == "High":
        r2.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)
    elif overall == "Medium":
        r2.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    elif overall == "Low":
        r2.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)
    p.add_run(". ({} High, {} Medium)".format(high_count, medium_count)).font.size = Pt(10)

    if not risks:
        return
    headers = ["ID", "Category", "Description", "Likelihood", "Consequence",
               "Risk Level", "Mitigation"]
    data_rows = []
    for risk in risks:
        data_rows.append([
            risk.get("id", ""),
            risk.get("category", ""),
            risk.get("description", ""),
            risk.get("likelihood", ""),
            risk.get("consequence", ""),
            risk.get("risk_level", ""),
            risk.get("mitigation", ""),
        ])
    table = add_word_table(doc, headers, data_rows)
    # Color-code risk level column (column 5, 0-indexed)
    for row in table.rows[1:]:
        cell = row.cells[5]
        text = cell.text.strip()
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                if text == "High":
                    run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)
                elif text == "Medium":
                    run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
                elif text == "Low":
                    run.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)


def add_uncertainty_word_tables(doc, results):
    """Add uncertainty analysis as styled Word tables."""
    unc = results.get("uncertainty", {})
    if not unc:
        return

    # Summary paragraph
    p = doc.add_paragraph()
    p.add_run("{} with {} simulations".format(
        unc.get("method", "Monte Carlo analysis"),
        unc.get("n_simulations", "N/A"))).font.size = Pt(10)
    engine = unc.get("simulation_engine", "")
    if engine:
        p.add_run(" using {}.".format(engine)).font.size = Pt(10)

    # Input parameters table
    params = unc.get("input_parameters", [])
    if params:
        doc.add_heading("Input Parameter Ranges", level=2)
        headers = ["Parameter", "Unit", "Low", "Base", "High", "Distribution"]
        data_rows = []
        for param in params:
            data_rows.append([
                param.get("name", ""),
                param.get("unit", ""),
                _fmt_num(param.get("low", "")),
                _fmt_num(param.get("base", "")),
                _fmt_num(param.get("high", "")),
                param.get("distribution", ""),
            ])
        add_word_table(doc, headers, data_rows)

    # Output P10/P50/P90 table
    out_param = unc.get("output_parameter", "")
    out_params = unc.get("output_parameters", {})
    if out_param or out_params:
        doc.add_heading("Output Distribution (P10 / P50 / P90)", level=2)
        headers = ["Output Parameter", "P10", "P50", "P90"]
        data_rows = []
        if out_param:
            data_rows.append([
                out_param,
                _fmt_num(unc.get("p10", "")),
                _fmt_num(unc.get("p50", "")),
                _fmt_num(unc.get("p90", "")),
            ])
        for key, val in out_params.items():
            label = key.replace("_", " ").title()
            data_rows.append([
                label,
                _fmt_num(val.get("p10", "")),
                _fmt_num(val.get("p50", "")),
                _fmt_num(val.get("p90", "")),
            ])
        add_word_table(doc, headers, data_rows,
                       col_widths=[Inches(2.5), Inches(1.0), Inches(1.0), Inches(1.0)])

    # Tornado sensitivity table
    tornado = unc.get("tornado", [])
    if tornado:
        doc.add_heading("Sensitivity Ranking (Tornado)", level=2)
        first = tornado[0]
        cols = [k for k in first.keys() if k != "parameter"]
        headers = ["Parameter"] + [c.replace("_", " ").title() for c in cols]
        data_rows = []
        for item in tornado:
            row = [item.get("parameter", "")]
            for col in cols:
                row.append(_fmt_num(item.get(col, "")))
            data_rows.append(row)
        add_word_table(doc, headers, data_rows)


def add_benchmark_word_table(doc, results):
    """Add benchmark validation as a styled Word table."""
    bv = results.get("benchmark_validation", {})
    if not bv:
        return
    headers = ["Test", "Description", "Status", "Details"]
    data_rows = []
    for key, val in bv.items():
        label = key.replace("_", " ").title()
        desc = val.get("description", key)
        status = val.get("status", "N/A")
        details = []
        for dk, dv in val.items():
            if dk in ("description", "status"):
                continue
            dl = dk.replace("_", " ").title()
            if isinstance(dv, float):
                details.append("{}: {:.4g}".format(dl, dv))
            else:
                details.append("{}: {}".format(dl, dv))
        data_rows.append([label, desc, status, "; ".join(details)])
    table = add_word_table(doc, headers, data_rows)
    # Color-code status column (column 2, 0-indexed)
    for row in table.rows[1:]:
        cell = row.cells[2]
        text = cell.text.strip()
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                if text == "PASS":
                    run.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)
                elif text == "FAIL":
                    run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)


def add_word_table(doc, headers, data_rows, col_widths=None):
    """Add a professionally styled table to a Word document.

    Args:
        doc: Document object.
        headers: list of column header strings.
        data_rows: list of lists (each inner list = one row of cell values).
        col_widths: optional list of Inches widths per column.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(text)
        # Style header: bold, white text on dark blue background
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(
            '<w:shd {} w:fill="2F5496"/>'.format(nsdecls('w'))
        )
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for row_data in data_rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Apply column widths if specified
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    doc.add_paragraph("")  # spacing after table
    return table


def add_results_word_table(doc, results):
    """Add key_results as a styled Word table with units column."""
    key_results = results.get("key_results", {})
    if not key_results:
        return
    headers = ["Parameter", "Value", "Unit"]
    data_rows = []
    for key, value in key_results.items():
        label, unit = _parse_key_name(key)
        if isinstance(value, float):
            val_str = "{:.4g}".format(value)
        else:
            val_str = str(value)
        data_rows.append([label, val_str, unit])
    add_word_table(doc, headers, data_rows,
                   col_widths=[Inches(3.0), Inches(1.5), Inches(1.5)])


def add_validation_word_table(doc, results):
    """Add validation checks as a styled Word table."""
    validation = results.get("validation", {})
    if not validation:
        return
    headers = ["Check", "Result"]
    data_rows = []
    for check, outcome in validation.items():
        label = check.replace("_", " ").title()
        if isinstance(outcome, bool):
            status = "PASS" if outcome else "FAIL"
        elif isinstance(outcome, (int, float)):
            status = "{:.4g}".format(outcome)
        else:
            status = str(outcome)
        data_rows.append([label, status])
    table = add_word_table(doc, headers, data_rows,
                           col_widths=[Inches(4.0), Inches(2.0)])
    # Color-code PASS/FAIL cells
    for row in table.rows[1:]:
        cell = row.cells[1]
        text = cell.text.strip()
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                if text == "PASS":
                    run.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)
                elif text == "FAIL":
                    run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)


def add_custom_word_tables(doc, results):
    """Add custom tables from results.json 'tables' key."""
    tables = results.get("tables", [])
    if not tables:
        return
    for tbl in tables:
        title = tbl.get("title", "")
        headers = tbl.get("headers", [])
        data_rows = tbl.get("rows", [])
        if not headers or not data_rows:
            continue
        if title:
            doc.add_heading(title, level=2)
        # Format numeric values
        formatted_rows = []
        for row in data_rows:
            formatted = []
            for cell in row:
                if isinstance(cell, float):
                    formatted.append("{:.4g}".format(cell))
                else:
                    formatted.append(str(cell))
            formatted_rows.append(formatted)
        add_word_table(doc, headers, formatted_rows)


def format_discussion_html(results):
    """Format figure_discussion from results.json as styled HTML discussion blocks.

    Each discussion entry links a figure to its observation, physical mechanism,
    engineering implication, and recommendation — creating traceability from
    calculation to conclusion.
    """
    discussions = results.get("figure_discussion", [])
    if not discussions:
        return ""
    h = ""
    for i, disc in enumerate(discussions, 1):
        fig_file = disc.get("figure", "")
        title = disc.get("title", fig_file.replace("_", " ").replace(".png", "").title())
        observation = disc.get("observation", "")
        mechanism = disc.get("mechanism", "")
        implication = disc.get("implication", "")
        recommendation = disc.get("recommendation", "")
        linked = disc.get("linked_results", [])
        insight_ref = disc.get("insight_question_ref", "")

        h += '<div class="discussion-block">\n'
        h += '<h3>Discussion {}: {}</h3>\n'.format(i, title)
        if observation:
            h += '<p><strong>Observation:</strong> {}</p>\n'.format(observation)
        if mechanism:
            h += '<p><strong>Physical Mechanism:</strong> {}</p>\n'.format(mechanism)
        if implication:
            h += '<p><strong>Engineering Implication:</strong> {}</p>\n'.format(implication)
        if recommendation:
            h += '<p class="recommendation"><strong>Recommendation:</strong> {}</p>\n'.format(
                recommendation)
        # Traceability footer
        trace_parts = []
        if linked:
            trace_parts.append("Linked results: {}".format(", ".join(linked)))
        if insight_ref:
            trace_parts.append("Answers: {}".format(insight_ref))
        if trace_parts:
            h += '<p class="traceability"><em>{}</em></p>\n'.format(" | ".join(trace_parts))
        h += '</div>\n'
    return h


def add_discussion_word(doc, results):
    """Add figure discussion entries as styled Word content.

    Each discussion block has: observation, physical mechanism, engineering
    implication, and recommendation with traceability references.
    """
    discussions = results.get("figure_discussion", [])
    if not discussions:
        return
    for i, disc in enumerate(discussions, 1):
        fig_file = disc.get("figure", "")
        title = disc.get("title", fig_file.replace("_", " ").replace(".png", "").title())
        observation = disc.get("observation", "")
        mechanism = disc.get("mechanism", "")
        implication = disc.get("implication", "")
        recommendation = disc.get("recommendation", "")
        linked = disc.get("linked_results", [])
        insight_ref = disc.get("insight_question_ref", "")

        doc.add_heading("Discussion {}: {}".format(i, title), level=2)

        if observation:
            p = doc.add_paragraph()
            r = p.add_run("Observation: ")
            r.bold = True
            r.font.size = Pt(10)
            p.add_run(observation).font.size = Pt(10)

        if mechanism:
            p = doc.add_paragraph()
            r = p.add_run("Physical Mechanism: ")
            r.bold = True
            r.font.size = Pt(10)
            p.add_run(mechanism).font.size = Pt(10)

        if implication:
            p = doc.add_paragraph()
            r = p.add_run("Engineering Implication: ")
            r.bold = True
            r.font.size = Pt(10)
            p.add_run(implication).font.size = Pt(10)

        if recommendation:
            p = doc.add_paragraph()
            r = p.add_run("Recommendation: ")
            r.bold = True
            r.font.size = Pt(10)
            r2 = p.add_run(recommendation)
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x1A, 0x53, 0x7A)

        # Traceability line
        trace_parts = []
        if linked:
            trace_parts.append("Linked results: {}".format(", ".join(linked)))
        if insight_ref:
            trace_parts.append("Answers: {}".format(insight_ref))
        if trace_parts:
            p = doc.add_paragraph()
            r = p.add_run(" | ".join(trace_parts))
            r.font.size = Pt(8)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")  # spacing


# ══════════════════════════════════════════════════════════
# Build sections (auto-populated where possible)
# ══════════════════════════════════════════════════════════

def build_sections(results, task_spec, study_config_warnings=None):
    """Build report sections, auto-populating from results.json and task_spec.md."""
    sections = []
    if study_config_warnings is None:
        study_config_warnings = []

    # 1. Executive Summary
    sections.append({
        "heading": "1. Executive Summary",
        "content": MANUAL_SECTIONS["executive_summary"],
    })

    # 2. Problem Description
    sections.append({
        "heading": "2. Problem Description",
        "content": MANUAL_SECTIONS["problem_description"],
    })

    # 3. Scope and Standards (auto-populated from task_spec.md)
    scope_parts = []
    standards = extract_spec_section(task_spec, "Applicable Standards")
    if standards:
        scope_parts.append("Applicable Standards:\n" + standards)
    methods = extract_spec_section(task_spec, "Calculation Methods")
    if methods:
        scope_parts.append("Calculation Methods:\n" + methods)
    criteria = extract_spec_section(task_spec, "Acceptance Criteria")
    if criteria:
        scope_parts.append("Acceptance Criteria:\n" + criteria)
    envelope = extract_spec_section(task_spec, "Operating Envelope")
    if envelope:
        scope_parts.append("Operating Envelope:\n" + envelope)

    scope_content = "\n\n".join(scope_parts) if scope_parts else (
        "[Auto-populated from task_spec.md when filled in. "
        "Edit step1_scope_and_research/task_spec.md and re-run.]"
    )
    sections.append({
        "heading": "3. Scope and Standards",
        "content": scope_content,
        "has_scope": True,
    })

    # 4. Approach
    approach = MANUAL_SECTIONS["approach"]
    if results and results.get("approach"):
        approach = results["approach"]
    sections.append({
        "heading": "4. Approach",
        "content": approach,
        "has_equations": True,
    })

    # 5. Results (auto-populated from results.json)
    if results and results.get("key_results"):
        results_text = format_results_table(results)
    else:
        results_text = (
            "[Auto-populated from results.json when created by notebook. "
            "Save results with the pattern shown in the task README.]"
        )
    sections.append({
        "heading": "5. Results",
        "content": results_text,
        "has_figures": True,
    })

    # 6. Discussion (auto-populated from results.json figure_discussion)
    if results and results.get("figure_discussion"):
        sections.append({
            "heading": "6. Discussion",
            "content": "",
            "has_discussion": True,
        })
        next_validation_num = 7
    else:
        next_validation_num = 6

    # N. Validation Summary (auto-populated from results.json)
    if results and results.get("validation"):
        validation_text = format_validation_table(results)
    else:
        validation_text = (
            "[Auto-populated from results.json validation section. "
            "Add validation checks to your notebook results output.]"
        )
    sections.append({
        "heading": "{}. Validation Summary".format(next_validation_num),
        "content": validation_text,
    })

    # N+1. Benchmark Validation (if data available)
    next_num = next_validation_num + 1

    if study_config_warnings:
        warning_lines = ["- {}".format(warning) for warning in study_config_warnings]
        sections.append({
            "heading": "{}. Study Configuration Warnings".format(next_num),
            "content": "\n".join(warning_lines),
        })
        next_num += 1

    if results and results.get("benchmark_validation"):
        sections.append({
            "heading": "{}. Benchmark Validation".format(next_num),
            "content": "",
            "has_benchmark": True,
        })
        next_num += 1

    # N. Uncertainty Analysis (if data available)
    if results and results.get("uncertainty"):
        sections.append({
            "heading": "{}. Uncertainty Analysis".format(next_num),
            "content": "",
            "has_uncertainty": True,
        })
        next_num += 1

    # N. Risk Assessment (if data available)
    if results and results.get("risk_evaluation"):
        sections.append({
            "heading": "{}. Risk Assessment".format(next_num),
            "content": "",
            "has_risk": True,
        })
        next_num += 1

    # N. Conclusions and Recommendations
    conclusions = MANUAL_SECTIONS["conclusions"]
    if results and results.get("conclusions"):
        conclusions = results["conclusions"]
    sections.append({
        "heading": "{}. Conclusions and Recommendations".format(next_num),
        "content": conclusions,
    })
    next_num += 1

    # N. References (auto-populated from results.json if available)
    refs_content = MANUAL_SECTIONS["references"]
    if results and results.get("references"):
        ref_lines = []
        for i, ref in enumerate(results["references"], 1):
            ref_id = ref.get("id", "")
            ref_text = ref.get("text", "")
            if ref_id:
                ref_lines.append("[{}] {}".format(i, ref_text))
            else:
                ref_lines.append("[{}] {}".format(i, ref_text))
        refs_content = "\n".join(ref_lines)
    sections.append({
        "heading": "{}. References".format(next_num),
        "content": refs_content,
        "has_references": True,
    })

    return sections


# ══════════════════════════════════════════════════════════
# Word report
# ══════════════════════════════════════════════════════════

def _auto_doc_number():
    """Generate a document number from the task folder name if DOC_NUMBER is blank."""
    if DOC_NUMBER:
        return DOC_NUMBER
    folder_name = os.path.basename(TASK_DIR)
    # Extract date prefix (YYYY-MM-DD) if present
    if len(folder_name) >= 10 and folder_name[4] == "-" and folder_name[7] == "-":
        date_part = folder_name[:10].replace("-", "")
        return "REP-{}".format(date_part)
    return "REP-DRAFT"


def _add_cover_page(doc):
    """Add a professional cover page with title, metadata, and revision history."""
    doc_num = _auto_doc_number()

    # Spacer
    for _ in range(4):
        doc.add_paragraph("")

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(TITLE)
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(47, 84, 150)
    run.bold = True

    doc.add_paragraph("")

    # Subtitle line
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("NeqSim Engineering Report")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    for _ in range(3):
        doc.add_paragraph("")

    # Metadata table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document Number", doc_num),
        ("Revision", REVISION),
        ("Date", TASK_DATE),
        ("Author", AUTHOR or "(not specified)"),
        ("Classification", CLASSIFICATION),
    ]
    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        for cell in meta_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
        # Bold the label column
        for run in meta_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph("")

    # Revision history table (if entries exist)
    rev_entries = REVISION_HISTORY or [
        {"rev": REVISION, "date": TASK_DATE,
         "description": "Initial issue", "author": AUTHOR or ""}
    ]
    rev_heading = doc.add_paragraph()
    rev_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = rev_heading.add_run("Revision History")
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)

    rev_table = doc.add_table(rows=1 + len(rev_entries), cols=4)
    rev_table.style = "Table Grid"
    rev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Rev", "Date", "Description", "Author"]
    for j, h in enumerate(headers):
        cell = rev_table.rows[0].cells[j]
        cell.text = h
        shading = parse_xml(
            '<w:shd {} w:fill="2F5496"/>'.format(nsdecls("w"))
        )
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
        cell._tc.get_or_add_tcPr().append(shading)
    for i, entry in enumerate(rev_entries, 1):
        rev_table.rows[i].cells[0].text = str(entry.get("rev", ""))
        rev_table.rows[i].cells[1].text = str(entry.get("date", ""))
        rev_table.rows[i].cells[2].text = str(entry.get("description", ""))
        rev_table.rows[i].cells[3].text = str(entry.get("author", ""))

    doc.add_page_break()


def _add_word_toc(doc):
    """Add a Table of Contents field to the Word document."""
    # Add TOC heading
    doc.add_heading("Table of Contents", level=1)
    # Insert a Word TOC field (updates when user presses F9 in Word)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = parse_xml(
        '<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls("w"))
    )
    run._r.append(fldChar1)
    run2 = paragraph.add_run()
    instrText = parse_xml(
        '<w:instrText {} xml:space="preserve"> TOC \\o "1-2" \\h \\z \\u </w:instrText>'.format(
            nsdecls("w")
        )
    )
    run2._r.append(instrText)
    run3 = paragraph.add_run()
    fldChar2 = parse_xml(
        '<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls("w"))
    )
    run3._r.append(fldChar2)
    run4 = paragraph.add_run("(Right-click and select 'Update Field' to populate)")
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.italic = True
    run5 = paragraph.add_run()
    fldChar3 = parse_xml(
        '<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls("w"))
    )
    run5._r.append(fldChar3)
    doc.add_page_break()


def build_word_report(sections, results=None):
    """Build the Word document with cover page, TOC, numbered figures, and equations."""
    doc = Document()

    # Cover page with metadata and revision history
    _add_cover_page(doc)

    # Table of Contents
    _add_word_toc(doc)

    # Add all sections
    for section in sections:
        doc.add_heading(section["heading"], level=1)

        # Results section: use Word table instead of plain text
        if section.get("has_figures") and results and results.get("key_results"):
            add_results_word_table(doc, results)
            # Custom tables
            if results.get("tables"):
                add_custom_word_tables(doc, results)
        elif section.get("has_figures"):
            # No results data — show placeholder text
            for para_text in section["content"].split("\n\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        elif section.get("has_scope"):
            # Scope section: parse markdown tables, bold, and lists
            render_scope_to_word(doc, section["content"])
        elif "Validation" in section["heading"] and results and results.get("validation"):
            # Validation section: use Word table
            add_validation_word_table(doc, results)
        elif section.get("has_benchmark") and results:
            # Benchmark Validation section: styled table
            add_benchmark_word_table(doc, results)
        elif section.get("has_uncertainty") and results:
            # Uncertainty Analysis section: styled tables
            add_uncertainty_word_tables(doc, results)
        elif section.get("has_risk") and results:
            # Risk Assessment section: styled table with color-coded levels
            add_risk_word_table(doc, results)
        elif section.get("has_discussion") and results:
            # Discussion section: figure-by-figure interpretation
            add_discussion_word(doc, results)
        else:
            # Regular text content
            for para_text in section["content"].split("\n\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

        # Embed figures after Results section
        if section.get("has_figures"):
            figures = get_figures()
            if figures:
                for fig_idx, fig_path in enumerate(figures, 1):
                    caption_text = get_figure_caption(fig_path, results, fig_idx)
                    doc.add_picture(fig_path, width=Inches(6.0))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = doc.add_paragraph(caption_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].font.size = Pt(9)
                    caption.runs[0].font.italic = True
                    doc.add_paragraph("")
            else:
                doc.add_paragraph(
                    "[No figures found in figures/ directory. "
                    "Save plots as PNG files there and re-run this script.]"
                )

        # Embed equations after Approach section
        if section.get("has_equations"):
            equations = get_equations(results)
            if equations:
                doc.add_heading("Key Equations", level=2)
                eq_img_dir = os.path.join(REPORT_DIR, "_eq_images")
                if not os.path.exists(eq_img_dir):
                    os.makedirs(eq_img_dir)
                for eq_idx, eq in enumerate(equations, 1):
                    label = eq.get("label", "Equation {}".format(eq_idx))
                    latex = eq.get("latex", "")
                    if not latex:
                        continue
                    # Try to render equation as image
                    eq_img_path = os.path.join(eq_img_dir, "eq_{}.png".format(eq_idx))
                    if render_equation_to_image(latex, eq_img_path):
                        doc.add_paragraph("")
                        doc.add_picture(eq_img_path, width=Inches(5.5))
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption = doc.add_paragraph(
                            "Equation {}: {}".format(eq_idx, label)
                        )
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.runs[0].font.size = Pt(9)
                        caption.runs[0].font.italic = True
                    else:
                        # Fallback: text representation
                        doc.add_paragraph("{}: {}".format(label, latex))
                    doc.add_paragraph("")

    # Save
    doc.save(DOCX_FILE)
    print("Word report saved: {}".format(DOCX_FILE))


# ══════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════
# HTML report
# ══════════════════════════════════════════════════════════

def _build_rev_rows_html():
    """Build HTML table rows for revision history in the HTML report."""
    rev_entries = REVISION_HISTORY or [
        {"rev": REVISION, "date": TASK_DATE,
         "description": "Initial issue", "author": AUTHOR or ""}
    ]
    rows = ""
    for entry in rev_entries:
        rows += "<tr><td>{}</td><td>{}</td><td>{}</td><td>{}</td></tr>\n".format(
            entry.get("rev", ""), entry.get("date", ""),
            entry.get("description", ""), entry.get("author", ""))
    return rows


def build_html_report(sections, results=None):
    """Build an HTML report with embedded figures, KaTeX equations, and navigation."""
    figures = get_figures()

    # Build figure HTML with base64-embedded images and numbered captions
    figure_html = ""
    if figures:
        for fig_idx, fig_path in enumerate(figures, 1):
            fig_name = os.path.basename(fig_path)
            caption_text = get_figure_caption(fig_path, results, fig_idx)
            # Determine MIME type
            if fig_path.endswith(".svg"):
                mime = "image/svg+xml"
            else:
                mime = "image/png"
            with open(fig_path, "rb") as f:
                img_data = base64.b64encode(f.read()).decode("utf-8")
            figure_html += """
            <div class="figure">
                <img src="data:{};base64,{}" alt="{}">
                <p class="caption">{}</p>
            </div>
            """.format(mime, img_data, caption_text, caption_text)
    else:
        figure_html = "<p><em>No figures found in figures/ directory.</em></p>"

    # Build equation HTML (KaTeX rendering with embedded image fallbacks)
    equation_html = ""
    equations = get_equations(results)
    if equations:
        equation_html += '<h3>Key Equations</h3>\n'
        # Pre-render equation images for offline fallback
        eq_img_dir = os.path.join(REPORT_DIR, "_eq_images")
        if not os.path.exists(eq_img_dir):
            os.makedirs(eq_img_dir)
        for eq_idx, eq in enumerate(equations, 1):
            label = eq.get("label", "Equation {}".format(eq_idx))
            latex = eq.get("latex", "")
            if not latex:
                continue
            # Render fallback image
            fallback_img = ""
            eq_img_path = os.path.join(eq_img_dir, "eq_{}.png".format(eq_idx))
            if render_equation_to_image(latex, eq_img_path):
                with open(eq_img_path, "rb") as imgf:
                    img_b64 = base64.b64encode(imgf.read()).decode("utf-8")
                fallback_img = (
                    '<img class="eq-fallback" '
                    'src="data:image/png;base64,{}" '
                    'alt="{}" style="display:none; max-width:90%;">'.format(
                        img_b64, label)
                )
            equation_html += """
            <div class="equation-block">
                <div class="equation katex-eq">$${}$$</div>
                {}
                <p class="equation-label">Equation {}: {}</p>
            </div>
            """.format(latex, fallback_img, eq_idx, label)

    # Build validation HTML
    validation_html = ""
    if results and results.get("validation"):
        validation_html = format_validation_html(results)

    # Build key results HTML table
    results_table_html = ""
    if results and results.get("key_results"):
        results_table_html = format_results_html(results)

    # Build custom tables HTML
    custom_tables_html = ""
    if results and results.get("tables"):
        custom_tables_html = format_custom_tables_html(results)

    # Build section HTML and navigation
    nav_items = ""
    section_html = ""
    for section in sections:
        section_id = section["heading"].lower().replace(" ", "-").replace(".", "")
        nav_items += '    <li><a href="#{}">{}</a></li>\n'.format(
            section_id, section["heading"]
        )
        # Convert scope section markdown to HTML
        if section.get("has_scope"):
            content = scope_content_to_html(section["content"])
        else:
            content = section["content"].replace("\n", "<br>")

        # Insert auto-generated HTML for special sections
        if section.get("has_figures"):
            if results_table_html:
                content = results_table_html + custom_tables_html + figure_html
            else:
                content += figure_html

        if section.get("has_equations") and equation_html:
            content += equation_html

        if "Validation" in section["heading"] and validation_html:
            content = validation_html

        if section.get("has_benchmark") and results:
            content = format_benchmark_html(results)

        if section.get("has_uncertainty") and results:
            content = format_uncertainty_html(results)

        if section.get("has_risk") and results:
            content = format_risk_html(results)

        if section.get("has_discussion") and results:
            content = format_discussion_html(results)

        if section.get("has_references") and results and results.get("references"):
            content = format_references_html(results)

        section_html += """
        <section id="{}">
            <h2>{}</h2>
            <div>{}</div>
        </section>
        """.format(section_id, section["heading"], content)

    # KaTeX CDN for equation rendering (only if equations exist)
    katex_head = ""
    katex_body_script = ""
    if equations:
        katex_head = """
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.css">
    <script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.js"></script>
    <script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/contrib/auto-render.min.js"></script>"""
        katex_body_script = """
    <script>
        document.addEventListener("DOMContentLoaded", function() {
            if (typeof renderMathInElement === "function") {
                renderMathInElement(document.body, {
                    delimiters: [
                        {left: "$$", right: "$$", display: true},
                        {left: "$", right: "$", display: false}
                    ],
                    throwOnError: false
                });
            } else {
                // KaTeX not available (offline) — show fallback images
                var eqs = document.querySelectorAll(".katex-eq");
                for (var i = 0; i < eqs.length; i++) {
                    eqs[i].style.display = "none";
                }
                var imgs = document.querySelectorAll(".eq-fallback");
                for (var j = 0; j < imgs.length; j++) {
                    imgs[j].style.display = "inline";
                }
            }
        });
    </script>"""

    html = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>{katex_head}
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
               display: flex; line-height: 1.6; color: #333; }}
        nav {{ width: 260px; min-height: 100vh; background: #f5f5f5; padding: 1.5rem;
              position: fixed; overflow-y: auto; border-right: 1px solid #ddd; }}
        nav h3 {{ margin-bottom: 1rem; color: #555; font-size: 0.9rem;
                  text-transform: uppercase; letter-spacing: 0.05em; }}
        nav ul {{ list-style: none; }}
        nav li {{ margin-bottom: 0.5rem; }}
        nav a {{ color: #0366d6; text-decoration: none; font-size: 0.9rem; }}
        nav a:hover {{ text-decoration: underline; }}
        main {{ margin-left: 260px; max-width: 900px; padding: 2rem 3rem; }}
        h1 {{ margin-bottom: 0.5rem; color: #1a1a1a; }}
        h2 {{ margin-top: 2rem; margin-bottom: 1rem; color: #1a1a1a;
             border-bottom: 1px solid #eee; padding-bottom: 0.3rem; }}
        h3 {{ margin-top: 1.5rem; margin-bottom: 0.5rem; color: #333; }}
        .meta {{ color: #666; margin-bottom: 2rem; }}
        .cover-page {{ text-align: center; padding: 3rem 0; margin-bottom: 2rem;
                       border-bottom: 3px solid #2F5496; }}
        .cover-page h1 {{ font-size: 2.2rem; color: #2F5496; margin-bottom: 0.5rem; }}
        .cover-page .subtitle {{ font-size: 1.1rem; color: #888; margin-bottom: 2rem; }}
        .cover-meta {{ display: inline-block; text-align: left; margin: 1rem auto;
                       background: #f8f9fa; padding: 1rem 2rem; border-radius: 6px;
                       border: 1px solid #e0e0e0; }}
        .cover-meta td {{ padding: 0.2rem 0.8rem; }}
        .cover-meta td:first-child {{ font-weight: bold; color: #555; }}
        .rev-table {{ margin: 1rem auto; font-size: 0.9rem; max-width: 700px; }}
        .rev-table th {{ background: #2F5496; color: #fff; padding: 0.4rem 0.8rem; }}
        .rev-table td {{ padding: 0.3rem 0.8rem; border: 1px solid #e0e0e0; }}
        section {{ margin-bottom: 2rem; }}
        .figure {{ text-align: center; margin: 1.5rem 0; }}
        .figure img {{ max-width: 100%; border: 1px solid #ddd; border-radius: 4px; }}
        .caption {{ font-size: 0.85rem; color: #666; font-style: italic;
                    margin-top: 0.3rem; }}
        .equation-block {{ margin: 1.5rem 0; text-align: center; }}
        .equation {{ font-size: 1.2rem; padding: 0.5rem 0; }}
        .equation-label {{ font-size: 0.85rem; color: #666; font-style: italic;
                           margin-top: 0.2rem; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1.5rem 0;
                font-size: 0.92rem; box-shadow: 0 1px 3px rgba(0,0,0,0.08); }}
        thead th {{ background: #2F5496; color: #fff; font-weight: 600;
                    padding: 0.6rem 0.75rem; text-align: left;
                    border: 1px solid #2a4a85; }}
        tbody td {{ border: 1px solid #e0e0e0; padding: 0.5rem 0.75rem;
                    text-align: left; }}
        tbody tr:nth-child(even) {{ background: #f8f9fa; }}
        tbody tr:hover {{ background: #e9ecef; }}
        td.num {{ text-align: right; font-variant-numeric: tabular-nums; }}
        .pass {{ color: #28a745; font-weight: bold; }}
        .fail {{ color: #dc3545; font-weight: bold; }}
        .results-table {{ max-width: 600px; }}
        .validation-table {{ max-width: 500px; }}
        .custom-table {{ margin-top: 0.5rem; }}
        .scope-table {{ margin: 0.5rem 0 1rem 0; }}
        section h3 {{ color: #2F5496; margin-top: 1.2rem; margin-bottom: 0.4rem;
            font-size: 1.1rem; border-bottom: 1px solid #ddd; padding-bottom: 0.2rem; }}
        section ul {{ margin: 0.3rem 0 0.8rem 1.5rem; }}
        section ul li {{ margin-bottom: 0.3rem; }}
        .reference-list {{ list-style: none; padding-left: 0; }}
        .reference-list li {{ margin-bottom: 0.6rem; padding: 0.4rem 0.6rem;
            border-left: 3px solid #2F5496; background: #f8f9fa; }}
        .reference-list li strong {{ color: #2F5496; }}
        /* Risk assessment styles */
        .risk-summary-card {{ background: #f8f9fa; border-left: 4px solid #2F5496;
            padding: 0.8rem 1.2rem; margin-bottom: 1.2rem; border-radius: 0 4px 4px 0; }}
        .risk-badge {{ display: inline-block; padding: 0.15rem 0.6rem; border-radius: 3px;
            font-weight: 600; font-size: 0.85rem; color: #fff; }}
        .risk-badge.risk-high {{ background: #dc3545; }}
        .risk-badge.risk-medium {{ background: #e67e22; }}
        .risk-badge.risk-low {{ background: #28a745; }}
        .risk-table {{ font-size: 0.88rem; }}
        .risk-table .mitigation-cell {{ font-size: 0.85rem; color: #555; }}
        .risk-level {{ font-weight: bold; text-align: center; }}
        .risk-level.risk-high {{ color: #dc3545; }}
        .risk-level.risk-medium {{ color: #e67e22; }}
        .risk-level.risk-low {{ color: #28a745; }}
        /* Uncertainty analysis styles */
        .uncertainty-summary {{ background: #e8f4fd; border-left: 4px solid #0366d6;
            padding: 0.8rem 1.2rem; margin-bottom: 1.2rem; border-radius: 0 4px 4px 0; }}
        .uncertainty-table {{ font-size: 0.88rem; }}
        .tornado-table {{ font-size: 0.88rem; }}
        /* Benchmark validation styles */
        .benchmark-table {{ font-size: 0.88rem; }}
        /* Discussion section styles */
        .discussion-block {{ background: #f9fafb; border-left: 4px solid #2d6a4f;
            padding: 1rem 1.2rem; margin-bottom: 1.5rem; border-radius: 0 4px 4px 0; }}
        .discussion-block h3 {{ color: #2d6a4f; margin-top: 0; }}
        .discussion-block .recommendation {{ background: #e8f5e9; padding: 0.5rem 0.8rem;
            border-radius: 3px; border-left: 3px solid #28a745; }}
        .discussion-block .traceability {{ font-size: 0.8rem; color: #888;
            border-top: 1px dashed #ccc; padding-top: 0.4rem; margin-top: 0.6rem; }}
        @media (max-width: 768px) {{
            nav {{ position: static; width: 100%; min-height: auto; }}
            main {{ margin-left: 0; padding: 1rem; }}
        }}
    </style>
</head>
<body>
    <nav>
        <h3>Contents</h3>
        <ul>
{nav}
        </ul>
        <hr style="margin: 1rem 0;">
        <p style="font-size: 0.8rem; color: #999;">{doc_num}</p>
        <p style="font-size: 0.8rem; color: #999;">Rev {rev} | {date}</p>
    </nav>
    <main>
        <div class="cover-page">
            <h1>{title}</h1>
            <p class="subtitle">NeqSim Engineering Report</p>
            <table class="cover-meta">
                <tr><td>Document No.</td><td>{doc_num}</td></tr>
                <tr><td>Revision</td><td>{rev}</td></tr>
                <tr><td>Date</td><td>{date}</td></tr>
                <tr><td>Author</td><td>{author}</td></tr>
                <tr><td>Classification</td><td>{classification}</td></tr>
            </table>
            <h3 style="margin-top: 2rem; color: #2F5496;">Revision History</h3>
            <table class="rev-table">
                <thead><tr><th>Rev</th><th>Date</th><th>Description</th><th>Author</th></tr></thead>
                <tbody>{rev_rows}</tbody>
            </table>
        </div>
{sections}
    </main>{katex_body_script}
</body>
</html>""".format(
        title=TITLE,
        author=AUTHOR or "(not specified)",
        date=TASK_DATE,
        doc_num=_auto_doc_number(),
        rev=REVISION,
        classification=CLASSIFICATION,
        rev_rows=_build_rev_rows_html(),
        nav=nav_items,
        sections=section_html,
        katex_head=katex_head,
        katex_body_script=katex_body_script,
    )

    with open(HTML_FILE, "w", encoding="utf-8") as f:
        f.write(html)
    print("HTML report saved: {}".format(HTML_FILE))


# ══════════════════════════════════════════════════════════
# Scientific Paper Generation
# ══════════════════════════════════════════════════════════

def build_paper_sections(results, task_spec):
    """Build scientific paper sections from results.json and task_spec.md.

    Maps task data to a standard engineering paper structure:
    Abstract, Introduction, Methodology, Results & Discussion,
    Uncertainty Analysis, Conclusions, Acknowledgments, References.
    """
    sections = []
    paper_title = PAPER_TITLE or TITLE

    # Abstract
    abstract = PAPER_SECTIONS["abstract"]
    if results and results.get("approach") and results.get("conclusions"):
        # Auto-generate abstract from approach + key results + conclusions
        parts = []
        parts.append(results["approach"])
        kr = results.get("key_results", {})
        if kr:
            highlights = []
            for key, value in list(kr.items())[:5]:
                label, unit = _parse_key_name(key)
                if isinstance(value, float):
                    highlights.append("{}: {:.4g} {}".format(label, value, unit).strip())
                else:
                    highlights.append("{}: {} {}".format(label, value, unit).strip())
            parts.append("Key results: " + "; ".join(highlights) + ".")
        parts.append(results["conclusions"])
        abstract = " ".join(parts)
    sections.append({
        "type": "abstract",
        "heading": "Abstract",
        "content": abstract,
    })

    # 1. Introduction
    intro = PAPER_SECTIONS["introduction"]
    # Try to auto-populate from problem_description + task_spec background
    if MANUAL_SECTIONS.get("problem_description") and not MANUAL_SECTIONS[
            "problem_description"].startswith("["):
        intro = MANUAL_SECTIONS["problem_description"]
        background = extract_spec_section(task_spec, "Background")
        if background:
            intro = background + "\n\n" + intro
    sections.append({
        "type": "numbered",
        "number": 1,
        "heading": "Introduction",
        "content": intro,
    })

    # 2. Methodology
    methodology = PAPER_SECTIONS["methodology"]
    if results and results.get("approach"):
        methodology = results["approach"]
    # Append standards info from task_spec
    standards = extract_spec_section(task_spec, "Applicable Standards")
    methods = extract_spec_section(task_spec, "Calculation Methods")
    if standards:
        methodology += "\n\nApplicable Standards:\n" + standards
    if methods:
        methodology += "\n\nCalculation Methods:\n" + methods
    sections.append({
        "type": "numbered",
        "number": 2,
        "heading": "Methodology",
        "content": methodology,
        "has_equations": True,
    })

    # 3. Results and Discussion
    results_text = PAPER_SECTIONS["results_discussion"]
    if results and results.get("key_results"):
        results_text = format_results_table(results)
    sections.append({
        "type": "numbered",
        "number": 3,
        "heading": "Results and Discussion",
        "content": results_text,
        "has_figures": True,
        "has_discussion": bool(results and results.get("figure_discussion")),
    })

    # 3.1 Validation (sub-section if data available)
    if results and results.get("validation"):
        sections.append({
            "type": "numbered",
            "number": 3.1,
            "heading": "Validation",
            "content": format_validation_table(results),
            "is_subsection": True,
        })

    # 3.2 Benchmark comparison (sub-section if data available)
    if results and results.get("benchmark_validation"):
        bv = results["benchmark_validation"]
        bv_lines = []
        for key, val in bv.items():
            desc = val.get("description", key)
            status = val.get("status", "N/A")
            bv_lines.append("- {}: {}".format(desc, status))
            if "max_deviation_pct" in val:
                bv_lines.append("  Max deviation: {:.4f}%".format(
                    val["max_deviation_pct"]))
            if "deviation_pct" in val:
                bv_lines.append("  Deviation: {:.2f}%".format(
                    val["deviation_pct"]))
        sections.append({
            "type": "numbered",
            "number": 3.2,
            "heading": "Benchmark Comparison",
            "content": "\n".join(bv_lines),
            "is_subsection": True,
        })

    # 4. Uncertainty Analysis (if data available)
    if results and results.get("uncertainty"):
        unc = results["uncertainty"]
        unc_lines = []
        unc_lines.append("A {} was performed with {} simulations using {}.".format(
            unc.get("method", "Monte Carlo analysis"),
            unc.get("n_simulations", "N/A"),
            unc.get("simulation_engine", "NeqSim")))
        unc_lines.append("")
        unc_lines.append("Input parameters and ranges:")
        for param in unc.get("input_parameters", []):
            unc_lines.append("  - {} [{}]: {}-{} ({}, base={})".format(
                param["name"], param.get("unit", ""),
                param["low"], param["high"],
                param.get("distribution", "uniform"), param["base"]))
        unc_lines.append("")
        # Output P10/P50/P90
        out_param = unc.get("output_parameter", "")
        if out_param:
            unc_lines.append("Results for {}:".format(out_param))
            unc_lines.append("  P10: {}, P50: {}, P90: {}".format(
                unc.get("p10", "N/A"), unc.get("p50", "N/A"),
                unc.get("p90", "N/A")))
        # Also handle output_parameters dict (mercury-style)
        for out_key, out_val in unc.get("output_parameters", {}).items():
            unc_lines.append("Results for {}:".format(out_key))
            unc_lines.append("  P10: {}, P50: {}, P90: {}".format(
                out_val.get("p10", "N/A"), out_val.get("p50", "N/A"),
                out_val.get("p90", "N/A")))
        sections.append({
            "type": "numbered",
            "number": 4,
            "heading": "Uncertainty Analysis",
            "content": "\n".join(unc_lines),
            "has_uncertainty": True,
        })

    # 5. Risk Assessment (if data available)
    if results and results.get("risk_evaluation"):
        re_data = results["risk_evaluation"]
        re_lines = []
        re_lines.append("Risk assessment using {} framework.".format(
            re_data.get("risk_matrix_used", "5x5 (ISO 31000)")))
        re_lines.append("Overall risk level: {}.".format(
            re_data.get("overall_risk_level", "N/A")))
        re_lines.append("")
        for risk in re_data.get("risks", []):
            re_lines.append("- {} ({}): {} [{}]".format(
                risk["id"], risk["category"], risk["description"],
                risk["risk_level"]))
            re_lines.append("  Mitigation: {}".format(risk["mitigation"]))
        next_num = 5
        sections.append({
            "type": "numbered",
            "number": next_num,
            "heading": "Risk Assessment",
            "content": "\n".join(re_lines),
            "has_risk": True,
        })

    # N. Conclusions
    conclusions = PAPER_SECTIONS["conclusions"]
    if results and results.get("conclusions"):
        conclusions = results["conclusions"]
    elif not MANUAL_SECTIONS["conclusions"].startswith("["):
        conclusions = MANUAL_SECTIONS["conclusions"]
    # Determine next section number
    last_num = max((s.get("number", 0) for s in sections
                    if isinstance(s.get("number"), int)), default=3)
    sections.append({
        "type": "numbered",
        "number": last_num + 1,
        "heading": "Conclusions",
        "content": conclusions,
    })

    # Acknowledgments (unnumbered)
    ack = PAPER_SECTIONS["acknowledgments"] or PAPER_ACKNOWLEDGMENTS
    if ack:
        sections.append({
            "type": "unnumbered",
            "heading": "Acknowledgments",
            "content": ack,
        })

    # References (unnumbered)
    refs_content = ""
    if results and results.get("references"):
        ref_lines = []
        for i, ref in enumerate(results["references"], 1):
            ref_text = ref.get("text", "")
            ref_lines.append("[{}] {}".format(i, ref_text))
        refs_content = "\n".join(ref_lines)
    elif not MANUAL_SECTIONS["references"].startswith("["):
        refs_content = MANUAL_SECTIONS["references"]
    if refs_content:
        sections.append({
            "type": "references",
            "heading": "References",
            "content": refs_content,
            "has_references": True,
        })

    return sections


def _paper_section_heading(section):
    """Format a section heading with number for the paper."""
    stype = section.get("type", "numbered")
    if stype == "abstract":
        return "Abstract"
    elif stype in ("unnumbered", "references"):
        return section["heading"]
    elif section.get("is_subsection"):
        return "{} {}".format(section["number"], section["heading"])
    else:
        return "{}. {}".format(section["number"], section["heading"])


def build_paper_docx(sections, results=None):
    """Build a scientific paper in Word format.

    Uses standard academic formatting: Times New Roman, single-column,
    numbered sections, centered title/author block, italic abstract,
    numbered figures and equations.
    """
    doc = Document()

    # ── Page setup ──
    for doc_section in doc.sections:
        doc_section.top_margin = Inches(1.0)
        doc_section.bottom_margin = Inches(1.0)
        doc_section.left_margin = Inches(1.0)
        doc_section.right_margin = Inches(1.0)

    paper_title = PAPER_TITLE or TITLE

    # ── Title (centered, bold, 16pt) ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(paper_title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    doc.add_paragraph("")

    # ── Authors and affiliations (centered) ──
    authors = PAPER_AUTHORS
    if authors:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        names = []
        affiliations = []
        seen_aff = {}
        for i, a in enumerate(authors):
            name = a.get("name", "")
            aff = a.get("affiliation", "")
            if aff and aff not in seen_aff:
                seen_aff[aff] = len(seen_aff) + 1
                affiliations.append(aff)
            sup = str(seen_aff.get(aff, "")) if aff else ""
            names.append(name + ("" if not sup else sup))
        author_run = author_para.add_run(", ".join(names))
        author_run.font.size = Pt(12)
        author_run.font.name = "Times New Roman"
        if affiliations:
            aff_para = doc.add_paragraph()
            aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for idx, aff in enumerate(affiliations, 1):
                aff_run = aff_para.add_run("{}{}".format(idx, aff))
                aff_run.font.size = Pt(10)
                aff_run.font.name = "Times New Roman"
                aff_run.font.italic = True
                if idx < len(affiliations):
                    aff_para.add_run("; ")
    elif AUTHOR:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(AUTHOR)
        author_run.font.size = Pt(12)
        author_run.font.name = "Times New Roman"

    # ── Date (centered) ──
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(TASK_DATE)
    date_run.font.size = Pt(10)
    date_run.font.name = "Times New Roman"
    date_run.font.italic = True

    doc.add_paragraph("")  # spacing

    # ── Keywords (if provided) ──
    if PAPER_KEYWORDS:
        kw_para = doc.add_paragraph()
        kw_bold = kw_para.add_run("Keywords: ")
        kw_bold.bold = True
        kw_bold.font.size = Pt(10)
        kw_bold.font.name = "Times New Roman"
        kw_text = kw_para.add_run(", ".join(PAPER_KEYWORDS))
        kw_text.font.size = Pt(10)
        kw_text.font.name = "Times New Roman"
        kw_text.font.italic = True
        doc.add_paragraph("")

    # Track figure and equation counters for the whole paper
    fig_counter = [0]
    eq_counter = [0]

    # ── Sections ──
    for section in sections:
        heading_text = _paper_section_heading(section)
        is_sub = section.get("is_subsection", False)
        stype = section.get("type", "numbered")

        # Heading level
        if stype == "abstract":
            h = doc.add_heading(heading_text, level=1)
        elif is_sub:
            h = doc.add_heading(heading_text, level=2)
        else:
            h = doc.add_heading(heading_text, level=1)

        # Style heading runs as Times New Roman
        for run in h.runs:
            run.font.name = "Times New Roman"

        # Abstract is italic
        if stype == "abstract":
            for para_text in section["content"].split("\n\n"):
                if para_text.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(para_text.strip())
                    r.font.italic = True
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"
        elif stype == "references" and results and results.get("references"):
            # Numbered reference list
            for i, ref in enumerate(results["references"], 1):
                ref_text = ref.get("text", "")
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.3)
                bracket_run = p.add_run("[{}] ".format(i))
                bracket_run.bold = True
                bracket_run.font.size = Pt(10)
                bracket_run.font.name = "Times New Roman"
                text_run = p.add_run(ref_text)
                text_run.font.size = Pt(10)
                text_run.font.name = "Times New Roman"
        elif section.get("has_figures") and results and results.get("key_results"):
            # Results section: add results table
            add_results_word_table(doc, results)
            if results.get("tables"):
                add_custom_word_tables(doc, results)
            # Discussion text
            disc = section["content"]
            if disc and not disc.startswith("["):
                for para_text in disc.split("\n\n"):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
        elif "Validation" in section["heading"] and results and results.get("validation"):
            add_validation_word_table(doc, results)
        elif section.get("has_benchmark") and results:
            add_benchmark_word_table(doc, results)
        elif section.get("has_uncertainty") and results:
            add_uncertainty_word_tables(doc, results)
        elif section.get("has_risk") and results:
            add_risk_word_table(doc, results)
        elif section.get("has_scope", False):
            render_scope_to_word(doc, section["content"])
        else:
            # Regular text content
            for para_text in section["content"].split("\n\n"):
                if para_text.strip():
                    p = doc.add_paragraph()
                    _add_bold_runs(p, para_text.strip())
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        if not run.font.size:
                            run.font.size = Pt(11)

        # Embed equations after Methodology section
        if section.get("has_equations"):
            equations = get_equations(results)
            if equations:
                eq_img_dir = os.path.join(REPORT_DIR, "_eq_images")
                if not os.path.exists(eq_img_dir):
                    os.makedirs(eq_img_dir)
                for eq in equations:
                    eq_counter[0] += 1
                    label = eq.get("label", "")
                    latex = eq.get("latex", "")
                    if not latex:
                        continue
                    eq_img_path = os.path.join(
                        eq_img_dir, "eq_{}.png".format(eq_counter[0]))
                    if render_equation_to_image(latex, eq_img_path):
                        doc.add_paragraph("")
                        doc.add_picture(eq_img_path, width=Inches(5.0))
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap = doc.add_paragraph(
                            "({}){}".format(
                                eq_counter[0],
                                "  " + label if label else ""))
                        cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for run in cap.runs:
                            run.font.size = Pt(10)
                            run.font.name = "Times New Roman"
                    else:
                        doc.add_paragraph("{}: {}".format(label, latex))

        # Embed figures after Results section
        if section.get("has_figures"):
            figures = get_figures()
            if figures:
                for fig_path in figures:
                    fig_counter[0] += 1
                    caption_text = get_figure_caption(
                        fig_path, results, fig_counter[0])
                    doc.add_paragraph("")
                    doc.add_picture(fig_path, width=Inches(5.5))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap = doc.add_paragraph(caption_text)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cap.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Times New Roman"
                        run.font.italic = True
                    doc.add_paragraph("")

        # Embed figure discussion after figures in Results & Discussion section
        if section.get("has_discussion") and results:
            add_discussion_word(doc, results)

    doc.save(PAPER_DOCX_FILE)
    print("Scientific paper (Word) saved: {}".format(PAPER_DOCX_FILE))


def build_paper_html(sections, results=None):
    """Build a scientific paper in HTML format with academic styling.

    Single-column layout, no sidebar, centered title/author block,
    properly numbered sections, KaTeX equations, and embedded figures.
    """
    paper_title = PAPER_TITLE or TITLE
    figures = get_figures()

    # Build author block
    author_html = ""
    authors = PAPER_AUTHORS
    if authors:
        names = []
        affiliations = []
        seen_aff = {}
        for a in authors:
            name = a.get("name", "")
            aff = a.get("affiliation", "")
            if aff and aff not in seen_aff:
                seen_aff[aff] = len(seen_aff) + 1
                affiliations.append(aff)
            sup = str(seen_aff.get(aff, "")) if aff else ""
            names.append("{}<sup>{}</sup>".format(name, sup) if sup else name)
        author_html = '<p class="authors">{}</p>\n'.format(", ".join(names))
        if affiliations:
            aff_items = []
            for idx, aff in enumerate(affiliations, 1):
                aff_items.append("<sup>{}</sup>{}".format(idx, aff))
            author_html += '<p class="affiliations">{}</p>\n'.format(
                "; ".join(aff_items))
    elif AUTHOR:
        author_html = '<p class="authors">{}</p>\n'.format(AUTHOR)

    # Keywords
    keywords_html = ""
    if PAPER_KEYWORDS:
        keywords_html = (
            '<p class="keywords"><strong>Keywords:</strong> '
            '<em>{}</em></p>\n'.format(", ".join(PAPER_KEYWORDS)))

    # Build figure HTML (reuse base64 embedding)
    fig_counter = [0]

    def _embed_figure(fig_path):
        fig_counter[0] += 1
        fig_name = os.path.basename(fig_path)
        caption = get_figure_caption(fig_path, results, fig_counter[0])
        mime = "image/svg+xml" if fig_path.endswith(".svg") else "image/png"
        with open(fig_path, "rb") as f:
            img_data = base64.b64encode(f.read()).decode("utf-8")
        return """
        <div class="paper-figure">
            <img src="data:{};base64,{}" alt="{}">
            <p class="fig-caption">{}</p>
        </div>""".format(mime, img_data, caption, caption)

    # Build equation HTML
    eq_counter = [0]

    def _embed_equation(eq):
        eq_counter[0] += 1
        label = eq.get("label", "")
        latex = eq.get("latex", "")
        if not latex:
            return ""
        fallback_img = ""
        eq_img_dir = os.path.join(REPORT_DIR, "_eq_images")
        if not os.path.exists(eq_img_dir):
            os.makedirs(eq_img_dir)
        eq_img_path = os.path.join(eq_img_dir, "eq_{}.png".format(eq_counter[0]))
        if render_equation_to_image(latex, eq_img_path):
            with open(eq_img_path, "rb") as imgf:
                img_b64 = base64.b64encode(imgf.read()).decode("utf-8")
            fallback_img = (
                '<img class="eq-fallback" src="data:image/png;base64,{}" '
                'alt="{}" style="display:none; max-width:90%;">'.format(
                    img_b64, label))
        return """
        <div class="paper-equation">
            <div class="eq-content katex-eq">$${}$$</div>
            {}
            <span class="eq-number">({}){}</span>
        </div>""".format(latex, fallback_img, eq_counter[0],
                         "  " + label if label else "")

    # Build sections HTML
    section_html = ""
    for section in sections:
        heading_text = _paper_section_heading(section)
        stype = section.get("type", "numbered")
        is_sub = section.get("is_subsection", False)
        section_id = heading_text.lower().replace(" ", "-").replace(".", "")

        # Heading tag
        if stype == "abstract":
            tag = "h2"
        elif is_sub:
            tag = "h3"
        else:
            tag = "h2"

        # Content formatting
        if stype == "abstract":
            content = '<div class="abstract-text">{}</div>'.format(
                section["content"].replace("\n\n", "</p><p>").replace("\n", "<br>"))
        elif stype == "references" and results and results.get("references"):
            content = format_references_html(results)
        elif section.get("has_figures") and results and results.get("key_results"):
            content = format_results_html(results)
            if results.get("tables"):
                content += format_custom_tables_html(results)
            # Add discussion text
            disc = section["content"]
            if disc and not disc.startswith("["):
                content += "<p>{}</p>".format(
                    disc.replace("\n\n", "</p><p>").replace("\n", "<br>"))
            # Add figures
            if figures:
                for fig_path in figures:
                    content += _embed_figure(fig_path)
        elif "Validation" in section["heading"] and results and results.get("validation"):
            content = format_validation_html(results)
        elif section.get("has_benchmark") and results:
            content = format_benchmark_html(results)
        elif section.get("has_uncertainty") and results:
            content = format_uncertainty_html(results)
        elif section.get("has_risk") and results:
            content = format_risk_html(results)
        elif section.get("has_scope", False):
            content = scope_content_to_html(section["content"])
        else:
            content = section["content"].replace("\n\n", "</p><p>").replace(
                "\n", "<br>")

        # Add figure discussion after figures in Results & Discussion
        if section.get("has_discussion") and results:
            content += format_discussion_html(results)

        # Add equations after methodology
        if section.get("has_equations"):
            equations = get_equations(results)
            if equations:
                for eq in equations:
                    content += _embed_equation(eq)

        section_html += """
        <section id="{}">
            <{} class="section-heading">{}</{}>
            <div class="section-content">{}</div>
        </section>""".format(section_id, tag, heading_text, tag, content)

    # KaTeX CDN
    equations = get_equations(results)
    katex_head = ""
    katex_body_script = ""
    if equations:
        katex_head = """
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.css">
    <script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.js"></script>
    <script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/contrib/auto-render.min.js"></script>"""
        katex_body_script = """
    <script>
        document.addEventListener("DOMContentLoaded", function() {{
            if (typeof renderMathInElement === "function") {{
                renderMathInElement(document.body, {{
                    delimiters: [
                        {{left: "$$", right: "$$", display: true}},
                        {{left: "$", right: "$", display: false}}
                    ],
                    throwOnError: false
                }});
            }} else {{
                var eqs = document.querySelectorAll(".katex-eq");
                for (var i = 0; i < eqs.length; i++) {{ eqs[i].style.display = "none"; }}
                var imgs = document.querySelectorAll(".eq-fallback");
                for (var j = 0; j < imgs.length; j++) {{ imgs[j].style.display = "inline"; }}
            }}
        }});
    </script>"""

    html = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>{katex_head}
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Times New Roman', Times, 'DejaVu Serif', Georgia, serif;
            line-height: 1.6; color: #1a1a1a;
            max-width: 800px; margin: 0 auto; padding: 2rem 2.5rem;
            background: #fff;
        }}
        /* ── Title block ── */
        .paper-header {{ text-align: center; margin-bottom: 2rem;
            border-bottom: 2px solid #333; padding-bottom: 1.5rem; }}
        .paper-header h1 {{ font-size: 1.6rem; font-weight: bold;
            margin-bottom: 0.8rem; line-height: 1.3; }}
        .authors {{ font-size: 1.05rem; margin-bottom: 0.3rem; }}
        .affiliations {{ font-size: 0.9rem; font-style: italic;
            color: #555; margin-bottom: 0.3rem; }}
        .paper-date {{ font-size: 0.9rem; color: #666;
            font-style: italic; margin-top: 0.5rem; }}
        .keywords {{ font-size: 0.9rem; margin-top: 0.8rem;
            text-align: left; padding: 0.5rem 1rem;
            background: #f9f9f9; border-left: 3px solid #333; }}
        /* ── Abstract ── */
        .abstract-text {{ font-style: italic; padding: 0.5rem 1.5rem;
            border-left: 3px solid #666; margin: 0.5rem 0 1.5rem 0;
            color: #333; font-size: 0.95rem; }}
        .abstract-text p {{ margin-bottom: 0.5rem; }}
        /* ── Section headings ── */
        h2.section-heading {{ font-size: 1.2rem; font-weight: bold;
            margin-top: 1.8rem; margin-bottom: 0.6rem; color: #1a1a1a;
            border-bottom: 1px solid #ccc; padding-bottom: 0.2rem; }}
        h3.section-heading {{ font-size: 1.05rem; font-weight: bold;
            margin-top: 1.2rem; margin-bottom: 0.4rem; color: #333; }}
        .section-content {{ margin-bottom: 1rem; }}
        .section-content p {{ margin-bottom: 0.6rem; text-align: justify; }}
        /* ── Figures ── */
        .paper-figure {{ text-align: center; margin: 1.5rem 0;
            page-break-inside: avoid; }}
        .paper-figure img {{ max-width: 100%; border: 1px solid #ddd; }}
        .fig-caption {{ font-size: 0.85rem; color: #333;
            margin-top: 0.4rem; font-style: italic; }}
        /* ── Equations ── */
        .paper-equation {{ display: flex; align-items: center;
            justify-content: center; margin: 1rem 0;
            position: relative; }}
        .eq-content {{ flex: 1; text-align: center; font-size: 1.1rem; }}
        .eq-number {{ position: absolute; right: 0; font-size: 0.95rem;
            color: #333; }}
        /* ── Tables ── */
        table {{ border-collapse: collapse; width: 100%; margin: 1rem 0;
            font-size: 0.9rem; }}
        thead th {{ background: #2F5496; color: #fff; font-weight: 600;
            padding: 0.5rem; text-align: left; border: 1px solid #2a4a85; }}
        tbody td {{ border: 1px solid #ddd; padding: 0.4rem 0.5rem; }}
        tbody tr:nth-child(even) {{ background: #f8f9fa; }}
        td.num {{ text-align: right; font-variant-numeric: tabular-nums; }}
        .pass {{ color: #28a745; font-weight: bold; }}
        .fail {{ color: #dc3545; font-weight: bold; }}
        .results-table {{ max-width: 600px; margin: 1rem auto; }}
        .validation-table {{ max-width: 500px; margin: 1rem auto; }}
        .custom-table {{ margin: 0.5rem auto; }}
        /* ── References ── */
        .reference-list {{ list-style: none; padding-left: 0; }}
        .reference-list li {{ margin-bottom: 0.5rem; padding-left: 2rem;
            text-indent: -2rem; font-size: 0.9rem; }}
        .reference-list li strong {{ color: #333; }}
        /* ── Risk assessment styles ── */
        .risk-summary-card {{ background: #f8f9fa; border-left: 4px solid #333;
            padding: 0.8rem 1.2rem; margin-bottom: 1.2rem; }}
        .risk-badge {{ display: inline-block; padding: 0.15rem 0.6rem; border-radius: 3px;
            font-weight: 600; font-size: 0.85rem; color: #fff; }}
        .risk-badge.risk-high {{ background: #dc3545; }}
        .risk-badge.risk-medium {{ background: #e67e22; }}
        .risk-badge.risk-low {{ background: #28a745; }}
        .risk-table {{ font-size: 0.88rem; }}
        .risk-table .mitigation-cell {{ font-size: 0.85rem; color: #555; }}
        .risk-level {{ font-weight: bold; text-align: center; }}
        .risk-level.risk-high {{ color: #dc3545; }}
        .risk-level.risk-medium {{ color: #e67e22; }}
        .risk-level.risk-low {{ color: #28a745; }}
        /* ── Uncertainty analysis styles ── */
        .uncertainty-summary {{ background: #f5f5f5; border-left: 4px solid #333;
            padding: 0.8rem 1.2rem; margin-bottom: 1.2rem; }}
        .uncertainty-table {{ font-size: 0.88rem; }}
        .tornado-table {{ font-size: 0.88rem; }}
        /* ── Benchmark validation styles ── */
        .benchmark-table {{ font-size: 0.88rem; }}
        /* ── Discussion section styles ── */
        .discussion-block {{ background: #f9fafb; border-left: 4px solid #2d6a4f;
            padding: 1rem 1.2rem; margin-bottom: 1.5rem; border-radius: 0 4px 4px 0; }}
        .discussion-block h3 {{ color: #2d6a4f; margin-top: 0; }}
        .discussion-block .recommendation {{ background: #e8f5e9; padding: 0.5rem 0.8rem;
            border-radius: 3px; border-left: 3px solid #28a745; }}
        .discussion-block .traceability {{ font-size: 0.8rem; color: #888;
            border-top: 1px dashed #ccc; padding-top: 0.4rem; margin-top: 0.6rem; }}
        /* ── Footer ── */
        .paper-footer {{ margin-top: 3rem; padding-top: 1rem;
            border-top: 1px solid #ccc; font-size: 0.8rem;
            color: #999; text-align: center; }}
        /* ── Print styles ── */
        @media print {{
            body {{ padding: 0; max-width: none; }}
            .paper-figure {{ page-break-inside: avoid; }}
            table {{ page-break-inside: avoid; }}
        }}
    </style>
</head>
<body>
    <div class="paper-header">
        <h1>{title}</h1>
        {author_block}
        <p class="paper-date">{date}</p>
        {keywords}
    </div>
{sections}
    <div class="paper-footer">
        <p>Generated {date} using NeqSim task-solving workflow</p>
    </div>{katex_body_script}
</body>
</html>""".format(
        title=paper_title,
        author_block=author_html,
        date=TASK_DATE,
        keywords=keywords_html,
        sections=section_html,
        katex_head=katex_head,
        katex_body_script=katex_body_script,
    )

    with open(PAPER_HTML_FILE, "w", encoding="utf-8") as f:
        f.write(html)
    print("Scientific paper (HTML) saved: {}".format(PAPER_HTML_FILE))


# ══════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════

if __name__ == "__main__":
    generate_paper = "--paper" in sys.argv or "--paper-only" in sys.argv
    paper_only = "--paper-only" in sys.argv

    print("Generating outputs for: {}".format(TITLE))
    print("")

    # Auto-read task data
    study_config = load_study_config()
    results = load_results()
    task_spec = load_task_spec()
    study_config_warnings = validate_study_config(study_config, results, task_spec)
    if study_config_warnings:
        print("")
        print("Study configuration warnings:")
        for warning in study_config_warnings:
            print("  - {}".format(warning))

    if not paper_only:
        # Build report sections and generate technical report
        sections = build_sections(results, task_spec, study_config_warnings)
        print("")
        build_word_report(sections, results)
        build_html_report(sections, results)
        print("")
        print("Technical reports generated.")
        print("  Open Report.html in a browser for navigable view.")
        print("  Open Report.docx for formal distribution.")

    if generate_paper:
        # Build paper sections and generate scientific paper
        paper_sections = build_paper_sections(results, task_spec)
        print("")
        build_paper_docx(paper_sections, results)
        build_paper_html(paper_sections, results)
        print("")
        print("Scientific papers generated.")
        print("  Open Paper.html in a browser for reading.")
        print("  Open Paper.docx for journal submission / distribution.")

    if not generate_paper and not paper_only:
        print("")
        print("TIP: Add --paper flag to also generate a scientific paper.")
        print("     python step3_report/generate_report.py --paper")

    if not results:
        print("")
        print("TIP: Create results.json in the task root to auto-populate")
        print("     the Results and Validation sections. See the task README")
        print("     for the results.json pattern.")
