#!/usr/bin/env python3
"""
Build publication-quality Word document from paper.md.

Generates a formatted .docx following Computers & Chemical Engineering style:
- Double-spaced, line-numbered body text
- Proper heading hierarchy (section/subsection numbering)
- Tables with three-line style (booktabs equivalent)
- Embedded figures with captions
- Equation rendering as formatted text
- Citation handling (cite key -> [N])
- Highlights box, abstract, keywords
"""

import re
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──
PAPER_DIR = Path(__file__).parent
FIGURES_DIR = PAPER_DIR / "figures"
SUBMISSION_DIR = PAPER_DIR / "submission"
SUBMISSION_DIR.mkdir(exist_ok=True)

# ── BibTeX citation key -> number mapping ──
# Build from refs.bib order
REFS_FILE = PAPER_DIR / "refs.bib"


def build_citation_map():
    """Parse refs.bib and assign sequential numbers."""
    if not REFS_FILE.exists():
        return {}
    text = REFS_FILE.read_text(encoding="utf-8")
    keys = re.findall(r'@\w+\{(\w+),', text)
    # Sort alphabetically (matching elsarticle-num behavior for author-year sorted bib)
    keys_sorted = sorted(set(keys), key=str.lower)
    return {k: i + 1 for i, k in enumerate(keys_sorted)}


CITE_MAP = build_citation_map()


def resolve_cite(match):
    """Replace \\cite{key1, key2} with [N1, N2]."""
    raw = match.group(1)
    keys = [k.strip() for k in raw.split(",")]
    nums = []
    for k in keys:
        if k in CITE_MAP:
            nums.append(str(CITE_MAP[k]))
        else:
            nums.append(f"?{k}")
    return "[" + ", ".join(sorted(nums, key=lambda x: int(x) if x.isdigit() else 999)) + "]"


def clean_latex_math_inline(text):
    """Convert simple LaTeX math to readable Unicode/text for Word."""
    # Common substitutions for inline math
    # Greek letters
    greek = {
        r'\alpha': '\u03b1', r'\beta': '\u03b2', r'\gamma': '\u03b3',
        r'\delta': '\u03b4', r'\epsilon': '\u03b5', r'\zeta': '\u03b6',
        r'\eta': '\u03b7', r'\theta': '\u03b8', r'\iota': '\u03b9',
        r'\kappa': '\u03ba', r'\lambda': '\u03bb', r'\mu': '\u03bc',
        r'\nu': '\u03bd', r'\xi': '\u03be', r'\pi': '\u03c0',
        r'\rho': '\u03c1', r'\sigma': '\u03c3', r'\tau': '\u03c4',
        r'\phi': '\u03c6', r'\chi': '\u03c7', r'\psi': '\u03c8',
        r'\omega': '\u03c9', r'\Delta': '\u0394', r'\Sigma': '\u03a3',
        r'\Omega': '\u03a9', r'\Lambda': '\u039b',
    }
    for cmd, char in greek.items():
        text = text.replace(cmd, char)

    # Math operators
    text = text.replace(r'\times', '\u00d7')
    text = text.replace(r'\cdot', '\u00b7')
    text = text.replace(r'\leq', '\u2264')
    text = text.replace(r'\geq', '\u2265')
    text = text.replace(r'\neq', '\u2260')
    text = text.replace(r'\approx', '\u2248')
    text = text.replace(r'\sim', '\u223c')
    text = text.replace(r'\infty', '\u221e')
    text = text.replace(r'\partial', '\u2202')
    text = text.replace(r'\nabla', '\u2207')
    text = text.replace(r'\pm', '\u00b1')
    text = text.replace(r'\to', '\u2192')
    text = text.replace(r'\rightarrow', '\u2192')
    text = text.replace(r'\leftarrow', '\u2190')
    text = text.replace(r'\sum', '\u2211')
    text = text.replace(r'\prod', '\u220f')
    text = text.replace(r'\in', '\u2208')

    # Clean up \text{} and \mathrm{}
    text = re.sub(r'\\text\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\hat\{([^}]*)\}', lambda m: m.group(1) + '\u0302', text)
    text = re.sub(r'\\tilde\{([^}]*)\}', lambda m: m.group(1) + '\u0303', text)
    text = re.sub(r'\\bar\{([^}]*)\}', lambda m: m.group(1) + '\u0304', text)

    # Fractions: \frac{a}{b} -> (a)/(b)
    text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', text)

    # Superscript/subscript markers (simplified)
    text = re.sub(r'\^{([^}]*)}', r'^\1', text)  # keep ^ for superscript
    text = re.sub(r'_{([^}]*)}', r'_\1', text)    # keep _ for subscript

    # Remove remaining \commands
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    # Clean braces
    text = text.replace('{', '').replace('}', '')

    return text.strip()


def process_inline_text(paragraph, text):
    """Add text to a paragraph, handling bold, italic, code, math, and citations."""
    # Resolve citations first
    text = re.sub(r'\\cite\{([^}]+)\}', resolve_cite, text)

    # Remove HTML comments
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)

    # Split on formatting markers
    # Pattern handles: **bold**, *italic*, `code`, $math$
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`|\$[^$]+\$)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
        elif part.startswith('$') and part.endswith('$'):
            # Inline math
            math_text = clean_latex_math_inline(part[1:-1])
            run = paragraph.add_run(math_text)
            run.italic = True
            run.font.name = 'Cambria Math'
            run.font.size = Pt(10.5)
        else:
            paragraph.add_run(part)


def set_cell_shading(cell, color_hex):
    """Set background/shading color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell (width in 1/8 pt, e.g. 12 = 1.5pt)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if spec:
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{spec}" w:space="0" w:color="000000"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_table_from_markdown(doc, md_lines):
    """Parse markdown table lines and add a formatted Word table."""
    header_line = md_lines[0]
    # Skip separator line (index 1)
    data_lines = md_lines[2:] if len(md_lines) > 2 else []

    # Parse cells
    def parse_row(line):
        cells = [c.strip() for c in line.split('|')]
        return [c for c in cells if c]  # remove empty from leading/trailing |

    headers = parse_row(header_line)
    rows = [parse_row(line) for line in data_lines]

    n_cols = len(headers)
    n_rows = len(rows)

    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Set column widths
    # First column wider for labels
    total_width = Inches(6.5)

    # Header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clean_latex_math_inline(h))
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        set_cell_shading(cell, 'E3F2FD')
        set_cell_borders(cell, top=12, bottom=8)

    # Data rows
    for i, row_cells in enumerate(rows):
        for j, cell_text in enumerate(row_cells):
            if j >= n_cols:
                break
            cell = table.cell(i + 1, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            cleaned = clean_latex_math_inline(cell_text)
            run = p.add_run(cleaned)
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            # Bottom border on last row
            if i == n_rows - 1:
                set_cell_borders(cell, bottom=12)

    # Remove internal vertical borders for clean booktabs look
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    return table


def add_equation_block(doc, eq_text):
    """Add a display equation as a centered italic paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    cleaned = clean_latex_math_inline(eq_text)
    run = p.add_run(cleaned)
    run.italic = True
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    return p


def add_figure(doc, img_path, caption_text, fig_num):
    """Add a figure with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)

    if Path(img_path).exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
    else:
        run = p.add_run(f"[Missing figure: {img_path}]")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # Caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(12)
    run = cap_p.add_run(f"Figure {fig_num}. ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    # Process rest of caption with inline formatting
    process_inline_text(cap_p, caption_text)
    for r in cap_p.runs[1:]:
        r.font.size = Pt(9)
        r.font.name = 'Arial'

    return p


def set_paragraph_format(paragraph, font_name='Times New Roman', font_size=11,
                         line_spacing=WD_LINE_SPACING.DOUBLE, space_after=Pt(0),
                         space_before=Pt(0)):
    """Apply formatting to a paragraph and its runs."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = line_spacing
    pf.space_after = space_after
    pf.space_before = space_before
    for run in paragraph.runs:
        run.font.name = font_name
        if run.font.size is None:
            run.font.size = Pt(font_size)


def build_word_document():
    """Build the complete Word document from paper.md."""
    paper_md = (PAPER_DIR / "paper.md").read_text(encoding="utf-8")

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)

    # ── Custom styles ──
    # Title style
    if 'Paper Title' not in [s.name for s in doc.styles]:
        title_style = doc.styles.add_style('Paper Title', 1)  # paragraph
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        title_style.paragraph_format.space_before = Pt(24)

    # ── Parse paper line by line ──
    lines = paper_md.split('\n')

    # Track state
    fig_counter = 0
    table_counter = 0
    eq_counter = 0
    in_equation = False
    eq_buffer = []
    in_table = False
    table_buffer = []
    table_caption = None
    in_list = False
    list_type = None  # 'bullet' or 'number'
    skip_line = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip HTML comments
        if '<!--' in line:
            if '-->' in line:
                i += 1
                continue
            # Multi-line comment
            while i < len(lines) and '-->' not in lines[i]:
                i += 1
            i += 1
            continue

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # ── Display equations ($$...$$) ──
        if line.strip().startswith('$$') and not in_equation:
            if line.strip().endswith('$$') and len(line.strip()) > 4:
                # Single-line equation
                eq_counter += 1
                eq_text = line.strip()[2:-2].strip()
                p = add_equation_block(doc, eq_text)
                # Add equation number
                run = p.add_run(f'    ({eq_counter})')
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                i += 1
                continue
            else:
                in_equation = True
                eq_buffer = []
                i += 1
                continue

        if in_equation:
            if line.strip() == '$$':
                eq_counter += 1
                eq_text = '\n'.join(eq_buffer)
                p = add_equation_block(doc, eq_text)
                run = p.add_run(f'    ({eq_counter})')
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                in_equation = False
                eq_buffer = []
                i += 1
                continue
            else:
                eq_buffer.append(line)
                i += 1
                continue

        # ── Tables (| ... | ... |) ──
        if re.match(r'^\s*\|', line) and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_buffer = []
            table_buffer.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            table_counter += 1
            add_table_from_markdown(doc, table_buffer)
            doc.add_paragraph()  # spacer
            in_table = False
            table_buffer = []
            # Don't increment i - process current line in next iteration
            continue

        # ── Images: ![caption](path) ──
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_rel = img_match.group(2)
            img_path = PAPER_DIR / img_rel
            fig_counter += 1

            # Check if next line is an italic caption
            caption = alt_text
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # Match *Figure N: caption text*
                cap_match = re.match(r'^\*Figure\s+\d+:\s*(.+)\*$', next_line)
                if cap_match:
                    caption = cap_match.group(1)
                    i += 1  # skip the caption line
                elif next_line.startswith('*') and next_line.endswith('*'):
                    # Generic italic line following figure
                    inner = next_line[1:-1]
                    fig_cap_match = re.match(r'Figure\s+\d+[.:]\s*(.*)', inner)
                    if fig_cap_match:
                        caption = fig_cap_match.group(1)
                    else:
                        caption = inner
                    i += 1  # skip

            add_figure(doc, str(img_path), caption, fig_counter)
            i += 1
            continue

        # ── Headings ──
        h1_match = re.match(r'^#\s+(.+)$', line)
        h2_match = re.match(r'^##\s+(.+)$', line)
        h3_match = re.match(r'^###\s+(.+)$', line)

        if h1_match:
            title_text = h1_match.group(1).strip()
            # Paper title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

            # Author placeholder
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('[Author Name]')
            run.font.size = Pt(12)
            run.font.name = 'Arial'

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(24)
            run = p.add_run('[Department, Institution, City, Country]')
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.italic = True

            i += 1
            continue

        if h2_match:
            heading_text = h2_match.group(1).strip()

            if heading_text.lower() == 'highlights':
                # Highlights box
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                run = p.add_run('Highlights')
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

                # Read highlights items
                i += 1
                while i < len(lines):
                    hl = lines[i].strip()
                    if hl.startswith('- ') or hl.startswith('* '):
                        item_text = hl[2:].strip()
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.3)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        run = p.add_run('\u2022  ')
                        run.font.size = Pt(10)
                        process_inline_text(p, item_text)
                        for r in p.runs:
                            r.font.size = Pt(10)
                            r.font.name = 'Arial'
                    elif hl.startswith('##') or (hl and not hl.startswith('-') and not hl.startswith('*')):
                        break
                    i += 1
                doc.add_paragraph()  # spacer
                continue

            elif heading_text.lower() == 'abstract':
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                run = p.add_run('Abstract')
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Arial'

                # Gather abstract text
                i += 1
                abstract_lines = []
                while i < len(lines) and not lines[i].startswith('##'):
                    abstract_lines.append(lines[i])
                    i += 1

                abs_text = ' '.join(abstract_lines).strip()
                if abs_text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.right_indent = Inches(0.3)
                    process_inline_text(p, abs_text)
                    set_paragraph_format(p, font_size=10, line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE)

                doc.add_paragraph()
                continue

            elif heading_text.lower().startswith('keyword'):
                p = doc.add_paragraph()
                run = p.add_run('Keywords: ')
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Arial'

                i += 1
                kw_lines = []
                while i < len(lines) and not lines[i].startswith('#') and lines[i].strip() != '---':
                    if lines[i].strip():
                        kw_lines.append(lines[i].strip())
                    i += 1

                kw_text = ' '.join(kw_lines)
                run = p.add_run(kw_text)
                run.italic = True
                run.font.size = Pt(10)
                run.font.name = 'Arial'

                doc.add_paragraph()
                # Page break before main body
                doc.add_page_break()
                continue

            elif heading_text.lower() == 'references':
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(18)
                run = p.add_run('References')
                run.bold = True
                run.font.size = Pt(13)
                run.font.name = 'Arial'

                # Build reference list from refs.bib
                i += 1
                if REFS_FILE.exists():
                    add_reference_list(doc)
                # Skip remaining lines in References section
                while i < len(lines) and not lines[i].startswith('##'):
                    i += 1
                continue

            else:
                # Regular H2 section heading
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(13)
                run.font.name = 'Arial'
                i += 1
                continue

        if h3_match:
            heading_text = h3_match.group(1).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            i += 1
            continue

        # ── Bullet lists ──
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            item_text = bullet_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4 + indent * 0.2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = p.add_run('\u2022  ')
            process_inline_text(p, item_text)
            for r in p.runs:
                if r.font.name is None:
                    r.font.name = 'Times New Roman'
                if r.font.size is None:
                    r.font.size = Pt(11)
            i += 1
            continue

        # ── Numbered lists ──
        num_match = re.match(r'^(\d+)\.\s+(.+)', line)
        if num_match:
            num = num_match.group(1)
            item_text = num_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = p.add_run(f'{num}. ')
            run.bold = True
            process_inline_text(p, item_text)
            for r in p.runs:
                if r.font.name is None:
                    r.font.name = 'Times New Roman'
                if r.font.size is None:
                    r.font.size = Pt(11)
            i += 1
            continue

        # ── Table caption (bold line starting with "Table") ──
        table_cap_match = re.match(r'^\*\*Table\s+\d+', line)
        if table_cap_match:
            table_counter += 1
            cap_text = re.sub(r'^\*\*(.+?)\*\*:?\s*', r'\1: ', line).strip()
            cap_text = re.sub(r'\*\*', '', cap_text)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(cap_text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            i += 1
            continue

        # ── Italic line following figure (already handled in image section) ──
        if line.strip().startswith('*') and line.strip().endswith('*') and 'Figure' in line:
            # Skip - already consumed by figure handler above
            i += 1
            continue

        # ── Regular paragraph ──
        if line.strip():
            # Accumulate paragraph lines (until empty line or special line)
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_l = lines[i]
                # Stop at: headings, empty lines, lists, tables, images, equations, HRs
                if (next_l.strip() == '' or
                    next_l.startswith('#') or
                    re.match(r'^\s*[-*]\s', next_l) or
                    re.match(r'^\d+\.\s', next_l) or
                    re.match(r'^\s*\|', next_l) or
                    next_l.strip().startswith('$$') or
                    next_l.strip() == '---' or
                    next_l.strip().startswith('![') or
                    next_l.strip().startswith('**Table')):
                    break
                para_lines.append(next_l)
                i += 1

            full_text = ' '.join(l.strip() for l in para_lines)
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            process_inline_text(p, full_text)
            for r in p.runs:
                if r.font.name is None:
                    r.font.name = 'Times New Roman'
                if r.font.size is None:
                    r.font.size = Pt(11)
            continue

        i += 1

    # Flush any remaining table
    if in_table and table_buffer:
        add_table_from_markdown(doc, table_buffer)

    return doc


def add_reference_list(doc):
    """Parse refs.bib and add formatted reference list."""
    text = REFS_FILE.read_text(encoding="utf-8")

    # Parse each entry
    entries = []
    for match in re.finditer(r'@(\w+)\{(\w+),\s*(.*?)\n\}', text, flags=re.DOTALL):
        entry_type = match.group(1)
        key = match.group(2)
        body = match.group(3)

        # Parse fields
        fields = {}
        for fm in re.finditer(r'(\w+)\s*=\s*\{(.*?)\}', body, flags=re.DOTALL):
            fields[fm.group(1).lower()] = fm.group(2).strip()

        entries.append((key, entry_type, fields))

    # Sort alphabetically by key
    entries.sort(key=lambda x: x[0].lower())

    for idx, (key, etype, fields) in enumerate(entries):
        num = idx + 1
        author = fields.get('author', 'Unknown')
        title = fields.get('title', '')
        year = fields.get('year', '')
        journal = fields.get('journal', '')
        volume = fields.get('volume', '')
        pages = fields.get('pages', '')
        publisher = fields.get('publisher', '')
        booktitle = fields.get('booktitle', '')
        edition = fields.get('edition', '')
        url = fields.get('url', '')

        # Format reference string
        ref_parts = []
        # Shorten author: "Last, F. and Last, F." style already in bib
        ref_parts.append(f'[{num}]  {author}')
        if year:
            ref_parts[-1] += f' ({year}).'
        if title:
            ref_parts.append(f'  {title}.')
        if journal:
            ref_parts.append(f'  {journal}')
            if volume:
                ref_parts[-1] += f', {volume}'
            if pages:
                ref_parts[-1] += f', {pages}'
            ref_parts[-1] += '.'
        elif booktitle:
            ref_parts.append(f'  In: {booktitle}.')
        elif publisher:
            ref_parts.append(f'  {publisher}.')
        if edition:
            ref_parts[-1] = ref_parts[-1].rstrip('.') + f', {edition} ed.'

        ref_text = ''.join(ref_parts)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(ref_text)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'


def main():
    print("Building Word document...")
    print(f"  Paper dir: {PAPER_DIR}")
    print(f"  Figures:   {len(list(FIGURES_DIR.glob('*.png')))} PNG files")
    print(f"  Citations: {len(CITE_MAP)} references")

    doc = build_word_document()

    out_path = SUBMISSION_DIR / "paper.docx"
    # If file is locked (open in Word), save with timestamp suffix
    if out_path.exists():
        try:
            doc.save(str(out_path))
        except PermissionError:
            from datetime import datetime
            ts = datetime.now().strftime("%H%M%S")
            alt_path = SUBMISSION_DIR / f"paper_{ts}.docx"
            doc.save(str(alt_path))
            out_path = alt_path
            print(f"  NOTE: paper.docx was locked, saved as {alt_path.name}")
    else:
        doc.save(str(out_path))

    print(f"\n  Output: {out_path}")
    print(f"  Size:   {out_path.stat().st_size / 1024:.0f} KB")
    print("\nDone!")


if __name__ == '__main__':
    main()
