"""
Generate Word document from the Wei Yan response paper.
"""
import pathlib
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

SCRIPT_DIR = pathlib.Path(__file__).resolve().parent
OUTPUT = SCRIPT_DIR / "paper_response_to_weiyan.docx"


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

    # Title style
    ts = doc.styles.add_style("PaperTitle", WD_STYLE_TYPE.PARAGRAPH)
    ts.font.name = "Times New Roman"
    ts.font.size = Pt(14)
    ts.font.bold = True
    ts.paragraph_format.space_after = Pt(12)
    ts.paragraph_format.space_before = Pt(24)
    ts.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author style
    aus = doc.styles.add_style("Author", WD_STYLE_TYPE.PARAGRAPH)
    aus.font.name = "Times New Roman"
    aus.font.size = Pt(11)
    aus.font.italic = True
    aus.paragraph_format.space_after = Pt(4)
    aus.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section heading
    for i, sz in [(1, 13), (2, 12), (3, 11)]:
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Times New Roman"
        hs.font.size = Pt(sz)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(6)

    # Quote style
    qs = doc.styles.add_style("BlockQuote", WD_STYLE_TYPE.PARAGRAPH)
    qs.font.name = "Times New Roman"
    qs.font.size = Pt(10)
    qs.font.italic = True
    qs.paragraph_format.left_indent = Cm(1.5)
    qs.paragraph_format.right_indent = Cm(1.0)
    qs.paragraph_format.space_before = Pt(6)
    qs.paragraph_format.space_after = Pt(6)


def add_para(doc, text, style="Normal", bold=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    return p


def add_numbered_item(doc, number, text):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(f"{number}. ")
    run.bold = True
    p.add_run(text)
    return p


def add_quote(doc, text):
    return add_para(doc, text, style="BlockQuote")


def build_document():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    setup_styles(doc)

    # ── Title ──
    add_para(
        doc,
        "On the Relationship Between Michelsen's (2006) Solution Procedures "
        "and the Fully Implicit CPA Algorithm of Igben et al. (2026): "
        "A Technical Clarification",
        style="PaperTitle",
    )

    add_para(doc, "Even Solbraa", style="Author")
    add_para(
        doc,
        "Department of Energy and Process Engineering, "
        "Norwegian University of Science and Technology (NTNU), Trondheim, Norway",
        style="Author",
    )
    add_para(doc, "April 2026", style="Author")

    # ── 1. Purpose ──
    doc.add_heading("1. Purpose", level=1)
    add_para(
        doc,
        "This note addresses a question raised by Professor Wei Yan (DTU) regarding "
        "the novelty of the fully implicit CPA algorithm proposed by Igben et al. [1] "
        "relative to the earlier work of Michelsen [2]. Specifically, Professor Yan "
        "observed that Michelsen (2006) already described solving volume and site "
        "fractions as coupled independent variables, and expressed concern that Igben "
        "et al. may have presented Michelsen's method as their own novelty. After "
        "careful re-reading of both papers and independent implementation of all "
        "algorithms in the NeqSim library, we offer the following analysis.",
    )

    # ── 2. Summary of Michelsen (2006) ──
    doc.add_heading("2. Summary of Michelsen (2006)", level=1)
    add_para(
        doc,
        "Michelsen [2] presented a comprehensive treatment of solution procedures for "
        "association models, covering four distinct approaches in a five-page paper:",
    )

    doc.add_heading("2.1 Successive substitution (his primary method)", level=2)
    add_para(
        doc,
        "The site fractions X_k are updated iteratively at fixed volume using a simple "
        "fixed-point iteration. This is rapidly convergent when the association constants "
        "K_kl are small (gas-phase, weak association) but can require thousands of "
        "iterations when K_kl is large (liquid organic acids, strong association).",
    )

    doc.add_heading("2.2 Damped successive substitution", level=2)
    add_para(
        doc,
        'To address slow convergence, Michelsen introduced a damping parameter \u03c9 '
        'in the range 0.2\u20130.3 as a "reasonable compromise" between convergence '
        "speed and stability.",
    )

    doc.add_heading("2.3 Second-order (Newton) method for site fractions at fixed volume", level=2)
    add_para(
        doc,
        "Michelsen formulated the Q-function maximization as an unconstrained optimization "
        "problem and showed that a modified Hessian \u0124 provides a globally convergent "
        "Newton direction. His key contribution here was proving that \u0124 is always "
        "negative definite (guaranteeing ascent) and equals the true Hessian H at the "
        "solution (guaranteeing quadratic convergence).",
    )

    doc.add_heading("2.4 The nested-loop volume solver", level=2)
    add_para(
        doc,
        "For finding the molar volume, Michelsen described a nested-loop approach: an "
        "outer loop iterates on \u03b6 = b/v using a Newton-based method, and at each "
        "outer step the site fractions are resolved to convergence in an inner loop. "
        "Crucially, Michelsen stated (p. 8452):",
    )
    add_quote(
        doc,
        '"The tolerance for accepting an inner loop solution in the volume iteration '
        "is set fairly loose, and, as a consequence, only a single inner-loop iteration "
        "is necessary in most cases. This implies that the efficiency of the nested loop "
        "approach is only slightly lower than the more risky procedure of solving for V "
        'and X simultaneously."',
    )
    add_para(doc, "And further:")
    add_quote(
        doc,
        '"The safe procedure described earlier is only used in the few cases where a '
        "direct use of the Newton-based procedure fails to solve the association "
        'equations in a few iterations."',
    )

    doc.add_heading("2.5 What Michelsen described but did NOT detail", level=2)
    add_para(
        doc,
        "Michelsen clearly acknowledged the idea of solving V and X simultaneously. "
        'In his own words, he described it as "the more risky procedure" and contrasted '
        "it with his preferred nested approach. However, Michelsen did not:",
    )
    add_numbered_item(
        doc, 1,
        "Write down the coupled (n_s + 1)-dimensional system explicitly",
    )
    add_numbered_item(
        doc, 2,
        "Derive the Jacobian matrix for the coupled system",
    )
    add_numbered_item(
        doc, 3,
        "Describe the four-block structure of that Jacobian",
    )
    add_numbered_item(
        doc, 4,
        "Derive the off-diagonal blocks (\u2202R_k/\u2202\u03b6 and "
        "\u2202R_\u0394P/\u2202X_k) analytically",
    )
    add_numbered_item(doc, 5, "Implement and benchmark the coupled solver")
    add_numbered_item(
        doc, 6, "Describe restart criteria for the coupled system"
    )
    add_numbered_item(
        doc, 7, "Report computational timings or speedup factors"
    )
    add_para(
        doc,
        "His paper focused entirely on the inner-loop efficiency \u2014 successive "
        "substitution, damped SS, and second-order (Newton) methods for solving site "
        "fractions at fixed volume \u2014 and demonstrated that a loose-tolerance inner "
        "loop made the nested approach nearly as fast as direct coupling.",
    )

    # ── 3. Summary of Igben et al. (2026) ──
    doc.add_heading("3. Summary of Igben et al. (2026)", level=1)
    add_para(
        doc,
        "Igben et al. [1] presented four algorithms explicitly:",
    )
    for label, desc in [
        ("Algorithm 1", "Volume initialization (standard)"),
        ("Algorithm 2", "Successive substitution method (nested loops \u2014 Michelsen's basic approach)"),
        ("Algorithm 3", "Second-order method (nested loops with Newton inner solve \u2014 Michelsen's second-order method)"),
        ("Algorithm 4", "Fully implicit method (coupled Newton\u2013Raphson for V and X simultaneously)"),
    ]:
        p = doc.add_paragraph(style="Normal")
        run = p.add_run(f"\u2022 {label}: ")
        run.bold = True
        p.add_run(desc)

    add_para(
        doc,
        "Algorithm 4 is the authors' claimed contribution. It:",
    )
    add_numbered_item(
        doc, 1,
        "Writes the coupled (n_s + 1)-dimensional nonlinear system explicitly (their Eq. 18)",
    )
    add_numbered_item(
        doc, 2,
        "Derives all Jacobian entries analytically (their Eqs. 20\u201325)",
    )
    add_numbered_item(
        doc, 3,
        "Introduces a restart threshold parameter \u03b1 to suppress unnecessary root "
        "searches in supercritical regions (Algorithm 4, lines 8\u20139)",
    )
    add_numbered_item(
        doc, 4,
        "Applies step limiting with bisection (Algorithm 4, lines 10\u201311)",
    )
    add_numbered_item(
        doc, 5,
        "Benchmarks all four algorithms systematically on pure water, water\u2013methanol, "
        "and water\u2013ethanol\u2013acetic acid systems",
    )
    add_numbered_item(
        doc, 6,
        "Reports 30\u201380% computational cost reduction vs. the nested approaches",
    )

    # ── 4. Analysis ──
    doc.add_heading(
        "4. Analysis: Is the Igben et al. Contribution Genuinely Novel?", level=1
    )

    doc.add_heading("4.1 What is NOT novel", level=2)
    add_para(
        doc,
        "Professor Yan is correct that the idea of solving volume and site fractions "
        "simultaneously was known to Michelsen. This is explicitly stated in Michelsen "
        "(2006). The successive substitution (Algorithm 2) and second-order (Algorithm 3) "
        "methods in Igben et al. are straightforward reproductions of Michelsen's methods, "
        "which the authors acknowledge by citing Michelsen [2] (as their reference [22]).",
    )

    doc.add_heading(
        "4.2 What IS novel (or at least constitutes a useful engineering contribution)",
        level=2,
    )
    add_para(
        doc,
        "However, there is a meaningful distinction between mentioning the idea of "
        "coupled solution and working it out in detail. We identify the following "
        "contributions in Igben et al. that go beyond what Michelsen published:",
    )
    add_numbered_item(
        doc, 1,
        "Explicit derivation of the coupled Jacobian: Michelsen never wrote down the "
        "(n_s+1) \u00d7 (n_s+1) Jacobian. Igben et al. derived all four blocks "
        "analytically (Eqs. 20\u201325), including the cross-derivatives "
        "\u2202R_k/\u2202\u03b6 (their Eq. 21) and \u2202R_\u0394P/\u2202X_k "
        "(their Eq. 23). These derivatives are non-trivial because they involve "
        "the volume dependence of the radial distribution function g(\u03c1) through "
        "\u2202K_k\u2113/\u2202\u03b6.",
    )
    add_numbered_item(
        doc, 2,
        "Restart criterion for supercritical detection: The threshold parameter \u03b1 "
        "(Algorithm 4, lines 8\u20139) that detects when the first Newton correction "
        "produces non-physical site fractions and triggers a restart with the opposite "
        "phase initial guess is a practical robustness feature not discussed by Michelsen.",
    )
    add_numbered_item(
        doc, 3,
        "Systematic quantitative comparison: Michelsen stated qualitatively that the "
        'nested approach was "only slightly lower" in efficiency than direct coupling, '
        "but provided no timing data. Igben et al. provided detailed iteration counts, "
        "relative timing comparisons across T\u2013P grids, and demonstrated that the "
        "speedup is condition-dependent (larger for liquid-like roots, near critical "
        "points, and complex mixtures).",
    )
    add_numbered_item(
        doc, 4,
        "Step limiting strategy: The bisection-based step control for the coupled system "
        "adapts Michelsen's Q-function bisection to the coupled setting.",
    )

    doc.add_heading("4.3 A nuanced assessment", level=2)
    add_para(
        doc,
        'Michelsen\'s 2006 statement that the nested loop is "only slightly lower" in '
        "efficiency was based on the observation that with a loose inner-loop tolerance, "
        '"only a single inner-loop iteration is necessary in most cases." This is '
        "accurate for systems with moderate association strength where successive "
        "substitution converges rapidly.",
    )
    add_para(
        doc,
        "However, our own implementation experience (documented in [3]) reveals that "
        "this assessment does not hold for all systems:",
    )
    p = doc.add_paragraph(style="Normal")
    run = p.add_run("\u2022 Strongly associating liquids: ")
    run.bold = True
    p.add_run(
        "For acetic acid and glycols, the inner successive substitution can require "
        "15\u201330 iterations per outer step, making the nested loop substantially more expensive."
    )
    p = doc.add_paragraph(style="Normal")
    run = p.add_run("\u2022 Binary/ternary associating mixtures: ")
    run.bold = True
    p.add_run(
        "For water\u2013ethanol\u2013acetic acid, the inner Newton refinement adds "
        "significant cost because it requires matrix assembly and factorization at each "
        "outer step."
    )
    p = doc.add_paragraph(style="Normal")
    run = p.add_run("\u2022 Implicit function theorem derivatives: ")
    run.bold = True
    p.add_run(
        "The derivatives H\u00b7X_V = \u2212\u2202g/\u2202V needed by the outer "
        "Halley step in the nested approach require an additional linear system solve "
        "per outer iteration that is avoided entirely in the coupled formulation."
    )
    add_para(
        doc,
        "These observations explain why our benchmarks [3] found speedups of 2\u201333\u00d7 "
        'rather than Michelsen\'s qualitative assessment of "only slightly lower."',
    )

    doc.add_heading(
        "4.4 The key insight Michelsen may not have fully explored", level=2
    )
    add_para(
        doc,
        "The most important practical benefit of the fully implicit approach is not "
        "broadly treated by Michelsen: the complete elimination of the inner loop removes "
        "the need to compute volume derivatives of the association contribution "
        "(\u2202F_assoc/\u2202V, \u2202\u00b2F_assoc/\u2202V\u00b2) via the implicit "
        "function theorem. In the nested approach, these derivatives require solving a "
        "linear system at every outer iteration. In the coupled approach, the coupling "
        "between V and X is handled directly through the off-diagonal Jacobian blocks, "
        "eliminating this additional solve. For systems with many association sites "
        "(n_s \u2265 6), this saves substantial computation.",
    )

    # ── 5. Our Assessment ──
    doc.add_heading("5. Our Assessment and Recommendation", level=1)

    doc.add_heading("5.1 Fair characterization", level=2)
    add_para(doc, "We believe a fair characterization is:")
    p = doc.add_paragraph(style="Normal")
    p.add_run(
        "\u2022 Igben et al. did not invent the idea of solving V and X simultaneously "
        "\u2014 Michelsen clearly described this possibility in 2006."
    )
    p = doc.add_paragraph(style="Normal")
    p.add_run(
        "\u2022 Igben et al. did provide a complete, detailed, implementable algorithm "
        "with analytic Jacobian, robustness features, and quantitative benchmarks that "
        "Michelsen did not publish."
    )
    p = doc.add_paragraph(style="Normal")
    p.add_run(
        "\u2022 The characterization in Igben et al. of their Algorithm 3 (second-order "
        "method with nested loops) is a faithful reproduction of Michelsen's method, "
        "not a misattribution. They cite Michelsen appropriately."
    )
    p = doc.add_paragraph(style="Normal")
    p.add_run(
        '\u2022 The characterization of Algorithm 4 (fully implicit) as "new" is '
        "partially justified: the detailed algorithmic formulation, Jacobian derivation, "
        "restart strategy, and systematic benchmarking are new contributions, even though "
        "the underlying concept was known."
    )

    doc.add_heading("5.2 Where Igben et al. could have been more careful", level=2)
    add_para(
        doc,
        "The authors could have more explicitly acknowledged that Michelsen (2006) "
        "already discussed direct simultaneous solution as an alternative. Their abstract "
        'states the method "proposes a fully implicit solution strategy" without noting '
        "that the concept was mentioned (though not detailed) by Michelsen twenty years "
        "earlier. A more precise claim would be:",
    )
    add_quote(
        doc,
        '"This work provides the first complete algorithmic specification, analytic '
        "Jacobian derivation, and systematic quantitative evaluation of the fully "
        "implicit approach for CPA, building on the concept mentioned by Michelsen "
        '(2006)."',
    )

    doc.add_heading("5.3 Analogy", level=2)
    add_para(
        doc,
        "The situation is comparable to many cases in numerical methods where an idea "
        "is mentioned in passing by one author and later worked out in full detail by "
        "another. The second contribution has genuine engineering value \u2014 it takes "
        "the algorithm from concept to implementation \u2014 even though intellectual "
        "priority for the idea belongs to the first author.",
    )

    # ── 6. Impact ──
    doc.add_heading("6. Impact on Our Work", level=1)
    add_para(
        doc,
        "Our paper [3] already attributes the algorithm to Igben et al. [1] as the "
        "primary reference while citing Michelsen [2] as the foundational work on CPA "
        "solution procedures. Our paper focuses on:",
    )
    add_numbered_item(
        doc, 1,
        "An independent implementation in NeqSim with specific engineering optimizations "
        "(hand-coded Gaussian elimination, gas-phase fallback to nested solver)",
    )
    add_numbered_item(
        doc, 2,
        "Extension to 11 industrial fluid systems not covered by Igben et al.",
    )
    add_numbered_item(
        doc, 3,
        "Demonstration that the speedup advantages persist for realistic oil and gas "
        "compositions with glycol inhibitors (MEG, TEG)",
    )
    add_para(
        doc,
        "These contributions are independent of the priority question between Michelsen "
        "and Igben et al. Regardless of how the novelty of Method 4 is assessed, the "
        "practical benchmarking and industrial extension in our paper remain valid and useful.",
    )

    doc.add_heading("6.1 Suggested revision to our paper", level=2)
    add_para(
        doc,
        "To address Professor Yan's concern transparently, we propose adding the "
        "following paragraph to Section 1.3 of our paper:",
    )
    add_quote(
        doc,
        '"It should be noted that Michelsen [8] already mentioned the possibility of '
        "solving volume and site fractions simultaneously, describing it as 'the more "
        "risky procedure' compared to the nested loop approach. However, Michelsen did "
        "not derive the coupled Jacobian, implement the algorithm, or provide "
        "quantitative benchmarks. Igben et al. [10] provided the first complete "
        "algorithmic specification with analytic Jacobian, restart strategy, and "
        "systematic performance evaluation. Our implementation and industrial benchmarks "
        'build on the detailed formulation of Igben et al."',
    )

    # ── 7. Conclusion ──
    doc.add_heading("7. Conclusion", level=1)
    add_para(
        doc,
        "Professor Yan's observation is substantively correct: Michelsen (2006) was "
        "aware of and briefly discussed the idea of simultaneous volume\u2013site "
        "fraction solution. However, the detailed algorithmic formulation, Jacobian "
        "derivation, robustness strategies, and quantitative benchmarking provided by "
        "Igben et al. (2026) represent genuine engineering contributions beyond what "
        "Michelsen published. The situation is one of incremental novelty built on an "
        "acknowledged concept rather than a misunderstanding or misattribution of "
        "Michelsen's work. Both papers deserve appropriate credit for their respective "
        "contributions.",
    )

    # ── References ──
    doc.add_heading("References", level=1)
    refs = [
        "[1] O.N. Igben, W.Q. Barros, A.M. Moreno S.J., G. Malgaresi, S.M. Sheth, "
        "Fully implicit algorithm for the cubic-plus-association equation of state, "
        "Fluid Phase Equilib. 608 (2026) 114734.",
        "[2] M.L. Michelsen, Robust and efficient solution procedures for association "
        "models, Ind. Eng. Chem. Res. 45 (2006) 8449\u20138453.",
        "[3] E. Solbraa, Evaluation of a fully implicit CPA algorithm for industrially "
        "relevant associating fluid systems, manuscript (2026). NeqSim implementation: "
        "https://github.com/equinor/neqsim.",
        "[4] M.L. Michelsen, E.M. Hendriks, Physical properties from association models, "
        "Fluid Phase Equilib. 180 (2001) 165\u2013174.",
        "[5] G.M. Kontogeorgis, M.L. Michelsen, G.K. Folas, S.O. Derawi, N. von Solms, "
        "E.H. Stenby, Ten years with the CPA (Cubic-Plus-Association) equation of state. "
        "Part 1, Ind. Eng. Chem. Res. 45 (2006) 4855\u20134868.",
    ]
    for ref in refs:
        add_para(doc, ref)

    doc.save(str(OUTPUT))
    print(f"Word document saved to: {OUTPUT}")


if __name__ == "__main__":
    build_document()
