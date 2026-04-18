"""
book_render_word — Render a multi-chapter book to Word (.docx).

Scientific book formatting with:
- Native Word equation rendering via OMML (LaTeX -> MathML -> OMML)
- Chapter-numbered display equations (Eq. 5.1, 5.2, ...)
- Inline math rendered natively in Word
- Professional page setup for B5/A4 book format
- Justified paragraphs with proper line spacing
- Scientific figure captions (Figure X.Y: ...)
- Booktabs-style table formatting
- Auto-generated TOC field
- Page numbers
- Part separator pages
- Citation resolution and bibliography
"""

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

import book_builder
from citation_utils import (parse_bibtex, collect_all_cited_keys_from_chapters,
                            resolve_citations_numbered_plain, build_key_to_num,
                            format_bib_entry_plain)

try:
    from math_utils import latex_to_omml, latex_to_unicode
    HAS_MATH = True
except ImportError:
    HAS_MATH = False


# ---------------------------------------------------------------------------
# Equation counter
# ---------------------------------------------------------------------------

class EquationCounter:
    """Track display equation numbering per chapter."""

    def __init__(self):
        self.chapter = 0
        self.eq_num = 0

    def set_chapter(self, ch_num):
        self.chapter = ch_num
        self.eq_num = 0

    def next(self):
        self.eq_num += 1
        return "({}.{})".format(self.chapter, self.eq_num)


# ---------------------------------------------------------------------------
# Document setup helpers
# ---------------------------------------------------------------------------

def _setup_document_styles(doc, cfg):
    """Configure document-wide styles for a scientific book."""
    settings = cfg.get("settings", {})
    font_size = settings.get("font_size", 11)
    if isinstance(font_size, str):
        font_size = int(font_size.replace("pt", ""))
    line_spacing = float(settings.get("line_spacing", 1.2))

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(font_size)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.first_line_indent = Cm(0.7)

    # Heading 1 (Chapter titles)
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    h1pf = h1.paragraph_format
    h1pf.space_before = Pt(36)
    h1pf.space_after = Pt(18)
    h1pf.first_line_indent = Cm(0)
    h1pf.page_break_before = False
    h1pf.keep_with_next = True

    # Heading 2 (Sections)
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
    h2pf = h2.paragraph_format
    h2pf.space_before = Pt(18)
    h2pf.space_after = Pt(9)
    h2pf.first_line_indent = Cm(0)
    h2pf.keep_with_next = True

    # Heading 3 (Subsections)
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h3pf = h3.paragraph_format
    h3pf.space_before = Pt(12)
    h3pf.space_after = Pt(6)
    h3pf.first_line_indent = Cm(0)
    h3pf.keep_with_next = True

    # Heading 4
    h4 = doc.styles["Heading 4"]
    h4.font.name = "Times New Roman"
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = True
    h4pf = h4.paragraph_format
    h4pf.space_before = Pt(9)
    h4pf.space_after = Pt(4)
    h4pf.first_line_indent = Cm(0)
    h4pf.keep_with_next = True

    return font_size, line_spacing


def _setup_page(doc, cfg):
    """Configure page layout for the book."""
    settings = cfg.get("settings", {})
    page_size = settings.get("page_size", "a4")

    section = doc.sections[0]
    if page_size == "b5":
        section.page_width = Cm(17.6)
        section.page_height = Cm(25.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.0)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def _add_page_number_footer(doc):
    """Add centered page number to footer."""
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)

    run2 = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = para.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3._element.append(fld_end)

    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# TOC, page breaks
# ---------------------------------------------------------------------------

def _add_toc_field(doc):
    """Insert a TOC field code that auto-updates when opened in Word."""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0)

    run = para.add_run()
    fld_char_begin = parse_xml(
        r'<w:fldChar {} w:fldCharType="begin"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        )
    )
    run._element.append(fld_char_begin)

    run2 = para.add_run()
    instr = parse_xml(
        r'<w:instrText {} xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        )
    )
    run2._element.append(instr)

    run3 = para.add_run()
    fld_char_end = parse_xml(
        r'<w:fldChar {} w:fldCharType="end"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        )
    )
    run3._element.append(fld_char_end)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("[Right-click and Update Field to populate Table of Contents]")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return para


def _add_page_break(doc):
    """Add a simple page break."""
    doc.add_page_break()


def _strip_html_comments(text):
    """Remove HTML comments from text."""
    return re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)


# ---------------------------------------------------------------------------
# OMML equation insertion
# ---------------------------------------------------------------------------

def _insert_display_equation(doc, latex_str, eq_label=None):
    """Insert a display equation as a centered OMML math paragraph.

    Falls back to formatted Cambria Math text if OMML is unavailable.
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.first_line_indent = Cm(0)

    # Try native OMML equation
    if HAS_MATH:
        omml_el = latex_to_omml(latex_str)
        if omml_el is not None:
            para._element.append(omml_el)
            if eq_label:
                _add_equation_number(para, eq_label)
            return para

    # Fallback: Cambria Math formatted text
    clean = latex_str.strip()
    if HAS_MATH:
        clean = latex_to_unicode(clean)

    run = para.add_run(clean)
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    run.italic = True

    if eq_label:
        _add_equation_number(para, eq_label)

    return para


def _add_equation_number(para, label):
    """Add right-aligned equation number to a display equation paragraph."""
    run = para.add_run("\t" + label)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = False

    pPr = para._element.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_center = OxmlElement("w:tab")
    tab_center.set(qn("w:val"), "center")
    tab_center.set(qn("w:pos"), "4536")
    tabs.append(tab_center)
    tab_right = OxmlElement("w:tab")
    tab_right.set(qn("w:val"), "right")
    tab_right.set(qn("w:pos"), "9072")
    tabs.append(tab_right)
    pPr.append(tabs)


def _insert_inline_math(paragraph, latex_str):
    """Insert inline math into an existing paragraph."""
    if HAS_MATH:
        omml_el = latex_to_omml(latex_str)
        if omml_el is not None:
            paragraph._element.append(omml_el)
            return

    clean = latex_str.strip()
    if HAS_MATH:
        clean = latex_to_unicode(clean)
    run = paragraph.add_run(clean)
    run.font.name = "Cambria Math"
    run.italic = True


# ---------------------------------------------------------------------------
# Figure and table formatting
# ---------------------------------------------------------------------------

def _insert_figure(doc, img_path, caption, chapter_num=None, fig_counter=None):
    """Insert a figure with scientific-style caption."""
    try:
        doc.add_picture(str(img_path), width=Inches(4.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_p.paragraph_format.first_line_indent = Cm(0)
        last_p.paragraph_format.space_before = Pt(12)
    except Exception:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run("[Figure: {}]".format(img_path.name))
        run.italic = True

    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_before = Pt(3)
        cap_para.paragraph_format.space_after = Pt(12)
        cap_para.paragraph_format.first_line_indent = Cm(0)

        if not re.match(r"^Figure\s+\d", caption):
            if fig_counter and chapter_num:
                fig_num = fig_counter[0]
                fig_counter[0] += 1
                run_label = cap_para.add_run(
                    "Figure {}.{}: ".format(chapter_num, fig_num))
                run_label.bold = True
                run_label.font.size = Pt(9)
                run_label.font.name = "Times New Roman"

        run_text = cap_para.add_run(caption)
        run_text.italic = True
        run_text.font.size = Pt(9)
        run_text.font.name = "Times New Roman"


def _render_table_to_doc(doc, table_lines):
    """Render a markdown pipe table with booktabs-style scientific formatting."""
    if len(table_lines) < 2:
        return

    header_cells = [c.strip() for c in table_lines[0].split("|")[1:-1]]
    n_cols = len(header_cells)

    data_rows = []
    for row_line in table_lines[2:]:
        cells = [c.strip() for c in row_line.split("|")[1:-1]]
        if len(cells) == n_cols:
            data_rows.append(cells)

    if not data_rows:
        return

    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(cell_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

    _apply_booktabs_style(table)


def _apply_booktabs_style(table):
    """Apply booktabs-style borders: heavy top/bottom, medium header rule."""
    thick = "12"
    medium = "6"

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")

            if i == 0:
                _set_border(tcBorders, "top", thick)
                _set_border(tcBorders, "bottom", medium)
            elif i == len(table.rows) - 1:
                _set_border(tcBorders, "top", "0", val="none")
                _set_border(tcBorders, "bottom", thick)
            else:
                _set_border(tcBorders, "top", "0", val="none")
                _set_border(tcBorders, "bottom", "0", val="none")

            _set_border(tcBorders, "left", "0", val="none")
            _set_border(tcBorders, "right", "0", val="none")

            tcPr.append(tcBorders)


def _set_border(borders_element, side, size, val="single"):
    """Set a table cell border."""
    el = OxmlElement("w:{0}".format(side))
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), size)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "000000")
    borders_element.append(el)


# ---------------------------------------------------------------------------
# Markdown -> Word conversion (with equation support)
# ---------------------------------------------------------------------------

def _render_md_to_doc(doc, md_text, chapter_num=None, figures_dir=None,
                      eq_counter=None, fig_counter=None):
    """Parse markdown and render to Word with scientific formatting.

    Handles: headings, paragraphs, bold/italic/code, code blocks,
    images, lists, tables, display equations ($$...$$), inline math ($...$),
    and blockquotes.
    """
    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                _add_code_block(doc, code_text, code_lang)
                code_lines = []
                in_code_block = False
                code_lang = ""
            else:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Display math: $$ on its own line or $$...$$
        if stripped.startswith("$$"):
            if stripped == "$$":
                eq_lines = []
                i += 1
                while i < len(lines):
                    if lines[i].strip() == "$$":
                        break
                    eq_lines.append(lines[i])
                    i += 1
                i += 1
                latex = "\n".join(eq_lines).strip()
            elif stripped.endswith("$$") and len(stripped) > 4:
                latex = stripped[2:-2].strip()
                i += 1
            else:
                eq_lines = [stripped[2:]]
                i += 1
                while i < len(lines):
                    if lines[i].strip().endswith("$$"):
                        eq_lines.append(lines[i].strip()[:-2])
                        break
                    eq_lines.append(lines[i])
                    i += 1
                i += 1
                latex = "\n".join(eq_lines).strip()

            eq_label = eq_counter.next() if eq_counter else None
            _insert_display_equation(doc, latex, eq_label)
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.+)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            bq_lines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                bq_lines.append(lines[i].strip()[2:])
                i += 1
            bq_text = " ".join(bq_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _add_rich_text_with_math(p, bq_text)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Images
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            caption = img_match.group(1)
            img_path_str = img_match.group(2)
            if figures_dir:
                full_path = Path(figures_dir) / Path(img_path_str).name
                if full_path.exists():
                    _insert_figure(doc, full_path, caption,
                                   chapter_num, fig_counter)
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Cm(0)
                    run = p.add_run("[Figure: {}]".format(img_path_str))
                    run.italic = True
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.add_run("[Figure: {}]".format(img_path_str))
                run.italic = True
            i += 1
            continue

        # Italic caption lines: *Figure X.Y: ...*
        if stripped.startswith("*") and stripped.endswith("*") and len(stripped) > 2:
            inner = stripped[1:-1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(inner)
            run.italic = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            _clear_and_add_rich(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.first_line_indent = Cm(0)
            _clear_and_add_rich(p, text)
            i += 1
            continue

        # Table
        if "|" in stripped and stripped.startswith("|"):
            table_lines = []
            while (i < len(lines) and "|" in lines[i].strip()
                   and lines[i].strip().startswith("|")):
                table_lines.append(lines[i].strip())
                i += 1
            _render_table_to_doc(doc, table_lines)
            continue

        # Regular paragraph
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if (not next_line or next_line.startswith("#")
                    or next_line.startswith("```")
                    or re.match(r"^!\[", next_line)
                    or next_line.startswith("- ")
                    or next_line.startswith("* ")
                    or re.match(r"^\d+\.\s", next_line)
                    or next_line.startswith("$$")
                    or next_line.startswith("> ")
                    or ("|" in next_line and next_line.startswith("|"))):
                break
            para_lines.append(next_line)
            i += 1

        para_text = " ".join(para_lines)
        p = doc.add_paragraph()
        _add_rich_text_with_math(p, para_text)

    return doc


def _clear_and_add_rich(paragraph, text):
    """Clear existing runs and add rich text to a paragraph."""
    for run in paragraph.runs:
        run.text = ""
    _add_rich_text_with_math(paragraph, text)


def _add_code_block(doc, code_text, language=""):
    """Add a formatted code block to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0

    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)


def _add_rich_text_with_math(paragraph, text):
    """Add text with inline formatting and inline math to a paragraph.

    Handles: **bold**, *italic*, `code`, and $inline math$.
    """
    parts = _split_inline_math(text)
    for content, is_math in parts:
        if is_math:
            _insert_inline_math(paragraph, content)
        else:
            _add_rich_text(paragraph, content)


def _split_inline_math(text):
    """Split text into (content, is_math) pairs for inline $...$ math."""
    result = []
    pos = 0

    while pos < len(text):
        idx = text.find("$", pos)
        if idx == -1:
            result.append((text[pos:], False))
            break

        # Skip $$ (display math)
        if idx + 1 < len(text) and text[idx + 1] == "$":
            result.append((text[pos:idx + 2], False))
            pos = idx + 2
            continue

        # Skip escaped \$
        if idx > 0 and text[idx - 1] == "\\":
            result.append((text[pos:idx + 1], False))
            pos = idx + 1
            continue

        if idx > pos:
            result.append((text[pos:idx], False))

        close = text.find("$", idx + 1)
        if close == -1:
            result.append((text[idx:], False))
            break

        math_content = text[idx + 1:close]
        if len(math_content) < 200 and "\n" not in math_content:
            result.append((math_content, True))
        else:
            result.append((text[idx:close + 1], False))
        pos = close + 1

    return result


def _add_rich_text(paragraph, text):
    """Add text with bold, italic, and code formatting."""
    if not text:
        return
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Times New Roman"
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = "Times New Roman"
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(part)
            run.font.name = "Times New Roman"


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def render_book_word(book_dir, chapter_filter=None):
    """Render the book to a Word document with scientific formatting.

    Parameters
    ----------
    book_dir : str or Path
        Path to the book project directory.
    chapter_filter : str or None
        If given, render only the chapter whose ``dir`` matches.

    Returns
    -------
    Path or None
        Path to the generated .docx file, or None on failure.
    """
    if not HAS_DOCX:
        print("[book_render_word] Error: python-docx not installed. "
              "Run: pip install python-docx")
        return None

    book_dir = Path(book_dir)
    cfg = book_builder.load_book_config(book_dir)

    doc = Document()

    _setup_document_styles(doc, cfg)
    _setup_page(doc, cfg)

    eq_counter = EquationCounter()

    # Load bibliography and build citation numbering
    bib_path = book_dir / "refs.bib"
    bib_entries = parse_bibtex(bib_path)
    all_cited_keys = collect_all_cited_keys_from_chapters(book_dir, cfg)
    key_to_num = build_key_to_num(all_cited_keys)

    # Title page
    if not chapter_filter:
        title_md = book_dir / "frontmatter" / "title_page.md"
        if title_md.exists():
            text = _strip_html_comments(title_md.read_text(encoding="utf-8"))
            _render_md_to_doc(doc, text)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(120)
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(cfg.get("title", "Untitled"))
            run.bold = True
            run.font.size = Pt(24)
            run.font.name = "Times New Roman"

            subtitle = cfg.get("subtitle", "")
            if subtitle:
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p2.paragraph_format.first_line_indent = Cm(0)
                run2 = p2.add_run(subtitle)
                run2.font.size = Pt(14)
                run2.font.name = "Times New Roman"
                run2.italic = True

            authors = cfg.get("authors", [])
            if authors:
                p3 = doc.add_paragraph()
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p3.paragraph_format.space_before = Pt(36)
                p3.paragraph_format.first_line_indent = Cm(0)
                for a in authors:
                    run_name = p3.add_run(a.get("name", ""))
                    run_name.font.size = Pt(14)
                    run_name.font.name = "Times New Roman"
                    aff = a.get("affiliation", "")
                    if aff:
                        p3.add_run("\n")
                        run_aff = p3.add_run(aff)
                        run_aff.font.size = Pt(11)
                        run_aff.font.name = "Times New Roman"
                        run_aff.italic = True

        _add_page_break(doc)

        # Copyright page
        copyright_md = book_dir / "frontmatter" / "copyright.md"
        if copyright_md.exists():
            text = _strip_html_comments(copyright_md.read_text(encoding="utf-8"))
            _render_md_to_doc(doc, text)
            _add_page_break(doc)

        # Table of Contents
        doc.add_heading("Contents", level=1)
        _add_toc_field(doc)
        _add_page_break(doc)

        # Other frontmatter
        for fm in cfg.get("frontmatter", []):
            if fm in ("title_page", "copyright"):
                continue
            fm_path = book_dir / "frontmatter" / f"{fm}.md"
            if fm_path.exists():
                text = _strip_html_comments(fm_path.read_text(encoding="utf-8"))
                _render_md_to_doc(doc, text)
                _add_page_break(doc)

    # Chapters
    prev_part = None
    for ch_num, ch, part_title in book_builder.iter_chapters(cfg):
        if chapter_filter and ch["dir"] != chapter_filter:
            continue

        if not chapter_filter and part_title and part_title != prev_part:
            _add_page_break(doc)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(180)
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(part_title)
            run.bold = True
            run.font.size = Pt(22)
            run.font.name = "Times New Roman"
            _add_page_break(doc)
            prev_part = part_title

        eq_counter.set_chapter(ch_num)
        fig_counter = [1]

        ch_dir = book_builder.resolve_chapter_dir(book_dir, ch)
        ch_md = ch_dir / "chapter.md"

        if not ch_md.exists():
            doc.add_heading(
                "Chapter {}: {}".format(ch_num, ch.get("title", "Untitled")),
                level=1)
            p = doc.add_paragraph()
            run = p.add_run("Content not yet written.")
            run.italic = True
            continue

        text = ch_md.read_text(encoding="utf-8")
        text = _strip_html_comments(text)
        text = book_builder.strip_heading_numbers(text)

        # Resolve citations: \cite{key} -> [N]
        if key_to_num:
            text = resolve_citations_numbered_plain(text, key_to_num)
        # Strip empty "## References" placeholder (bibliography at end)
        text = re.sub(r'##\s+References\s*\n?(?:\s*\n)*', '', text)

        figures_dir = ch_dir / "figures"
        _render_md_to_doc(
            doc, text, chapter_num=ch_num,
            figures_dir=figures_dir if figures_dir.exists() else None,
            eq_counter=eq_counter, fig_counter=fig_counter)

    # Bibliography section
    if all_cited_keys and bib_entries:
        _add_page_break(doc)
        doc.add_heading("References", level=1)
        for key in all_cited_keys:
            num = key_to_num[key]
            fields = bib_entries.get(key, {})
            if fields:
                entry_text = format_bib_entry_plain(fields)
            else:
                entry_text = "[{}] -- not in refs.bib".format(key)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-1.0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run("[{}]  {}".format(num, entry_text))
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

    # Backmatter
    if not chapter_filter:
        for bm in cfg.get("backmatter", []):
            bm_path = book_dir / "backmatter" / f"{bm}.md"
            if bm_path.exists():
                _add_page_break(doc)
                text = _strip_html_comments(bm_path.read_text(encoding="utf-8"))
                _render_md_to_doc(doc, text)

    # Page numbers
    _add_page_number_footer(doc)

    # Save
    submission_dir = book_dir / "submission"
    submission_dir.mkdir(parents=True, exist_ok=True)

    out_name = "book.docx" if not chapter_filter else f"{chapter_filter}.docx"
    out_path = submission_dir / out_name

    try:
        doc.save(str(out_path))
    except PermissionError:
        from datetime import datetime
        ts = datetime.now().strftime("%H%M%S")
        out_path = submission_dir / f"book_{ts}.docx"
        doc.save(str(out_path))
        print("  NOTE: book.docx was locked, saved as {}".format(
            out_path.name))

    size_kb = out_path.stat().st_size / 1024
    print("[book_render_word] Word document generated: {}  ({:.0f} KB)".format(
        out_path, size_kb))
    return out_path
