"""
Render paper.md to all submission formats: HTML, LaTeX, and Word (.docx).

Usage:
    python render_all.py <paper_dir>
    python render_all.py papers/tpflash_algorithms_2026
    python render_all.py papers/gibbs_minimization_2026

Produces in <paper_dir>/submission/:
    paper.html   — Standalone HTML with KaTeX math rendering
    paper.tex    — Elsarticle LaTeX (requires journal profile)
    paper.docx   — Word document with styled headings, tables, and figures
"""
import re
import sys
import os
from pathlib import Path

# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------

def parse_sections(text):
    """Parse markdown into list of {level, title, content} dicts."""
    sections = []
    current = None
    lines = []

    for line in text.split("\n"):
        m = re.match(r'^(#{1,3})\s+(.+)$', line)
        if m:
            if current is not None:
                sections.append({
                    "level": current["level"],
                    "title": current["title"],
                    "content": "\n".join(lines).strip(),
                })
            current = {"level": len(m.group(1)), "title": m.group(2).strip()}
            lines = []
        else:
            lines.append(line)

    if current is not None:
        sections.append({
            "level": current["level"],
            "title": current["title"],
            "content": "\n".join(lines).strip(),
        })
    return sections


def strip_comments(text):
    """Remove HTML comments."""
    return re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)


def extract_title(sections):
    """Get the H1 title."""
    for s in sections:
        if s["level"] == 1:
            return s["title"]
    return "Untitled"


def find_figures(paper_dir):
    """Find all figure PNGs in the figures/ directory, sorted."""
    fig_dir = Path(paper_dir) / "figures"
    if not fig_dir.exists():
        return []
    figs = sorted(fig_dir.glob("fig*.png"))
    return figs


def get_figure_captions(paper_dir):
    """Load figure captions from results.json if available."""
    rj = Path(paper_dir) / "results.json"
    if rj.exists():
        import json
        data = json.loads(rj.read_text(encoding="utf-8"))
        return data.get("figure_captions", {})
    return {}


def get_journal_profile(paper_dir):
    """Load journal profile from plan.json if available."""
    plan_file = Path(paper_dir) / "plan.json"
    if plan_file.exists():
        import json
        plan = json.loads(plan_file.read_text(encoding="utf-8"))
        journal_name = plan.get("target_journal", "")
        if journal_name:
            journals_dir = Path(paper_dir).parent.parent / "journals"
            profile_path = journals_dir / f"{journal_name}.yaml"
            if profile_path.exists():
                try:
                    import yaml
                    with open(profile_path) as f:
                        return yaml.safe_load(f)
                except ImportError:
                    pass
    return {}


def parse_bibtex(bib_path):
    """Parse a BibTeX file into a dict of key -> {author, title, journal, year, ...}."""
    if not Path(bib_path).exists():
        return {}
    bib_text = Path(bib_path).read_text(encoding="utf-8")
    entries = {}
    for block in re.finditer(r'@\w+\{(\w+),\s*(.*?)\n\}', bib_text, flags=re.DOTALL):
        key = block.group(1)
        fields = {}
        for fld in re.finditer(r'(\w+)\s*=\s*\{(.*?)\}', block.group(2), flags=re.DOTALL):
            fields[fld.group(1).lower()] = fld.group(2).strip()
        entries[key] = fields
    return entries


def format_author_short(author_str):
    """Shorten 'Last, F. and Last2, F.' to 'Last et al.' or 'Last and Last2'."""
    author_str = author_str.replace("{", "").replace("}", "")
    authors = [a.strip() for a in author_str.split(" and ")]
    if len(authors) > 2:
        return authors[0].split(",")[0] + " et al."
    elif len(authors) == 2:
        return authors[0].split(",")[0] + " and " + authors[1].split(",")[0]
    else:
        return authors[0].split(",")[0]


def format_bib_entry_numbered(fields):
    """Format a BibTeX entry for numbered reference list."""
    author = format_author_short(fields.get("author", "Unknown"))
    title = fields.get("title", "Untitled").replace("{", "").replace("}", "")
    journal = fields.get("journal", fields.get("booktitle", "")).replace("{", "").replace("}", "")
    year = fields.get("year", "n.d.")
    volume = fields.get("volume", "")
    pages = fields.get("pages", "")

    parts = [f"{author},"]
    parts.append(f'"{title},"')
    if journal:
        parts.append(f"<em>{journal}</em>,")
    if volume:
        vol_str = f"vol. {volume}"
        if pages:
            vol_str += f", pp. {pages}"
        parts.append(vol_str + ",")
    parts.append(f"{year}.")
    return " ".join(parts)


def format_bib_entry_authoryear(fields):
    """Format a BibTeX entry for author-year reference list (alphabetical)."""
    author = fields.get("author", "Unknown").replace("{", "").replace("}", "")
    title = fields.get("title", "Untitled").replace("{", "").replace("}", "")
    journal = fields.get("journal", fields.get("booktitle", "")).replace("{", "").replace("}", "")
    year = fields.get("year", "n.d.")
    volume = fields.get("volume", "")
    pages = fields.get("pages", "")

    parts = [f"{author},"]
    parts.append(f"{year}.")
    parts.append(f"{title}.")
    if journal:
        parts.append(f"<em>{journal}</em>")
        if volume:
            parts.append(f"{volume}")
            if pages:
                parts[-1] += f", {pages}"
        parts[-1] += "."
    return " ".join(parts)


def resolve_citations(body, paper_dir, citation_style="numbered"):
    """Replace \\cite{key1, key2} with rendered citations and append reference list.

    For 'numbered' style: [1, 2] with numbered reference list.
    For 'authoryear' style: (Author1 et al., 2023; Author2, 2024) with
    alphabetical reference list.

    Returns (body_with_citations, reference_list_html).
    """
    bib_entries = parse_bibtex(Path(paper_dir) / "refs.bib")
    if not bib_entries:
        return body, ""

    if citation_style == "authoryear":
        return _resolve_authoryear(body, bib_entries)
    else:
        return _resolve_numbered(body, bib_entries)


def _resolve_numbered(body, bib_entries):
    """Numbered citation style: [1], [2, 3]."""
    # Build key→number mapping from order of first appearance
    keys_in_order = []
    for m in re.finditer(r'\\cite\{([^}]+)\}', body):
        for key in m.group(1).split(","):
            key = key.strip()
            if key not in keys_in_order:
                keys_in_order.append(key)
    key_to_num = {k: i + 1 for i, k in enumerate(keys_in_order)}

    def replace_cite(match):
        keys = [k.strip() for k in match.group(1).split(",")]
        nums = []
        for k in keys:
            n = key_to_num.get(k)
            if n:
                nums.append(f'<a href="#ref-{n}" class="cite">{n}</a>')
            else:
                nums.append(f'<span class="cite-unknown">{k}</span>')
        return "[" + ", ".join(nums) + "]"

    body = re.sub(r'\\cite\{([^}]+)\}', replace_cite, body)

    # Build reference list
    ref_html = '\n<h2>References</h2>\n<ol class="references">\n'
    for key in keys_in_order:
        n = key_to_num[key]
        fields = bib_entries.get(key, {})
        entry = format_bib_entry_numbered(fields) if fields else f"[{key}] — not in refs.bib"
        ref_html += f'  <li id="ref-{n}" value="{n}">{entry}</li>\n'
    ref_html += "</ol>\n"

    return body, ref_html


def _resolve_authoryear(body, bib_entries):
    """Author-year citation style: (Author et al., 2023)."""
    # Collect all cited keys
    all_keys = set()
    for m in re.finditer(r'\\cite\{([^}]+)\}', body):
        for key in m.group(1).split(","):
            all_keys.add(key.strip())

    def cite_text(key):
        fields = bib_entries.get(key, {})
        if not fields:
            return key
        author = format_author_short(fields.get("author", "Unknown"))
        year = fields.get("year", "n.d.")
        return f"{author}, {year}"

    def replace_cite(match):
        keys = [k.strip() for k in match.group(1).split(",")]
        citations = []
        for k in keys:
            anchor_id = k.replace(" ", "_")
            text = cite_text(k)
            citations.append(f'<a href="#ref-{anchor_id}" class="cite">{text}</a>')
        return "(" + "; ".join(citations) + ")"

    body = re.sub(r'\\cite\{([^}]+)\}', replace_cite, body)

    # Build alphabetical reference list
    def sort_key(key):
        fields = bib_entries.get(key, {})
        author = fields.get("author", "ZZZ")
        year = fields.get("year", "9999")
        return (author.lower(), year)

    sorted_keys = sorted(all_keys, key=sort_key)

    ref_html = '\n<h2>References</h2>\n<div class="references">\n'
    for key in sorted_keys:
        fields = bib_entries.get(key, {})
        entry = format_bib_entry_authoryear(fields) if fields else f"[{key}] — not in refs.bib"
        anchor_id = key.replace(" ", "_")
        ref_html += f'<p id="ref-{anchor_id}" class="ref-entry">{entry}</p>\n'
    ref_html += "</div>\n"

    return body, ref_html


# ---------------------------------------------------------------------------
# HTML renderer
# ---------------------------------------------------------------------------

def render_html(paper_dir):
    """Render paper.md to a standalone HTML file with KaTeX math."""
    paper_dir = Path(paper_dir)
    text = (paper_dir / "paper.md").read_text(encoding="utf-8")
    sections = parse_sections(strip_comments(text))
    title = extract_title(sections)
    figures = find_figures(paper_dir)
    captions = get_figure_captions(paper_dir)

    html = []
    html.append(f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{title}</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.css">
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.js"></script>
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/contrib/auto-render.min.js"
  onload="renderMathInElement(document.body, {{
    delimiters: [
      {{left: '$$', right: '$$', display: true}},
      {{left: '$', right: '$', display: false}}
    ]
  }})"></script>
<style>
  body {{ max-width: 900px; margin: 40px auto; font-family: 'Georgia', serif;
         line-height: 1.7; color: #333; padding: 0 20px; }}
  h1 {{ font-size: 1.8em; border-bottom: 2px solid #2196F3; padding-bottom: 10px; }}
  h2 {{ font-size: 1.4em; color: #1565C0; margin-top: 2em; }}
  h3 {{ font-size: 1.15em; color: #1976D2; }}
  table {{ border-collapse: collapse; margin: 1em auto; font-size: 0.95em; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 12px; text-align: left; }}
  th {{ background: #E3F2FD; font-weight: 600; }}
  tr:nth-child(even) {{ background: #FAFAFA; }}
  code {{ background: #f5f5f5; padding: 2px 5px; border-radius: 3px; font-size: 0.9em; }}
  pre {{ background: #263238; color: #eee; padding: 16px; border-radius: 6px;
        overflow-x: auto; font-size: 0.88em; line-height: 1.5; }}
  pre code {{ background: none; color: inherit; padding: 0; }}
  img {{ max-width: 100%; display: block; margin: 1em auto; }}
  .fig-container {{ text-align: center; margin: 2em 0; }}
  .fig-caption {{ font-style: italic; color: #666; margin-top: 0.5em; }}
  hr {{ border: none; border-top: 1px solid #ddd; margin: 2em 0; }}
  strong {{ color: #1565C0; }}
  a.cite {{ color: #1565C0; text-decoration: none; font-weight: 600; }}
  a.cite:hover {{ text-decoration: underline; }}
  .cite-unknown {{ color: red; }}
  ol.references {{ font-size: 0.92em; line-height: 1.6; }}
  ol.references li {{ margin-bottom: 4px; }}
  .references .ref-entry {{ font-size: 0.92em; line-height: 1.6; margin-bottom: 4px;
                            padding-left: 2em; text-indent: -2em; }}
</style>
</head>
<body>
""")

    # Convert markdown body
    body = strip_comments(text)
    body = re.sub(r'^### (.+)$', r'<h3>\1</h3>', body, flags=re.MULTILINE)
    body = re.sub(r'^## (.+)$', r'<h2>\1</h2>', body, flags=re.MULTILINE)
    body = re.sub(r'^# (.+)$', r'<h1>\1</h1>', body, flags=re.MULTILINE)
    body = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', body)
    body = re.sub(r'\*(.+?)\*', r'<em>\1</em>', body)
    body = re.sub(r'```(\w*)\n(.*?)```', r'<pre><code>\2</code></pre>', body, flags=re.DOTALL)
    body = re.sub(r'`([^`]+)`', r'<code>\1</code>', body)
    body = re.sub(r'^---$', '<hr>', body, flags=re.MULTILINE)

    # Inline markdown images: ![alt](src) -> <img>
    # Fix paths: paper.md uses "figures/..." but HTML goes to submission/,
    # so we need "../figures/..." for correct relative resolution.
    def _fix_img_path(m):
        alt = m.group(1)
        src = m.group(2)
        if src.startswith("figures/"):
            src = "../" + src
        return f'<div class="fig-container"><img src="{src}" alt="{alt}"></div>'
    body = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', _fix_img_path, body)

    # Tables
    def _html_table(match):
        rows = [l for l in match.group(0).strip().split("\n")
                if l.strip() and not re.match(r'^\|[\s:\-]+\|$', l)]
        out = '<table>\n'
        for i, row in enumerate(rows):
            cells = [c.strip() for c in row.split('|')[1:-1]]
            tag = 'th' if i == 0 else 'td'
            out += '<tr>' + ''.join('<{0}>{1}</{0}>'.format(tag, c) for c in cells) + '</tr>\n'
        out += '</table>'
        return out

    body = re.sub(r'(\|.+\|(?:\n\|.+\|)+)', _html_table, body)
    body = re.sub(r'^- (.+)$', r'<li>\1</li>', body, flags=re.MULTILINE)

    # Resolve citations (\cite{key} -> numbered or author-year)
    journal_profile = get_journal_profile(paper_dir)
    citation_style = journal_profile.get("citation_style", "numbered")
    body, ref_list_html = resolve_citations(body, paper_dir, citation_style)

    body = re.sub(r'\n\n+', '\n</p>\n<p>\n', body)

    # Figures are already rendered inline where they appear in the text.
    # No separate appendix section needed.

    html.append("<p>\n" + body + "\n</p>")
    html.append(ref_list_html)
    html.append("\n</body>\n</html>")

    out_dir = paper_dir / "submission"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_file = out_dir / "paper.html"
    out_file.write_text("\n".join(html), encoding="utf-8")
    return out_file


# ---------------------------------------------------------------------------
# OMML equation support (LaTeX → MathML → OMML for native Word equations)
# ---------------------------------------------------------------------------

_omml_transform = None

def _init_omml():
    """Lazily load the MML2OMML XSLT transform for converting MathML to OMML."""
    global _omml_transform
    if _omml_transform is not None:
        return True
    try:
        from lxml import etree
        import latex2mathml.converter  # noqa: F401
    except ImportError:
        return False

    xsl_candidates = [
        os.path.expandvars(r'%ProgramFiles%\Microsoft Office\root\Office16\MML2OMML.XSL'),
        os.path.expandvars(r'%ProgramFiles(x86)%\Microsoft Office\root\Office16\MML2OMML.XSL'),
        os.path.expandvars(r'%ProgramFiles%\Microsoft Office\Office16\MML2OMML.XSL'),
    ]
    for path in xsl_candidates:
        if os.path.exists(path):
            xslt_doc = etree.parse(path)
            _omml_transform = etree.XSLT(xslt_doc)
            return True
    return False


def _latex_to_omml(latex_str):
    """Convert a LaTeX math string to an OMML XML element for Word embedding.

    Returns None if the conversion pipeline is unavailable or fails.
    """
    if not _init_omml():
        return None
    try:
        import latex2mathml.converter
        from lxml import etree
        mathml = latex2mathml.converter.convert(latex_str)
        tree = etree.fromstring(mathml.encode('utf-8'))
        omml = _omml_transform(tree)
        return omml.getroot()
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Word renderer
# ---------------------------------------------------------------------------

def render_word(paper_dir):
    """Render paper.md to a Word document (.docx) with styled content."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("  [SKIP] python-docx not installed. Install with: pip install python-docx")
        return None

    paper_dir = Path(paper_dir)
    text = strip_comments((paper_dir / "paper.md").read_text(encoding="utf-8"))
    sections = parse_sections(text)
    title = extract_title(sections)
    figures = find_figures(paper_dir)
    captions = get_figure_captions(paper_dir)

    has_omml = _init_omml()
    if has_omml:
        print("  [INFO] Native Word equations enabled (LaTeX → MathML → OMML)")
    else:
        print("  [WARN] OMML equations unavailable — equations rendered as italic text")
        print("         Fix: pip install latex2mathml lxml  (and ensure MS Office is installed)")

    from docx.oxml.ns import qn as _qn
    try:
        from lxml import etree as _etree
    except ImportError:
        from xml.etree import ElementTree as _etree

    doc = Document()

    # -- Page layout: 1-inch margins, A4 --
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # -- Document styles --
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for lvl in range(1, 4):
        try:
            hs = doc.styles[f'Heading {lvl}']
            hs.font.name = 'Times New Roman'
            hs.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass

    # -- Footer with page numbers --
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add PAGE field
    run = fp.add_run()
    fld_char_begin = _etree.SubElement(run._element, _qn('w:fldChar'))
    fld_char_begin.set(_qn('w:fldCharType'), 'begin')
    run2 = fp.add_run()
    instr = _etree.SubElement(run2._element, _qn('w:instrText'))
    instr.set(_qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run3 = fp.add_run()
    fld_char_end = _etree.SubElement(run3._element, _qn('w:fldChar'))
    fld_char_end.set(_qn('w:fldCharType'), 'end')

    # Equation counter for numbered display equations
    eq_counter = [0]

    # -- Title --
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- Generation date --
    from datetime import datetime
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp.paragraph_format.space_before = Pt(0)
    dp.paragraph_format.space_after = Pt(12)
    run = dp.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(100, 100, 100)

    # -- Process each section --

    # ----- Citation support for Word output -----
    journal_profile = get_journal_profile(paper_dir)
    _word_citation_style = journal_profile.get("citation_style", "numbered")
    _word_bib_entries = parse_bibtex(paper_dir / "refs.bib")
    _word_cited_keys = []  # track order of appearance

    def _resolve_cite_for_word(text):
        """Replace \\cite{key1, key2} with plain-text citations for Word output."""
        if not _word_bib_entries:
            return text

        def _replace(match):
            keys = [k.strip() for k in match.group(1).split(",")]
            parts = []
            for k in keys:
                if k not in _word_cited_keys:
                    _word_cited_keys.append(k)
                fields = _word_bib_entries.get(k, {})
                if not fields:
                    parts.append(k)
                    continue
                if _word_citation_style == "authoryear":
                    author = format_author_short(fields.get("author", "Unknown"))
                    year = fields.get("year", "n.d.")
                    parts.append(f"{author}, {year}")
                else:
                    num = _word_cited_keys.index(k) + 1
                    parts.append(str(num))
            if _word_citation_style == "authoryear":
                return "(" + "; ".join(parts) + ")"
            else:
                return "[" + ", ".join(parts) + "]"

        return re.sub(r'\\cite\{([^}]+)\}', _replace, text)

    def _add_content(content_text):
        """Add markdown content as Word paragraphs, handling equations, tables, lists."""
        # Pre-process: resolve \cite{} to plain-text citations
        content_text = _resolve_cite_for_word(content_text)

        # Step 1: Extract fenced code blocks BEFORE any other splitting.
        # Code blocks can contain blank lines, $$, |, etc. that would confuse
        # downstream splitters.  Replace them with unique placeholders.
        code_blocks = {}
        code_idx = [0]

        def _stash_code(m):
            key = f'\x00CODE{code_idx[0]}\x00'
            code_blocks[key] = m.group(0)
            code_idx[0] += 1
            return key

        content_text = re.sub(
            r'```\w*\n.*?```', _stash_code, content_text, flags=re.DOTALL)

        # Step 2: Split out display equation blocks ($$...$$)
        parts = re.split(r'(\$\$.*?\$\$)', content_text, flags=re.DOTALL)

        for part in parts:
            part = part.strip()
            if not part:
                continue

            # Display equation block
            if part.startswith('$$') and part.endswith('$$'):
                latex = part[2:-2].strip()
                _add_display_equation(latex)
                continue

            # Step 3: Split remaining text into paragraph blocks
            blocks = re.split(r'\n\n+', part)
            for block in blocks:
                block = block.strip()
                if not block:
                    continue
                # Restore code block placeholders
                if block in code_blocks:
                    raw = code_blocks[block]
                    _add_code_block(raw)
                else:
                    _add_block(block)

    fig_counter = [0]

    def _add_block(block):
        """Add a single block (table, code, list, or paragraph)."""
        # Inline markdown image: ![alt](figures/path.png)
        # May be followed by a caption line like: *Figure 6.* Caption text here
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', block.strip())
        if img_match:
            fig_counter[0] += 1
            img_path = paper_dir / img_match.group(2)
            alt_text = img_match.group(1)
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(5.5))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[Image: {img_match.group(2)} — {e}]")
            else:
                doc.add_paragraph(f"[Image not found: {img_match.group(2)}]")
            # Extract caption from remaining text after the image markdown
            remaining = block[img_match.end():].strip()
            caption = alt_text
            # Match *Figure N.* Caption text  OR  *Figure N. Caption text*
            cap_match = re.match(r'^\*Figure\s+\d+[.:]\*\s*(.+)', remaining, re.DOTALL)
            if not cap_match:
                cap_match = re.match(r'^\*(?:Figure\s+\d+[.:]\s*)?(.+)\*$', remaining, re.DOTALL)
            if cap_match:
                caption = cap_match.group(1).strip()
            # Add numbered figure caption
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_after = Pt(12)
            run = cap_p.add_run(f"Figure {fig_counter[0]}. ")
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run = cap_p.add_run(caption)
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            return

        # Standalone italic figure caption (already handled above with image)
        stripped = block.strip()
        if stripped.startswith('*') and 'Figure' in stripped and re.match(r'^\*Figure\s+\d+', stripped):
            return

        # Table block
        if block.startswith('|') and '|' in block:
            _add_table(block)
            return

        # Code block (in case one slipped through without blank lines)
        if block.startswith('```'):
            _add_code_block(block)
            return

        # Horizontal rule
        if block == '---':
            return

        # Bullet list
        lines = block.split('\n')
        if all(re.match(r'^\s*[-*]\s', l) for l in lines if l.strip()):
            for line in lines:
                if not line.strip():
                    continue
                item = re.sub(r'^\s*[-*]\s+', '', line)
                _add_rich_list_item(item, 'List Bullet')
            return

        # Numbered list
        if all(re.match(r'^\s*\d+\.\s', l) for l in lines if l.strip()):
            for line in lines:
                if not line.strip():
                    continue
                item = re.sub(r'^\s*\d+\.\s+', '', line)
                _add_rich_list_item(item, 'List Number')
            return

        # Regular paragraph with possible inline math
        _add_rich_paragraph(block)

    def _add_code_block(raw):
        """Render a fenced code block as a shaded single-cell table (like HTML <pre>)."""
        from docx.oxml.ns import qn as _qn
        from docx.shared import Emu

        # Strip fences
        code = re.sub(r'^```\w*\n?', '', raw)
        code = re.sub(r'\n?```$', '', code)
        lines = code.split('\n')

        # Use a 1-cell table with shading to mimic <pre> background
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]

        # Set cell shading to light grey (#F0F0F0)
        tc_pr = cell._element.get_or_add_tcPr()
        shading = __import__('lxml').etree.SubElement(
            tc_pr, _qn('w:shd'))
        shading.set(_qn('w:val'), 'clear')
        shading.set(_qn('w:color'), 'auto')
        shading.set(_qn('w:fill'), 'F0F0F0')

        # Remove default empty paragraph
        for existing_p in cell.paragraphs:
            existing_p._element.getparent().remove(existing_p._element)

        # Add each line as its own paragraph with monospace font
        for line in lines:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(line if line else ' ')  # preserve blank lines
            run.font.name = 'Consolas'
            run.font.size = Pt(8.5)

        doc.add_paragraph()  # spacing after code block

    def _add_display_equation(latex):
        """Add a centered display equation as a native Word OMML math object."""
        eq_counter[0] += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        omml = _latex_to_omml(latex)
        if omml is not None:
            p._element.append(omml)
        else:
            # Fallback: italic text
            run = p.add_run(latex)
            run.italic = True
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)

        # Equation number
        run = p.add_run(f'    ({eq_counter[0]})')
        run.font.size = Pt(10)

    # ----- Core rich-text helper: renders $math$, **bold**, *italic*, `code` -----

    def _render_rich_into(p, text, bold_all=False):
        """Render text with inline $math$, **bold**, *italic*, `code` into paragraph p.

        If bold_all is True, all non-math runs are additionally bolded (for header rows).
        """
        text = re.sub(r'\n(?!\n)', ' ', text)  # join continuation lines

        # Split by inline math $...$, keeping delimiters
        parts = re.split(r'(?<!\$)(\$(?!\$).+?\$(?!\$))', text)

        for part in parts:
            if not part:
                continue

            # Inline math
            if part.startswith('$') and part.endswith('$') and not part.startswith('$$'):
                latex = part[1:-1]
                omml = _latex_to_omml(latex)
                if omml is not None:
                    p._element.append(omml)
                else:
                    run = p.add_run(latex)
                    run.italic = True
                    run.font.name = 'Cambria Math'
                continue

            # Formatted text runs: **bold**, *italic*, `code`, plain
            fmt_pattern = r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)'
            segments = re.split(fmt_pattern, part)

            for seg in segments:
                if not seg:
                    continue
                if seg.startswith('**') and seg.endswith('**'):
                    run = p.add_run(seg[2:-2])
                    run.bold = True
                elif seg.startswith('*') and seg.endswith('*'):
                    run = p.add_run(seg[1:-1])
                    run.italic = True
                elif seg.startswith('`') and seg.endswith('`'):
                    run = p.add_run(seg[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    if bold_all:
                        run.bold = True
                else:
                    run = p.add_run(seg)
                    if bold_all:
                        run.bold = True

    def _add_rich_paragraph(text):
        """Add a new paragraph with rich rendering (math, bold, italic, code)."""
        text = re.sub(r'\n(?!\n)', ' ', text)
        p = doc.add_paragraph()
        _render_rich_into(p, text)

    def _add_rich_list_item(text, style='List Bullet'):
        """Add a list item with rich rendering (math, bold, italic, code)."""
        p = doc.add_paragraph(style=style)
        _render_rich_into(p, text)

    def _add_table(block):
        """Parse a markdown table and add it to the Word doc with rich cell content."""
        # Filter out separator rows: |:---|:---|---:| etc.
        lines = [l for l in block.strip().split('\n')
                 if l.strip() and not re.match(r'^\|(?:[\s:\-]+\|)+$', l)]
        if not lines:
            return

        rows_data = []
        for line in lines:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows_data.append(cells)

        if len(rows_data) < 1:
            return

        ncols = max(len(r) for r in rows_data)
        nrows = len(rows_data)

        table = doc.add_table(rows=nrows, cols=ncols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                if j < ncols:
                    cell = table.rows[i].cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    _render_rich_into(p, cell_text, bold_all=(i == 0))
                    # Set cell font size
                    for run in p.runs:
                        run.font.size = Pt(9)

                    # Header row: light blue background
                    if i == 0:
                        tc_pr = cell._element.get_or_add_tcPr()
                        shading = _etree.SubElement(tc_pr, _qn('w:shd'))
                        shading.set(_qn('w:val'), 'clear')
                        shading.set(_qn('w:color'), 'auto')
                        shading.set(_qn('w:fill'), 'D6E4F0')

        doc.add_paragraph()  # spacing after table

    # Process sections
    skip_titles = {"highlights", "abstract", "keywords", "keyword"}
    highlights_content = None
    abstract_content = None
    keywords_content = None

    for sec in sections:
        tl = sec["title"].lower()
        if tl == "highlights":
            highlights_content = sec["content"]
        elif tl == "abstract":
            abstract_content = sec["content"]
        elif tl in ("keywords", "keyword"):
            keywords_content = sec["content"]

    # Add highlights
    if highlights_content:
        doc.add_heading("Highlights", level=1)
        for line in highlights_content.split('\n'):
            line = line.strip()
            if line.startswith('-') or line.startswith('*'):
                item = re.sub(r'^[-*]\s+', '', line)
                doc.add_paragraph(item, style='List Bullet')

    # Add abstract
    if abstract_content:
        doc.add_heading("Abstract", level=1)
        _add_rich_paragraph(abstract_content)

    # Add keywords
    if keywords_content:
        p = doc.add_paragraph()
        run = p.add_run("Keywords: ")
        run.bold = True
        p.add_run(keywords_content.strip())

    doc.add_page_break()

    # Add body sections
    for sec in sections:
        if sec["title"].lower() in skip_titles or sec["level"] == 1:
            continue

        heading_level = min(sec["level"], 3)
        doc.add_heading(sec["title"], level=heading_level)

        if sec["content"]:
            _add_content(sec["content"])

    # Add References section
    if _word_cited_keys and _word_bib_entries:
        doc.add_page_break()
        doc.add_heading("References", level=1)

        if _word_citation_style == "authoryear":
            # Sort alphabetically by author then year
            def _sort_key(key):
                fields = _word_bib_entries.get(key, {})
                author = fields.get("author", "ZZZ")
                year = fields.get("year", "9999")
                return (author.lower(), year)
            sorted_keys = sorted(_word_cited_keys, key=_sort_key)
        else:
            sorted_keys = _word_cited_keys

        for idx, key in enumerate(sorted_keys, 1):
            fields = _word_bib_entries.get(key, {})
            if not fields:
                p = doc.add_paragraph(f"[{key}] — not in refs.bib")
                continue

            author = fields.get("author", "Unknown").replace("{", "").replace("}", "")
            title = fields.get("title", "Untitled").replace("{", "").replace("}", "")
            journal = fields.get("journal", fields.get("booktitle", "")).replace("{", "").replace("}", "")
            year = fields.get("year", "n.d.")
            volume = fields.get("volume", "")
            pages = fields.get("pages", "")
            doi = fields.get("doi", "")

            p = doc.add_paragraph()
            # Hanging indent: left=0.5in, first_line=-0.5in
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after = Pt(4)

            # Author (bold)
            run = p.add_run(f"{author}, ")
            run.bold = True
            run.font.size = Pt(10)
            # Year
            run = p.add_run(f"{year}. ")
            run.font.size = Pt(10)
            # Title
            run = p.add_run(f"{title}. ")
            run.font.size = Pt(10)
            # Journal (italic)
            if journal:
                run = p.add_run(f"{journal}")
                run.italic = True
                run.font.size = Pt(10)
                if volume:
                    run = p.add_run(f" {volume}")
                    run.font.size = Pt(10)
                    if pages:
                        run = p.add_run(f", {pages}")
                        run.font.size = Pt(10)
                run = p.add_run(". ")
                run.font.size = Pt(10)
            # DOI
            if doi:
                run = p.add_run(f"https://doi.org/{doi}")
                run.font.size = Pt(10)

    # Add figures
    if figures:
        doc.add_page_break()
        doc.add_heading("Figures", level=1)
        for i, fp in enumerate(figures, 1):
            cap = captions.get(fp.name, "")
            try:
                doc.add_picture(str(fp), width=Inches(6.0))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                p = doc.add_paragraph(f"[Figure {i}: {fp.name} — {e}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Fig. {i}. ")
            run.bold = True
            run.font.size = Pt(10)
            run2 = p.add_run(cap)
            run2.font.size = Pt(10)
            run2.italic = True
            doc.add_paragraph()  # spacing

    # Save
    out_dir = paper_dir / "submission"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_file = out_dir / "paper.docx"
    doc.save(str(out_file))
    return out_file


# ---------------------------------------------------------------------------
# LaTeX renderer (delegates to paper_renderer.py)
# ---------------------------------------------------------------------------

def render_latex(paper_dir):
    """Render paper.md to LaTeX using paper_renderer.py if available."""
    paper_dir = Path(paper_dir)
    tools_dir = Path(__file__).resolve().parent
    sys.path.insert(0, str(tools_dir))

    try:
        from paper_renderer import render_latex as _render_latex, load_journal_profile
        # Try to find a journal profile
        journals_dir = tools_dir.parent / "journals"
        if journals_dir.exists():
            profiles = list(journals_dir.glob("*.yaml"))
            if profiles:
                try:
                    profile = load_journal_profile(profiles[0].stem, str(journals_dir))
                    return _render_latex(str(paper_dir), profile)
                except Exception as e:
                    print(f"  [WARN] LaTeX render failed: {e}")
    except ImportError:
        pass

    print("  [SKIP] LaTeX renderer not available (missing paper_renderer.py or PyYAML)")
    return None


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    paper_dir = Path(sys.argv[1])
    if not (paper_dir / "paper.md").exists():
        print(f"Error: {paper_dir / 'paper.md'} not found")
        sys.exit(1)

    print(f"Rendering: {paper_dir.name}")
    print("=" * 50)

    # HTML
    html_file = render_html(paper_dir)
    print(f"  [OK] HTML:  {html_file}  ({html_file.stat().st_size:,} bytes)")

    # Word
    docx_file = render_word(paper_dir)
    if docx_file:
        print(f"  [OK] Word:  {docx_file}  ({docx_file.stat().st_size:,} bytes)")

    # LaTeX
    tex_file = render_latex(paper_dir)
    if tex_file:
        print(f"  [OK] LaTeX: {tex_file}")

    print("=" * 50)
    print("Done.")


if __name__ == "__main__":
    main()
