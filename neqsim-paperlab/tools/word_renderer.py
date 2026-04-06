"""
word_renderer.py - Publication-quality Word document renderer for paperflow.

Combines the best of render_all.py (native OMML equations) and build_word.py
(booktabs tables, citation handling, journal-specific formatting).

Features:
  - Native Word equations via LaTeX -> MathML -> OMML pipeline
  - Booktabs-style tables (three-line: top/header/bottom, no vertical rules)
  - BibTeX citation resolution (\\cite{key} -> [N])
  - Journal profile integration (spacing, margins, fonts, required sections)
  - Embedded figures with numbered captions
  - Page numbers in footer
  - Rich inline formatting: bold, italic, code, inline math
  - Code blocks rendered as shaded table cells (Consolas)

Usage:
  from word_renderer import render_word_document
  render_word_document("papers/my_paper/", journal_profile=profile)

  Or via paperflow:
  python paperflow.py format papers/my_paper/ --journal computers_chem_eng
"""

import os
import re
import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    raise ImportError("python-docx required. Install with: pip install python-docx")

try:
    from lxml import etree
except ImportError:
    etree = None

try:
    import latex2mathml.converter
except ImportError:
    latex2mathml = None


# ═══════════════════════════════════════════════════════════════════════════
# OMML Equation Pipeline (LaTeX -> MathML -> OMML)
# ═══════════════════════════════════════════════════════════════════════════

_omml_transform = None
_omml_available = None


def _init_omml():
    """Lazily initialize the MML2OMML XSLT transform."""
    global _omml_transform, _omml_available
    if _omml_available is not None:
        return _omml_available

    if etree is None or latex2mathml is None:
        _omml_available = False
        return False

    xsl_candidates = [
        os.path.expandvars(r'%ProgramFiles%\Microsoft Office\root\Office16\MML2OMML.XSL'),
        os.path.expandvars(r'%ProgramFiles(x86)%\Microsoft Office\root\Office16\MML2OMML.XSL'),
        os.path.expandvars(r'%ProgramFiles%\Microsoft Office\Office16\MML2OMML.XSL'),
        os.path.expandvars(r'%ProgramFiles%\Microsoft Office\root\Office15\MML2OMML.XSL'),
    ]
    for path in xsl_candidates:
        if os.path.exists(path):
            xslt_doc = etree.parse(path)
            _omml_transform = etree.XSLT(xslt_doc)
            _omml_available = True
            return True

    _omml_available = False
    return False


def latex_to_omml(latex_str):
    """Convert LaTeX math string to OMML XML element for Word.

    Always returns an ``<m:oMath>`` element suitable for inline use.
    The XSLT may wrap it in ``<m:oMathPara>``; this function unwraps it.
    Returns None if the pipeline is unavailable or conversion fails.
    """
    if not _init_omml():
        return None
    try:
        mathml = latex2mathml.converter.convert(latex_str)
        tree = etree.fromstring(mathml.encode('utf-8'))
        omml = _omml_transform(tree)
        root = omml.getroot()
        # Always return <m:oMath> so it works inline inside <w:p>.
        # The XSLT sometimes produces <m:oMathPara> wrapping — unwrap it.
        ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        if root.tag == '{%s}oMathPara' % ns:
            oMath = root.find('{%s}oMath' % ns)
            if oMath is not None:
                return oMath
        return root
    except Exception:
        return None


# ═══════════════════════════════════════════════════════════════════════════
# Unicode Fallback for Equations (when OMML not available)
# ═══════════════════════════════════════════════════════════════════════════

_GREEK = {
    r'\alpha': '\u03b1', r'\beta': '\u03b2', r'\gamma': '\u03b3',
    r'\delta': '\u03b4', r'\epsilon': '\u03b5', r'\zeta': '\u03b6',
    r'\eta': '\u03b7', r'\theta': '\u03b8', r'\kappa': '\u03ba',
    r'\lambda': '\u03bb', r'\mu': '\u03bc', r'\nu': '\u03bd',
    r'\xi': '\u03be', r'\pi': '\u03c0', r'\rho': '\u03c1',
    r'\sigma': '\u03c3', r'\tau': '\u03c4', r'\phi': '\u03c6',
    r'\chi': '\u03c7', r'\psi': '\u03c8', r'\omega': '\u03c9',
    r'\Delta': '\u0394', r'\Sigma': '\u03a3', r'\Omega': '\u03a9',
    r'\Lambda': '\u039b', r'\Gamma': '\u0393', r'\Phi': '\u03a6',
}

_OPERATORS = {
    r'\times': '\u00d7', r'\cdot': '\u00b7', r'\leq': '\u2264',
    r'\geq': '\u2265', r'\neq': '\u2260', r'\approx': '\u2248',
    r'\sim': '\u223c', r'\infty': '\u221e', r'\partial': '\u2202',
    r'\nabla': '\u2207', r'\pm': '\u00b1', r'\to': '\u2192',
    r'\rightarrow': '\u2192', r'\leftarrow': '\u2190',
    r'\sum': '\u2211', r'\prod': '\u220f', r'\in': '\u2208',
    r'\ll': '\u226a', r'\gg': '\u226b',
}


def latex_to_unicode(text):
    """Convert LaTeX math to readable Unicode text (fallback when OMML unavailable)."""
    for cmd, char in _GREEK.items():
        text = text.replace(cmd, char)
    for cmd, char in _OPERATORS.items():
        text = text.replace(cmd, char)

    # Clean LaTeX commands
    text = re.sub(r'\\text\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\hat\{([^}]*)\}', lambda m: m.group(1) + '\u0302', text)
    text = re.sub(r'\\tilde\{([^}]*)\}', lambda m: m.group(1) + '\u0303', text)
    text = re.sub(r'\\bar\{([^}]*)\}', lambda m: m.group(1) + '\u0304', text)
    text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', text)
    text = re.sub(r'\^{([^}]*)}', r'^\1', text)
    text = re.sub(r'_{([^}]*)}', r'_\1', text)
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    text = text.replace('{', '').replace('}', '')
    return text.strip()


# ═══════════════════════════════════════════════════════════════════════════
# Citation Handling
# ═══════════════════════════════════════════════════════════════════════════

def build_citation_map(refs_bib_path):
    """Parse refs.bib and assign sequential citation numbers (alphabetical order)."""
    if not Path(refs_bib_path).exists():
        return {}
    text = Path(refs_bib_path).read_text(encoding="utf-8")
    keys = re.findall(r'@\w+\{(\w+),', text)
    keys_sorted = sorted(set(keys), key=str.lower)
    return {k: i + 1 for i, k in enumerate(keys_sorted)}


def resolve_citations(text, cite_map):
    """Replace \\cite{key1, key2} with [N1, N2]."""
    def _repl(m):
        keys = [k.strip() for k in m.group(1).split(",")]
        nums = []
        for k in keys:
            if k in cite_map:
                nums.append(str(cite_map[k]))
            else:
                nums.append(f"?{k}")
        return "[" + ", ".join(sorted(nums, key=lambda x: int(x) if x.isdigit() else 999)) + "]"
    return re.sub(r'\\cite\{([^}]+)\}', _repl, text)


def _extract_comment(text, key):
    """Extract a metadata value from an HTML comment like <!-- Key: value -->.

    Returns None if the key is not found.
    """
    m = re.search(r'<!--\s*' + re.escape(key) + r'\s*:\s*(.+?)\s*-->', text, re.IGNORECASE)
    return m.group(1).strip() if m else None


def clean_bibtex_latex(text):
    """Strip LaTeX markup from BibTeX field values to produce plain Unicode."""
    # Accented characters: {\o} -> ø, \"{u} -> ü, etc.
    _accent_map = {
        '`': '\u0300', "'": '\u0301', '^': '\u0302', '"': '\u0308',
        '~': '\u0303', '=': '\u0304', '.': '\u0307', 'u': '\u0306',
        'v': '\u030C', 'H': '\u030B', 'c': '\u0327', 'd': '\u0323',
        'b': '\u0332', 'r': '\u030A', 'k': '\u0328',
    }
    _special_map = {
        'o': 'ø', 'O': 'Ø', 'l': 'ł', 'L': 'Ł',
        'aa': 'å', 'AA': 'Å', 'ae': 'æ', 'AE': 'Æ',
        'oe': 'œ', 'OE': 'Œ', 'ss': 'ß', 'i': 'ı', 'j': 'ȷ',
    }
    # \"{u} or {\"{u}} or {\"u} patterns
    def _repl_accent(m):
        cmd = m.group(1)
        char = m.group(2)
        if cmd in _accent_map:
            import unicodedata
            return unicodedata.normalize('NFC', char + _accent_map[cmd])
        return char
    # Handle: \cmd{char}  e.g. \"{u}
    text = re.sub(r'\\([`\'^"~=.ubvHcdrk])\{(\w)\}', _repl_accent, text)
    # Handle: {\cmd char}  e.g. {\"u}
    text = re.sub(r'\{\\([`\'^"~=.ubvHcdrk])(\w)\}', _repl_accent, text)
    # Special letters: {\o} or \o{}
    for ltx, uni in _special_map.items():
        text = text.replace('{\\' + ltx + '}', uni)
        text = text.replace('\\' + ltx + '{}', uni)
        text = text.replace('\\' + ltx + ' ', uni + ' ')
    # Strip remaining protective braces: {CPA} -> CPA, {D}oes -> Does
    text = re.sub(r'\{([^{}]*)\}', r'\1', text)
    # Strip any remaining lone backslashes before letters (e.g. \& -> &)
    text = re.sub(r'\\([&%#_])', r'\1', text)
    return text


def _extract_brace_value(text, start):
    """Extract a brace-balanced value starting at the opening '{' at position *start*.

    Returns (value_content, end_pos) where end_pos is the index after the closing '}'.
    """
    depth = 0
    i = start
    while i < len(text):
        if text[i] == '{':
            depth += 1
        elif text[i] == '}':
            depth -= 1
            if depth == 0:
                return text[start + 1:i], i + 1
        i += 1
    # Unmatched — return to end
    return text[start + 1:], len(text)


def parse_bibtex_entries(refs_bib_path):
    """Parse refs.bib into a list of (key, type, fields) tuples."""
    if not Path(refs_bib_path).exists():
        return []
    text = Path(refs_bib_path).read_text(encoding="utf-8")
    entries = []
    for match in re.finditer(r'@(\w+)\{(\w+),\s*(.*?)\n\}', text, flags=re.DOTALL):
        body = match.group(3)
        fields = {}
        # Match field_name = { ... } with brace balancing
        for fm in re.finditer(r'(\w+)\s*=\s*\{', body):
            field_name = fm.group(1).lower()
            value, _ = _extract_brace_value(body, fm.end() - 1)
            fields[field_name] = clean_bibtex_latex(value.strip())
        entries.append((match.group(2), match.group(1), fields))
    entries.sort(key=lambda x: x[0].lower())
    return entries


# ═══════════════════════════════════════════════════════════════════════════
# Markdown Parser
# ═══════════════════════════════════════════════════════════════════════════

def parse_sections(text):
    """Parse markdown into list of {level, title, content} dicts."""
    sections = []
    current = None
    lines = []
    for line in text.split("\n"):
        m = re.match(r'^(#{1,3})\s+(.+)$', line)
        if m:
            if current is not None:
                sections.append({
                    "level": current["level"],
                    "title": current["title"],
                    "content": "\n".join(lines).strip(),
                })
            current = {"level": len(m.group(1)), "title": m.group(2).strip()}
            lines = []
        else:
            lines.append(line)
    if current is not None:
        sections.append({
            "level": current["level"],
            "title": current["title"],
            "content": "\n".join(lines).strip(),
        })
    return sections


def strip_comments(text):
    """Remove HTML comments."""
    return re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)


# ═══════════════════════════════════════════════════════════════════════════
# Word Document Builder
# ═══════════════════════════════════════════════════════════════════════════

class WordRenderer:
    """Renders paper.md to a publication-quality Word document.

    Supports journal-specific formatting via profile dicts.
    """

    # Default formatting (can be overridden by journal profile)
    DEFAULTS = {
        "font_body": "Times New Roman",
        "font_heading": "Arial",
        "font_math": "Cambria Math",
        "font_code": "Consolas",
        "font_size_body": 11,
        "font_size_heading1": 13,
        "font_size_heading2": 11,
        "font_size_table": 9,
        "font_size_caption": 9,
        "font_size_code": 8.5,
        "font_size_ref": 9,
        "line_spacing": 2.0,       # double
        "space_after_para": 0,
        "margin_inches": 1.0,
        "page_width": 8.5,
        "page_height": 11.0,
        "figure_width": 5.5,
        "heading_color": (0x15, 0x65, 0xC0),     # blue
        "table_header_fill": "E3F2FD",            # light blue
        "code_block_fill": "F0F0F0",              # light grey
    }

    def __init__(self, paper_dir, journal_profile=None):
        self.paper_dir = Path(paper_dir)
        self.figures_dir = self.paper_dir / "figures"
        self.refs_file = self.paper_dir / "refs.bib"
        self.profile = journal_profile or {}

        # Load configuration
        self.cfg = dict(self.DEFAULTS)
        if self.profile.get("double_spacing_required"):
            self.cfg["line_spacing"] = 2.0
        elif self.profile.get("line_spacing"):
            self.cfg["line_spacing"] = self.profile["line_spacing"]

        # Citations
        self.cite_map = build_citation_map(self.refs_file)

        # Counters
        self.eq_counter = 0
        self.fig_counter = 0
        self.table_counter = 0

        # OMML availability
        self.has_omml = _init_omml()

        # Document
        self.doc = None

    def get(self, key):
        """Get a config value."""
        return self.cfg[key]

    # ─── Document Setup ───────────────────────────────────────────────

    def _setup_document(self):
        """Create document with page layout, styles, and footer."""
        self.doc = Document()

        # Page layout
        section = self.doc.sections[0]
        m = self.get("margin_inches")
        section.page_width = Inches(self.get("page_width"))
        section.page_height = Inches(self.get("page_height"))
        section.top_margin = Inches(m)
        section.bottom_margin = Inches(m)
        section.left_margin = Inches(m)
        section.right_margin = Inches(m)

        # Default Normal style
        style = self.doc.styles['Normal']
        style.font.name = self.get("font_body")
        style.font.size = Pt(self.get("font_size_body"))
        pf = style.paragraph_format
        pf.space_after = Pt(self.get("space_after_para"))
        spacing = self.get("line_spacing")
        if spacing == 2.0:
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        elif spacing == 1.5:
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            pf.line_spacing = spacing

        # Heading styles
        for lvl in range(1, 4):
            try:
                hs = self.doc.styles[f'Heading {lvl}']
                hs.font.name = self.get("font_body")
                hs.font.color.rgb = RGBColor(0, 0, 0)
            except KeyError:
                pass

        # Footer with page numbers
        self._add_page_number_footer(section)

    def _add_page_number_footer(self, section):
        """Add centered page number field to footer."""
        if etree is None:
            return
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = fp.add_run()
        fld_begin = etree.SubElement(run._element, qn('w:fldChar'))
        fld_begin.set(qn('w:fldCharType'), 'begin')

        run2 = fp.add_run()
        instr = etree.SubElement(run2._element, qn('w:instrText'))
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '

        run3 = fp.add_run()
        fld_end = etree.SubElement(run3._element, qn('w:fldChar'))
        fld_end.set(qn('w:fldCharType'), 'end')

    # ─── Inline Text Rendering ────────────────────────────────────────

    @staticmethod
    def _safe_text(text):
        """Strip characters that are illegal in XML 1.0 (control chars except tab/LF/CR)."""
        return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    # Tokenizer regex: matches tokens in priority order.
    # Inline math $...$ is matched FIRST so that * inside math (e.g. K^*)
    # is never consumed by the bold/italic patterns.
    _TOKEN_RE = re.compile(
        r'(\$(?!\$).+?\$(?!\$))'    # group 1: inline math  $...$
        r'|(\*\*.+?\*\*)'           # group 2: bold         **...**
        r'|(\*(?!\*).+?\*(?!\*))'   # group 3: italic       *...*
        r'|(`[^`]+`)'               # group 4: code         `...`
    )

    def render_rich_text(self, paragraph, text, bold_all=False):
        """Render text with inline $math$, **bold**, *italic*, `code` into paragraph.

        Uses a single-pass tokenizer so that ``*`` inside ``$math$`` is never
        misinterpreted as markdown bold or italic.
        """
        # Resolve citations first
        text = resolve_citations(text, self.cite_map)
        # Join continuation lines (single newline → space)
        text = re.sub(r'\n(?!\n)', ' ', text)

        pos = 0
        for m in self._TOKEN_RE.finditer(text):
            # Plain text before this token
            if m.start() > pos:
                plain = self._safe_text(text[pos:m.start()])
                if plain:
                    run = paragraph.add_run(plain)
                    if bold_all:
                        run.bold = True

            if m.group(1):       # inline math $...$
                latex = m.group(1)[1:-1]
                omml = latex_to_omml(latex) if self.has_omml else None
                if omml is not None:
                    paragraph._element.append(omml)
                else:
                    math_text = self._safe_text(latex_to_unicode(latex))
                    run = paragraph.add_run(math_text)
                    run.italic = True
                    run.font.name = self.get("font_math")
            elif m.group(2):     # bold **...**
                run = paragraph.add_run(self._safe_text(m.group(2)[2:-2]))
                run.bold = True
            elif m.group(3):     # italic *...*
                run = paragraph.add_run(self._safe_text(m.group(3)[1:-1]))
                run.italic = True
            elif m.group(4):     # code `...`
                run = paragraph.add_run(self._safe_text(m.group(4)[1:-1]))
                run.font.name = self.get("font_code")
                run.font.size = Pt(self.get("font_size_body") - 1)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
                if bold_all:
                    run.bold = True

            pos = m.end()

        # Remaining plain text after last token
        if pos < len(text):
            plain = self._safe_text(text[pos:])
            if plain:
                run = paragraph.add_run(plain)
                if bold_all:
                    run.bold = True

    # ─── Display Equations ────────────────────────────────────────────

    def add_display_equation(self, latex_str):
        """Add a numbered display equation (centered, with OMML or fallback)."""
        self.eq_counter += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        omml = latex_to_omml(latex_str) if self.has_omml else None
        if omml is not None:
            p._element.append(omml)
        else:
            cleaned = self._safe_text(latex_to_unicode(latex_str))
            run = p.add_run(cleaned)
            run.italic = True
            run.font.name = self.get("font_math")
            run.font.size = Pt(self.get("font_size_body"))

        # Equation number
        run = p.add_run(f'    ({self.eq_counter})')
        run.font.size = Pt(10)
        run.font.name = self.get("font_body")

    # ─── Tables (Booktabs Style) ─────────────────────────────────────

    def add_table(self, md_lines):
        """Add a booktabs-style table from markdown pipe-table lines."""
        # Parse rows, skip separator
        def parse_row(line):
            cells = [c.strip() for c in line.split('|')]
            return [c for c in cells if c]

        header_line = md_lines[0]
        data_lines = md_lines[2:] if len(md_lines) > 2 else []

        headers = parse_row(header_line)
        rows = [parse_row(line) for line in data_lines]
        n_cols = len(headers)
        n_rows = len(rows)

        table = self.doc.add_table(rows=1 + n_rows, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        font_sz = Pt(self.get("font_size_table"))

        # Header row
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            self.render_rich_text(p, h, bold_all=True)
            for run in p.runs:
                run.font.size = font_sz
            self._set_cell_shading(cell, self.get("table_header_fill"))
            self._set_cell_borders(cell, top=12, bottom=8)

        # Data rows
        for i, row_cells in enumerate(rows):
            for j, cell_text in enumerate(row_cells):
                if j >= n_cols:
                    break
                cell = table.cell(i + 1, j)
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                self.render_rich_text(p, cell_text)
                for run in p.runs:
                    run.font.size = font_sz
                if i == n_rows - 1:
                    self._set_cell_borders(cell, bottom=12)

        # Set booktabs-style table borders (no vertical rules)
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            f'<w:tblPr {nsdecls("w")}></w:tblPr>'
        )
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        self.doc.add_paragraph()  # spacer

    def _set_cell_shading(self, cell, color_hex):
        """Set background shading on a table cell."""
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_cell_borders(self, cell, top=None, bottom=None):
        """Set borders on a table cell (width in 1/8 pt)."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        for side, spec in [('top', top), ('bottom', bottom)]:
            if spec:
                border = parse_xml(
                    f'<w:{side} {nsdecls("w")} w:val="single" '
                    f'w:sz="{spec}" w:space="0" w:color="000000"/>'
                )
                tcBorders.append(border)
        tcPr.append(tcBorders)

    # ─── Code Blocks ─────────────────────────────────────────────────

    def add_code_block(self, raw_text):
        """Render a fenced code block as a shaded single-cell table."""
        code = re.sub(r'^```\w*\n?', '', raw_text)
        code = re.sub(r'\n?```$', '', code)
        lines = code.split('\n')

        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        self._set_cell_shading(cell, self.get("code_block_fill"))

        # Remove default paragraph
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)

        for line in lines:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(self._safe_text(line) if line else ' ')
            run.font.name = self.get("font_code")
            run.font.size = Pt(self.get("font_size_code"))

        self.doc.add_paragraph()

    # ─── Figures ─────────────────────────────────────────────────────

    def add_figure(self, img_path, caption_text):
        """Add a figure with numbered caption."""
        self.fig_counter += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)

        img = Path(img_path)
        if img.exists():
            run = p.add_run()
            run.add_picture(str(img), width=Inches(self.get("figure_width")))
        else:
            run = p.add_run(f"[Missing figure: {img_path}]")
            run.font.color.rgb = RGBColor(0xFF, 0, 0)

        # Strip any existing "Fig. N." or "Figure N." prefix from caption text
        # to avoid duplication with the auto-generated prefix below.
        caption_text = re.sub(
            r'^(?:Fig\.|Figure)\s*\d+\s*[.:]\s*', '', caption_text).strip()

        # Caption paragraph
        cap_p = self.doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(12)

        run = cap_p.add_run(f"Figure {self.fig_counter}. ")
        run.bold = True
        run.font.size = Pt(self.get("font_size_caption"))
        run.font.name = self.get("font_heading")

        self.render_rich_text(cap_p, caption_text)
        for r in cap_p.runs[1:]:
            r.font.size = Pt(self.get("font_size_caption"))
            if r.font.name is None:
                r.font.name = self.get("font_heading")

    # ─── References ──────────────────────────────────────────────────

    def add_reference_list(self, md_content=None):
        """Generate formatted reference list.

        If md_content contains pre-numbered references (e.g. [1] Author...),
        use those directly to preserve the paper.md ordering. Otherwise
        fall back to generating from refs.bib (alphabetical order).
        """
        # Check if paper.md already has numbered references
        if md_content:
            ref_lines = [l.strip() for l in md_content.split('\n') if l.strip()]
            numbered = [l for l in ref_lines if re.match(r'^\[\d+\]', l)]
            if len(numbered) >= 3:  # at least 3 numbered refs → use paper.md order
                for ref_text in numbered:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    p.paragraph_format.first_line_indent = Inches(-0.4)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    self.render_rich_text(p, ref_text)
                    for r in p.runs:
                        r.font.size = Pt(self.get("font_size_ref"))
                        if r.font.name is None:
                            r.font.name = self.get("font_body")
                return

        # Fallback: generate from refs.bib (alphabetical order)
        entries = parse_bibtex_entries(self.refs_file)
        for idx, (key, etype, fields) in enumerate(entries):
            num = idx + 1
            author = fields.get('author', 'Unknown')
            title = fields.get('title', '')
            year = fields.get('year', '')
            journal = fields.get('journal', '')
            volume = fields.get('volume', '')
            pages = fields.get('pages', '')
            publisher = fields.get('publisher', '')
            booktitle = fields.get('booktitle', '')

            ref_parts = [f'[{num}]  {author}']
            if year:
                ref_parts[-1] += f' ({year}).'
            if title:
                ref_parts.append(f'  {title}.')
            if journal:
                ref_parts.append(f'  {journal}')
                if volume:
                    ref_parts[-1] += f', {volume}'
                if pages:
                    ref_parts[-1] += f', {pages}'
                ref_parts[-1] += '.'
            elif booktitle:
                ref_parts.append(f'  In: {booktitle}.')
            elif publisher:
                ref_parts.append(f'  {publisher}.')

            ref_text = ''.join(ref_parts)

            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(self._safe_text(ref_text))
            run.font.size = Pt(self.get("font_size_ref"))
            run.font.name = self.get("font_body")

    # ─── Section Heading Helpers ─────────────────────────────────────

    def add_title(self, title_text, md_content=""):
        """Add centered paper title, author, and affiliation.

        Author metadata is extracted from HTML comments in paper.md:
          <!-- Author: Name -->
          <!-- Affiliation: Dept, Institution, City, Country -->
          <!-- Email: user@example.com -->
        Falls back to placeholders if not found.
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = self.get("font_heading")
        r, g, b = self.get("heading_color")
        run.font.color.rgb = RGBColor(r, g, b)

        # Extract author metadata from HTML comments
        author = _extract_comment(md_content, "Author") or "[Author Name]"
        affiliation = _extract_comment(md_content, "Affiliation") or "[Department, Institution, City, Country]"
        email = _extract_comment(md_content, "Email")

        # Author name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(author)
        run.font.size = Pt(12)
        run.font.name = self.get("font_heading")

        # Affiliation
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2) if email else Pt(24)
        run = p.add_run(affiliation)
        run.font.size = Pt(10)
        run.font.name = self.get("font_heading")
        run.italic = True

        # Email (if present)
        if email:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(24)
            run = p.add_run(email)
            run.font.size = Pt(10)
            run.font.name = self.get("font_heading")
            run.italic = True

    def add_section_heading(self, text, level=2):
        """Add a section or subsection heading."""
        font_sz = self.get("font_size_heading1") if level <= 2 else self.get("font_size_heading2")
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
        p.paragraph_format.space_after = Pt(6 if level <= 2 else 4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_sz)
        run.font.name = self.get("font_heading")

    # ─── Content Block Processing ────────────────────────────────────

    def process_content(self, content_text):
        """Process a section's content: equations, tables, lists, code, paragraphs."""
        # Stash code blocks first
        code_blocks = {}
        code_idx = [0]

        def _stash(m):
            key = f'\x00CODE{code_idx[0]}\x00'
            code_blocks[key] = m.group(0)
            code_idx[0] += 1
            return key

        content_text = re.sub(r'```\w*\n.*?```', _stash, content_text, flags=re.DOTALL)

        # Split on display equations
        parts = re.split(r'(\$\$.*?\$\$)', content_text, flags=re.DOTALL)

        for part in parts:
            part = part.strip()
            if not part:
                continue

            if part.startswith('$$') and part.endswith('$$'):
                self.add_display_equation(part[2:-2].strip())
                continue

            # Split into paragraph blocks
            blocks = re.split(r'\n\n+', part)
            for block in blocks:
                block = block.strip()
                if not block:
                    continue
                if block in code_blocks:
                    self.add_code_block(code_blocks[block])
                else:
                    self._process_block(block)

    def _process_block(self, block):
        """Process a single block: table, list, image, or paragraph."""
        # Horizontal rule
        if block.strip() == '---':
            return

        # Table
        if block.startswith('|') and '|' in block[1:]:
            lines = [l for l in block.split('\n')
                     if l.strip() and not re.match(r'^\|[\s:\-]+\|$', l.strip())]
            if len(lines) >= 2:
                # Re-include separator for parsing
                all_lines = block.strip().split('\n')
                self.add_table(all_lines)
                return

        # Image — use .*? instead of [^\]]* to allow ] inside alt text (e.g. citations [10])
        img_match = re.match(r'^!\[(.*?)\]\(([^)]+)\)', block.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_rel = img_match.group(2)
            img_path = self.paper_dir / img_rel
            # Check if there's a following caption line
            remaining = block[img_match.end():].strip()
            caption = alt_text
            # Match *Figure N.* Caption text  (partial italic prefix)
            cap_match = re.match(r'^\*Figure\s+\d+[.:]\*\s*(.+)', remaining, re.DOTALL)
            if not cap_match:
                # Match *Figure N. Caption text*  (all italic)
                cap_match = re.match(r'^\*(?:Figure\s+\d+[.:]\s*)?(.+)\*$', remaining, re.DOTALL)
            if cap_match:
                caption = cap_match.group(1).strip()
            self.add_figure(str(img_path), caption)
            return

        # Bullet list
        lines = block.split('\n')
        if all(re.match(r'^\s*[-*]\s', l) for l in lines if l.strip()):
            for line in lines:
                if not line.strip():
                    continue
                item = re.sub(r'^\s*[-*]\s+', '', line)
                indent = len(re.match(r'^(\s*)', line).group(1))
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4 + indent * 0.15)
                p.paragraph_format.space_after = Pt(2)
                p.add_run('\u2022  ')
                self.render_rich_text(p, item)
                for r in p.runs:
                    if r.font.name is None:
                        r.font.name = self.get("font_body")
                    if r.font.size is None:
                        r.font.size = Pt(self.get("font_size_body"))
            return

        # Numbered list
        if all(re.match(r'^\s*\d+\.\s', l) for l in lines if l.strip()):
            for line in lines:
                if not line.strip():
                    continue
                m = re.match(r'^(\d+)\.\s+(.+)', line)
                if m:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(f'{m.group(1)}. ')
                    run.bold = True
                    self.render_rich_text(p, m.group(2))
                    for r in p.runs:
                        if r.font.name is None:
                            r.font.name = self.get("font_body")
                        if r.font.size is None:
                            r.font.size = Pt(self.get("font_size_body"))
            return

        # Table caption (bold **Table N:** lines handled inline)
        if re.match(r'^\*\*Table\s+\d+', block):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            self.render_rich_text(p, block)
            for r in p.runs:
                r.font.size = Pt(self.get("font_size_caption"))
                r.font.name = self.get("font_heading")
            return

        # Standalone italic figure caption line (skip - handled by image parser)
        stripped = block.strip()
        if stripped.startswith('*') and 'Figure' in stripped and re.match(r'^\*Figure\s+\d+', stripped):
            return

        # Regular paragraph
        p = self.doc.add_paragraph()
        self.render_rich_text(p, block)
        for r in p.runs:
            if r.font.name is None:
                r.font.name = self.get("font_body")
            if r.font.size is None:
                r.font.size = Pt(self.get("font_size_body"))

    # ─── Main Build Method ───────────────────────────────────────────

    def build(self):
        """Build the complete Word document from paper.md.

        Returns the Document object. Call .save() on it or use render_word_document().
        """
        raw_md = (self.paper_dir / "paper.md").read_text(encoding="utf-8")
        paper_md = strip_comments(raw_md)
        sections = parse_sections(paper_md)

        self._setup_document()

        # Extract title
        title = "Untitled"
        for sec in sections:
            if sec["level"] == 1:
                title = sec["title"]
                break
        self.add_title(title, md_content=raw_md)

        # Categorize front matter sections
        skip_titles = {"highlights", "abstract", "keywords", "keyword", "references"}

        # Highlights
        for sec in sections:
            if sec["title"].lower() == "highlights":
                self.add_section_heading("Highlights", level=2)
                for line in sec["content"].split('\n'):
                    line = line.strip()
                    if line.startswith('-') or line.startswith('*'):
                        item = re.sub(r'^[-*]\s+', '', line)
                        p = self.doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.3)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        run = p.add_run('\u2022  ')
                        run.font.size = Pt(10)
                        self.render_rich_text(p, item)
                        for r in p.runs:
                            r.font.size = Pt(10)
                            if r.font.name is None:
                                r.font.name = self.get("font_heading")
                self.doc.add_paragraph()
                break

        # Abstract
        for sec in sections:
            if sec["title"].lower() == "abstract":
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                run = p.add_run("Abstract")
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = self.get("font_heading")

                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.right_indent = Inches(0.3)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                self.render_rich_text(p, sec["content"])
                for r in p.runs:
                    r.font.size = Pt(10)
                    if r.font.name is None:
                        r.font.name = self.get("font_body")
                self.doc.add_paragraph()
                break

        # Keywords
        for sec in sections:
            if sec["title"].lower() in ("keywords", "keyword"):
                content = sec["content"].strip()
                # Remove --- HR markers
                content = re.sub(r'^\s*---\s*$', '', content, flags=re.MULTILINE).strip()
                if content:
                    p = self.doc.add_paragraph()
                    run = p.add_run("Keywords: ")
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = self.get("font_heading")
                    run = p.add_run(content)
                    run.italic = True
                    run.font.size = Pt(10)
                    run.font.name = self.get("font_heading")
                self.doc.add_paragraph()
                self.doc.add_page_break()
                break

        # Body sections
        for sec in sections:
            if sec["title"].lower() in skip_titles or sec["level"] == 1:
                continue

            self.add_section_heading(sec["title"], level=sec["level"])

            if sec["content"]:
                self.process_content(sec["content"])

        # References
        for sec in sections:
            if sec["title"].lower() == "references":
                self.add_section_heading("References", level=2)
                md_content = sec.get("content", "")
                if self.refs_file.exists():
                    self.add_reference_list(md_content=md_content)
                elif md_content.strip():
                    self.add_reference_list(md_content=md_content)
                break

        return self.doc


# ═══════════════════════════════════════════════════════════════════════════
# Public API
# ═══════════════════════════════════════════════════════════════════════════

def render_word_document(paper_dir, journal_profile=None, output_dir=None):
    """Render paper.md to a Word document.

    Args:
        paper_dir: Path to paper directory containing paper.md
        journal_profile: Optional journal profile dict (from YAML)
        output_dir: Output directory (default: paper_dir/submission/)

    Returns:
        Path to generated .docx file
    """
    paper_dir = Path(paper_dir)
    if output_dir is None:
        output_dir = paper_dir / "submission"
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    renderer = WordRenderer(paper_dir, journal_profile)

    print(f"  Paper dir: {paper_dir}")
    print(f"  Figures:   {len(list(renderer.figures_dir.glob('*.png')))} PNG files")
    print(f"  Citations: {len(renderer.cite_map)} references")
    if renderer.has_omml:
        print("  Equations: Native OMML (LaTeX -> MathML -> Office Math)")
    else:
        print("  Equations: Unicode fallback (install latex2mathml + lxml for native)")

    doc = renderer.build()

    out_path = output_dir / "paper.docx"
    try:
        doc.save(str(out_path))
    except PermissionError:
        from datetime import datetime
        ts = datetime.now().strftime("%H%M%S")
        out_path = output_dir / f"paper_{ts}.docx"
        doc.save(str(out_path))
        print(f"  NOTE: paper.docx was locked, saved as {out_path.name}")

    print(f"  Output: {out_path}")
    print(f"  Size:   {out_path.stat().st_size / 1024:.0f} KB")

    # Post-render validation
    issues = validate_word_output(out_path, paper_dir, journal_profile)
    if issues:
        print("\n  ⚠ POST-RENDER VALIDATION ISSUES:")
        for issue in issues:
            print(f"    [{issue['severity']}] {issue['message']}")
    else:
        print("  ✓ Post-render validation: all checks passed")

    return out_path


# ═══════════════════════════════════════════════════════════════════════════
# Post-Render Validation
# ═══════════════════════════════════════════════════════════════════════════

def validate_word_output(docx_path, paper_dir, journal_profile=None):
    """Validate the generated Word document against the source paper.md.

    Catches common rendering bugs:
    - Missing figures (images in paper.md not embedded in docx)
    - Reference count mismatch
    - Duplicate figure captions
    - Missing equations (display $$ blocks)

    Returns a list of issue dicts with 'severity' and 'message'.
    """
    issues = []
    paper_dir = Path(paper_dir)
    docx_path = Path(docx_path)

    # Read source paper.md
    paper_md = (paper_dir / "paper.md").read_text(encoding="utf-8")

    # --- Check 1: Figure count ---
    # Count image references in paper.md
    md_images = re.findall(r'!\[.*?\]\([^)]+\)', paper_md)
    md_fig_count = len(md_images)

    # Count images embedded in the docx
    try:
        doc = Document(str(docx_path))
        docx_img_count = sum(
            1 for rel in doc.part.rels.values()
            if "image" in rel.reltype
        )
    except Exception:
        docx_img_count = -1

    if docx_img_count >= 0 and docx_img_count != md_fig_count:
        issues.append({
            "severity": "ERROR",
            "message": f"Figure count mismatch: paper.md has {md_fig_count} images, "
                       f"docx has {docx_img_count} embedded images",
        })
    elif docx_img_count >= 0:
        issues.append({
            "severity": "INFO",
            "message": f"Figures OK: {docx_img_count}/{md_fig_count} images embedded",
        })

    # --- Check 2: Reference count ---
    # Count refs in paper.md (numbered [N] pattern)
    md_refs = re.findall(r'^\s*\[\d+\]', paper_md, flags=re.MULTILINE)
    refs_bib = paper_dir / "refs.bib"
    if refs_bib.exists():
        bib_text = refs_bib.read_text(encoding="utf-8")
        bib_entries = re.findall(r'@\w+\{', bib_text)
        bib_count = len(bib_entries)
    else:
        bib_count = 0

    expected_refs = len(md_refs) if md_refs else bib_count
    if expected_refs > 0:
        # Count [N] patterns in docx text
        docx_text = "\n".join(p.text for p in doc.paragraphs)
        docx_ref_lines = re.findall(r'^\s*\[\d+\]', docx_text, flags=re.MULTILINE)
        if len(docx_ref_lines) != expected_refs:
            issues.append({
                "severity": "WARN",
                "message": f"Reference count: expected {expected_refs}, "
                           f"found {len(docx_ref_lines)} in docx",
            })

    # --- Check 3: Duplicate figure captions ---
    fig_caps = re.findall(r'Figure\s+(\d+)\.', docx_text)
    seen = set()
    for num in fig_caps:
        if num in seen:
            issues.append({
                "severity": "ERROR",
                "message": f"Duplicate figure caption: 'Figure {num}.' appears more than once",
            })
        seen.add(num)

    # --- Check 4: Display equations ---
    md_display_eqs = len(re.findall(r'\$\$.*?\$\$', paper_md, flags=re.DOTALL))
    if md_display_eqs > 0:
        # Check for OMML elements in docx XML
        omml_count = 0
        for p in doc.paragraphs:
            for child in p._element:
                if child.tag.endswith('}oMath') or child.tag.endswith('}oMathPara'):
                    omml_count += 1
        # Also count Unicode equation paragraphs (fallback rendering)
        unicode_eq_count = sum(
            1 for p in doc.paragraphs
            if any(c in p.text for c in '\u03b1\u03b2\u03b3\u03b4\u2211\u220f')
            and len(p.text) < 200
        )
        total_eqs = omml_count + unicode_eq_count
        if total_eqs == 0:
            issues.append({
                "severity": "WARN",
                "message": f"paper.md has {md_display_eqs} display equations "
                           f"but no OMML/equation elements detected in docx",
            })

    # Filter out INFO-level for return (keep only WARN and ERROR)
    return [i for i in issues if i["severity"] != "INFO"]


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print("Usage: python word_renderer.py <paper_dir> [--journal <name>]")
        sys.exit(1)

    pdir = sys.argv[1]
    profile = None

    if '--journal' in sys.argv:
        idx = sys.argv.index('--journal')
        if idx + 1 < len(sys.argv):
            journals_dir = Path(__file__).parent.parent / "journals"
            try:
                import yaml
                ppath = journals_dir / f"{sys.argv[idx + 1]}.yaml"
                with open(ppath) as f:
                    profile = yaml.safe_load(f)
            except Exception as e:
                print(f"Warning: Could not load journal profile: {e}")

    render_word_document(pdir, journal_profile=profile)
